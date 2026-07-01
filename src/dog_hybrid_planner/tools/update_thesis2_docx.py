#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Update thesis2.docx with structure text, architecture figures and analysis.

This script edits a working copy and writes a new docx. It intentionally keeps
the existing manuscript content and only inserts or lightly corrects targeted
paragraphs requested by the user.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


INPUT = Path("/tmp/论文2_work.docx")
OUTPUT = Path("/tmp/论文2_补充完成.docx")
TF_IMAGE = Path("/media/cjx/D683-C616/四足机器人路径规划TF框架图.png")
NODE_IMAGE = Path("/media/cjx/D683-C616/四足机器人路径规划ROS节点图.png")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def body_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_para = insert_paragraph_after(paragraph, text, "Normal")
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para.paragraph_format.first_line_indent = Pt(24)
    new_para.paragraph_format.line_spacing = 1.25
    return new_para


def heading_after(paragraph: Paragraph, text: str, style: str = "Heading 2") -> Paragraph:
    new_para = insert_paragraph_after(paragraph, text, style)
    return new_para


def picture_after(paragraph: Paragraph, image_path: Path, caption: str) -> Paragraph:
    image_para = insert_paragraph_after(paragraph, "", "Normal")
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(image_path), width=Inches(5.75))

    caption_para = insert_paragraph_after(image_para, caption, "Normal")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.first_line_indent = Pt(0)
    caption_para.paragraph_format.line_spacing = 1.25
    return caption_para


def find_para(doc: Document, startswith: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip().startswith(startswith):
            return para
    raise RuntimeError(f"Paragraph not found: {startswith}")


def find_para_contains(doc: Document, text: str) -> Paragraph:
    for para in doc.paragraphs:
        if text in para.text:
            return para
    raise RuntimeError(f"Paragraph not found: {text}")


def replace_whole_paragraph(doc: Document, old_start: str, new_text: str) -> None:
    para = find_para(doc, old_start)
    para.text = new_text


def replace_whole_paragraph_contains(doc: Document, old_text: str, new_text: str) -> None:
    para = find_para_contains(doc, old_text)
    para.text = new_text


def add_update_fields_setting(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def add_chapter_one_structure(doc: Document) -> None:
    anchor = find_para(doc, "现有文献对上述差异关注不足")
    p = heading_after(anchor, "1.4  论文主要内容及结构安排")
    p = body_after(
        p,
        "本文以ASK-3四足机器人为研究对象，围绕室内复杂环境中的全局路径规划、局部实时避障、全局—局部融合以及仿真执行稳定性开展研究。参考已有室内机器人路径规划研究中“改进A*算法—改进DWA算法—混合路径规划算法—ROS/Gazebo仿真实验”的技术路线，本文并不直接沿用轮式机器人模型下的路径评价标准，而是进一步将四足机器人的机体尺寸、步态摆动空间、头部朝前运动习惯、侧移能力边界和低层速度响应延迟等因素纳入规划系统中。全文研究内容既包括算法层面的改进，也包括ROS仿真系统中坐标变换、传感器、定位、路径显示和底层速度接口之间的工程闭环验证。章节内容安排如下：",
    )
    p = body_after(
        p,
        "第一章为绪论。该章首先阐述四足机器人路径规划研究的背景与意义，说明四足机器人相较于轮式机器人在复杂室内场景中的应用价值；随后从四足机器人研究现状、全局路径规划算法、局部路径规划算法以及混合路径规划算法等方面梳理国内外研究进展；最后分析现有路径规划方法直接应用于四足机器人时可能出现的贴墙、穿角、频繁转向、侧移滥用和速度衰减等问题，并明确本文的研究内容与章节结构。",
    )
    p = body_after(
        p,
        "第二章研究基于已有改进A*算法的全局路径规划方法。该章首先对基础A*算法与参考论文中的改进A*算法进行实验对比，说明双向搜索、动态权重、二十四邻域扩展和路径平滑对搜索效率与路径形态的改善作用；随后针对ASK-3四足机器人进一步引入足迹膨胀与安全距离软代价、线段安全检测、转角惩罚、安全剪枝、圆角平滑与均匀重采样等改进方法，并通过消融实验分析各改进点对路径安全性、平滑性和后续局部规划稳定性的影响。",
    )
    p = body_after(
        p,
        "第三章研究适用于四足机器人运动特性的改进DWA局部路径规划算法。传统DWA主要面向轮式机器人，速度空间通常只包含前进速度和偏航角速度，难以完整表达四足机器人侧移、前进和转向组合运动的能力。为此，本章将局部速度空间扩展为前进速度、侧向速度和偏航角速度组成的三维动态窗口，并在评价函数中加入前向姿态偏好与路径头部对齐机制，使机器人在空间充足时优先头部朝前沿路径运动，在局部避障必要时再使用侧向运动。",
    )
    p = body_after(
        p,
        "第四章研究A*算法与DWA算法的融合改进方法。单独的A*算法能够提供全局可达路径，但难以根据实时障碍物信息直接输出可执行速度；单独的DWA算法具有局部避障能力，但容易受到局部最优和目标跳变影响。本章将A*路径转化为DWA参考走廊，并设计局部目标单调推进机制，使DWA局部轨迹既能保持实时避障能力，又能持续沿全局路径方向推进，从而减少墙角处反复旋转、路径切换后原地犹豫以及两侧路线跳变等问题。",
    )
    body_after(
        p,
        "第五章进行ROS与Gazebo平台下的系统仿真实验。该章首先介绍ASK-3四足机器人仿真模型、激光雷达、AMCL定位、TF坐标变换关系、RViz可视化和底层步态控制节点之间的关系；随后补充速度保持与受限Recovery机制，解决多次重规划后速度下降以及恢复行为过度触发的问题；最后结合各项消融实验数据，对本文改进算法在导航时间、轨迹长度、平均速度、累计转角、侧向运动比例、Recovery次数和碰撞风险等指标上的效果进行综合分析。",
    )


def add_chapter_five_figures(doc: Document) -> None:
    anchor = find_para(doc, "在仿真流程中")
    p = body_after(
        anchor,
        "为了更清晰地说明本文算法在ROS仿真系统中的部署方式，本文进一步给出四足机器人路径规划系统的TF框架图和ROS节点图。与仅描述算法流程不同，TF框架图能够说明导航节点获取机器人全局位姿、雷达坐标以及机器人本体坐标之间关系的方式；ROS节点图则展示地图、传感器、定位、规划、可视化和底层步态控制之间的话题通信关系。",
    )
    p = picture_after(p, TF_IMAGE, "图5.1 四足机器人路径规划TF框架图")
    p = body_after(
        p,
        "如图5.1所示，本文仿真系统中的核心坐标链为map→odom→base→laser。AMCL节点根据/map静态栅格地图、/scan激光雷达数据和/odom里程计信息发布map→odom变换；mydog_state_estimator_ros_node根据Gazebo模型状态发布/odom并广播odom→base变换；robot_state_publisher根据ASK-3机器人URDF模型和/ask_3/joint_states关节状态发布base→laser以及四足腿部关节TF。dog_hybrid_navigator通过TF缓冲查询map→base和map→laser，将机器人位姿、激光点和全局路径统一到map坐标系下，从而为A*全局规划、DWA局部轨迹评价和RViz可视化提供一致的坐标基础。",
    )
    p = picture_after(p, NODE_IMAGE, "图5.2 四足机器人路径规划ROS节点图")
    body_after(
        p,
        "如图5.2所示，Gazebo负责加载ASK-3四足机器人模型、实验地图、关节控制器和仿真激光雷达；map_server发布静态地图，AMCL完成定位并输出map→odom，状态估计节点发布/odom和odom→base，robot_state_publisher负责机器人本体、雷达和腿部关节TF。dog_hybrid_navigator是本文路径规划系统的核心节点，接收/map、/scan、/odom、/move_base_simple/goal和TF信息，输出/dog_global_path、/dog_dwa_path以及关键点Marker用于RViz显示，同时通过/ask/dog/forward_back、/ask/dog/left_right、/ask/dog/yaw等话题向底层四足步态控制节点发送速度命令。该结构将路径规划算法与底层运动控制解耦，使本文改进算法能够在不改动dog_sim原始步态控制代码的前提下完成仿真验证。",
    )


def expand_analysis(doc: Document) -> None:
    inserts = [
        (
            "对数据分析可知",
            "进一步从表2.1中的数值可以看出，改进A*算法将规划时间由0.01842 s降低至0.00104 s，访问节点数由2203个减少至47个，说明双向搜索、动态权重和邻域扩展显著提高了搜索效率；同时累计转角由14.137 rad降至8.497 rad，路径点数由90个减少至46个，说明路径形态更加简洁。但改进后最小障碍距离仍为1.0 cell，并出现2条风险线段，表明该算法更偏向轮式机器人路径效率优化，对四足机器人所需的安全余量仍考虑不足。因此，本文后续并非否定参考论文的改进，而是在保留其高效搜索和平滑优势的基础上继续加入四足机器人约束。",
        ),
        (
            "由图2.2和表2.2可知",
            "从具体数据看，加入足迹膨胀与安全距离软代价后，路径长度由109.98 cell增加至115.45 cell，规划时间也由0.01607 s增加至0.05622 s，说明算法主动放弃了部分贴近障碍物的短路径；但最小障碍距离由1.0 cell提高至2.41 cell，路径点数由141个减少至68个，累计转角由7.576 rad降低至6.136 rad。该结果说明安全代价并不是简单让路径变长，而是在可接受的搜索代价内换取更大的通行余量和更稳定的参考路径。对于四足机器人而言，适当远离墙体能够减少腿部摆动碰撞风险，也能降低后续DWA为了避开障碍边缘而频繁修正速度的概率。",
        ),
        (
            "消融实验中，去除线段安全检测后",
            "从表2.3可以看到，去除线段安全检测时路径长度为74.49 cell，表面上更短，但最小障碍距离降至0.0 cell，风险线段数达到2条，说明路径存在从栅格障碍边角处“穿过”的隐患。加入线段安全检测后，路径长度增加到81.98 cell，规划时间由0.00111 s增加到0.02316 s，但风险线段数降为0，最小障碍距离恢复到1.0 cell。这说明线段安全检测的作用主要体现在连续空间安全性而不是离散节点可达性上。对于四足机器人来说，两个安全栅格之间的连线如果穿过障碍边缘，机器人机体和腿部仍可能发生碰撞，因此该改进能够有效弥补栅格离散化造成的安全判断缺陷。",
        ),
        (
            "根据图2.4与表2.4可知",
            "从表2.4的数值进一步分析，加入转角惩罚后，路径长度保持为118.49 cell，访问节点数也仅由142个小幅增加至149个，但累计转角由12.556 rad下降至8.384 rad，显著转弯数由26个减少至18个。这说明转角惩罚并未以明显增加路径长度为代价，而是通过搜索阶段的代价引导，使算法在多条距离接近的候选路径中优先选择方向变化更少的路线。对于ASK-3四足机器人而言，减少急转弯能够降低频繁偏航和步态切换造成的停顿，使DWA局部轨迹在跟踪全局路径时更容易保持连续前进。",
        ),
        (
            "消融实验中，去除该后处理后",
            "结合表2.5可以看出，安全剪枝、圆角平滑与重采样对路径形态改善尤为明显。加入后处理后，路径长度由119.32 cell缩短至115.45 cell，累计转角由13.172 rad降低至6.136 rad，显著转弯数由27个减少至11个，而规划时间基本保持在0.055 s左右。该结果说明后处理并非额外增加系统负担的装饰性步骤，而是能够在保持安全距离不变的条件下压缩冗余折线、缓和局部尖角，并将路径转换为更适合DWA跟踪的均匀参考序列。对于四足机器人而言，均匀路径点能够减少局部目标突然跳变，圆角化路径则有助于保持头部姿态连续变化。",
        ),
        (
            "由图3.1以及表3.1可知",
            "由表3.1可进一步看出，三维速度空间使导航时间由35.31 s缩短至17.22 s，轨迹长度由8.85 m缩短至7.40 m，平均速度由0.251 m/s提高至0.430 m/s，Recovery次数由2次降为0次。需要注意的是，最终系统的侧向运动比例由0.04提高至0.23，最小障碍距离由0.25 m降至0.104 m，这并不表示安全性下降，而是说明机器人在可控范围内利用侧移能力完成更高效的局部通过。若完全限制侧移，机器人只能依靠大幅转向绕行，反而会增加停顿和恢复动作。因此，三维速度空间的价值在于让DWA拥有表达四足机器人真实运动能力的候选动作集合。",
        ),
        (
            "根据上述数据，消融实验中",
            "从表3.2可以看到，加入前向姿态偏好与DWA路径头部对齐后，导航时间由29.93 s下降至18.28 s，平均速度由0.276 m/s提升至0.460 m/s，侧向运动比例由0.42降低至0.19，Recovery次数由1次降至0次。虽然轨迹长度由8.26 m小幅增加至8.41 m，最小障碍距离由0.304 m降至0.200 m，但整体通行效率明显提高。这说明该改进不是单纯追求距离最短或离墙最远，而是在安全可行前提下减少无意义侧移，使机器人更多以头部朝前的方式推进。对于四足机器人，这种运动姿态更符合步态执行习惯，也能减少在狭窄路口附近因横向移动过多造成的时间损失。",
        ),
        (
            "由图4.1与表4.1数据可知",
            "从表4.1进一步分析，加入A星参考走廊与局部目标单调推进后，导航时间由49.33 s降至18.35 s，轨迹长度由11.38 m缩短至8.08 m，累计转角由9.268 rad降至3.248 rad，显著转弯数由8个减少至4个，侧向运动比例由0.31下降至0.16，Recovery次数由3次降为0次，同时最小障碍距离由0.200 m提高至0.271 m。该组数据说明，混合规划性能下降的主要原因并不一定是A*路径本身或DWA采样能力不足，而可能是两者之间缺少稳定的信息传递。参考走廊和单调推进机制能够抑制局部目标在路径两侧反复切换，使DWA在局部避障时仍保持明确的全局方向。",
        ),
        (
            "由图5.3和表5.1可以看出",
            "从表5.1可进一步看出，加入速度保持与受限Recovery机制后，导航时间由48.58 s下降至18.96 s，轨迹长度由9.75 m缩短至7.45 m，平均速度由0.201 m/s提高至0.393 m/s，累计转角由14.33 rad降低至2.362 rad，显著转弯数由9个减少至3个，Recovery触发次数由4次降低至1次。与此同时，最小障碍距离由0.20 m提高至0.29 m，侧向运动比例由0.27下降至0.18。该结果说明，速度保持机制不仅提高了移动速度，还间接降低了因低速停滞引起的误恢复；受限Recovery机制则使恢复动作主要发生在真正无进展或前方阻塞时，而不是在DWA仍能正常避障的情况下频繁抢占控制。",
        ),
    ]

    for marker, text in inserts:
        anchor = find_para(doc, marker)
        body_after(anchor, text)


def correct_existing_text(doc: Document) -> None:
    replace_whole_paragraph_contains(
        doc,
        "表2.4 线段安全检测消融实验数据对比表",
        "表2.4 转角惩罚消融实验数据对比表",
    )
    replace_whole_paragraph_contains(
        doc,
        "速度保持与受限Recovery消融实验对比图",
        "图5.3 速度保持与受限Recovery消融实验对比图",
    )
    for para in doc.paragraphs:
        if "由图5.1和表5.1可以看出" in para.text:
            para.text = para.text.replace("由图5.1和表5.1可以看出", "由图5.3和表5.1可以看出")


def main() -> None:
    doc = Document(INPUT)
    add_chapter_one_structure(doc)
    add_chapter_five_figures(doc)
    correct_existing_text(doc)
    expand_analysis(doc)
    add_update_fields_setting(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
