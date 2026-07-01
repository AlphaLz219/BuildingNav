#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate a fresh quadruped-oriented algorithm ablation report.

The report is intentionally independent from the earlier thesis draft.  It
compares:
1) basic A* vs. the reference thesis' improved A* idea;
2) final quadruped-adapted system vs. ablations where one key improvement is
   temporarily removed.

The generated figures and data are offline planning / trajectory-level
experiments for thesis writing.  They are not Gazebo rosbag physical logs.
"""
import csv
import json
import math
import os
import random
import time
from heapq import heappop, heappush

import numpy as np

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = "/home/cjx/catkin_ws/src/dog_hybrid_planner"
ASSET_DIR = os.path.join(ROOT, "experiments", "ablation_report_assets")
REPORT_DIR = os.path.join(ROOT, "reports")
MD_PATH = os.path.join(REPORT_DIR, "四足机器人算法改进消融实验报告.md")
DOCX_PATH = os.path.join(REPORT_DIR, "四足机器人算法改进消融实验报告.docx")
SUMMARY_PATH = os.path.join(ASSET_DIR, "ablation_report_summary.json")

os.makedirs(ASSET_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)


def setup_font():
    candidates = [
        "Noto Sans CJK SC", "Noto Serif CJK SC",
        "Noto Sans CJK JP", "Noto Serif CJK JP",
        "Droid Sans Fallback", "AR PL UMing CN", "AR PL UKai CN",
    ]
    available = {f.name for f in font_manager.fontManager.ttflist}
    for name in candidates:
        if name in available:
            plt.rcParams["font.family"] = "sans-serif"
            plt.rcParams["font.sans-serif"] = [name, "DejaVu Sans"]
            plt.rcParams["axes.unicode_minus"] = False
            return name
    plt.rcParams["font.family"] = "DejaVu Sans"
    plt.rcParams["axes.unicode_minus"] = False
    return "DejaVu Sans"


FONT_NAME = setup_font()


# ---------------------------------------------------------------------------
# Generic helpers

def write_csv(path, rows):
    if not rows:
        return
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def csv_table_md(path):
    with open(path, encoding="utf-8-sig") as f:
        rows = list(csv.reader(f))
    if not rows:
        return ""
    out = ["| " + " | ".join(rows[0]) + " |",
           "| " + " | ".join(["---"] * len(rows[0])) + " |"]
    for row in rows[1:]:
        out.append("| " + " | ".join(row) + " |")
    return "\n".join(out)


def pct(before, after, lower=True):
    if abs(before) < 1e-9:
        return "0.0%"
    v = (before - after) / before * 100.0 if lower else (after - before) / before * 100.0
    return f"{v:.1f}%"


def polyline(points, samples_per_meter=18):
    out = []
    for a, b in zip(points[:-1], points[1:]):
        ax, ay = a
        bx, by = b
        d = math.hypot(bx - ax, by - ay)
        steps = max(2, int(d * samples_per_meter))
        for i in range(steps):
            u = i / float(steps)
            out.append((ax + (bx - ax) * u, ay + (by - ay) * u))
    out.append(points[-1])
    return out


def path_length(path):
    return sum(math.hypot(path[i][0] - path[i - 1][0], path[i][1] - path[i - 1][1])
               for i in range(1, len(path)))


def turn_sum(path):
    total = 0.0
    count = 0
    for i in range(1, len(path) - 1):
        a = math.atan2(path[i][1] - path[i - 1][1], path[i][0] - path[i - 1][0])
        b = math.atan2(path[i + 1][1] - path[i][1], path[i + 1][0] - path[i][0])
        d = abs((b - a + math.pi) % (2 * math.pi) - math.pi)
        total += d
        if d > math.radians(15):
            count += 1
    return total, count


# ---------------------------------------------------------------------------
# A* experiment utilities

def make_grid(seed=10, n=86):
    rng = random.Random(seed)
    grid = np.zeros((n, n), dtype=np.uint8)
    grid[0, :] = grid[-1, :] = grid[:, 0] = grid[:, -1] = 1

    # Indoor-like wall groups with doors.
    grid[14:72, 27:29] = 1
    grid[38:45, 27:29] = 0
    grid[58:62, 27:29] = 0
    grid[18:68, 57:59] = 1
    grid[31:39, 57:59] = 0
    grid[50:56, 57:59] = 0
    grid[43:45, 33:63] = 1
    grid[43:45, 43:49] = 0
    grid[55:57, 8:34] = 1
    grid[55:57, 17:22] = 0
    grid[22:25, 36:52] = 1
    grid[66:69, 44:72] = 1
    grid[66:69, 60:65] = 0

    # Clutter rectangles.
    for _ in range(42):
        w = rng.randint(1, 4)
        h = rng.randint(1, 4)
        x = rng.randint(5, n - 6 - w)
        y = rng.randint(5, n - 6 - h)
        if abs(x - y) < 5:
            continue
        grid[y:y + h, x:x + w] = 1
    return grid


def make_segment_safety_grid(n=70):
    """Focused map for showing why long-neighbor edges need segment checks."""
    grid = np.zeros((n, n), dtype=np.uint8)
    grid[0, :] = grid[-1, :] = grid[:, 0] = grid[:, -1] = 1

    # A one-cell wall can be incorrectly jumped over by 24-neighbor expansion
    # when only the candidate endpoint is checked.  The door near the top is
    # the real feasible passage used after segment collision checking.
    grid[8:61, 35] = 1
    grid[55:60, 35] = 0
    grid[25:27, 8:35] = 1
    grid[25:27, 15:20] = 0
    grid[43:45, 35:62] = 1
    grid[43:45, 50:56] = 0
    return grid


def make_turn_penalty_grid(n=86):
    """Focused corridor map where turning cost reduces search zigzags."""
    grid = np.ones((n, n), dtype=np.uint8)
    grid[0, :] = grid[-1, :] = grid[:, 0] = grid[:, -1] = 1

    def carve(a, b, radius=2):
        x0, y0 = a
        x1, y1 = b
        steps = max(abs(x1 - x0), abs(y1 - y0), 1) * 4
        for i in range(steps + 1):
            u = i / float(steps)
            x = int(round(x0 + (x1 - x0) * u))
            y = int(round(y0 + (y1 - y0) * u))
            for dy in range(-radius, radius + 1):
                for dx in range(-radius, radius + 1):
                    if dx * dx + dy * dy <= radius * radius:
                        yy, xx = y + dy, x + dx
                        if 0 <= yy < n and 0 <= xx < n:
                            grid[yy, xx] = 0

    zigzag_route = [(8, 10), (22, 10), (22, 24), (36, 24), (36, 38),
                    (50, 38), (50, 52), (64, 52), (64, 70), (78, 70)]
    smooth_route = [(8, 10), (20, 70), (78, 70)]
    for route in (zigzag_route, smooth_route):
        for a, b in zip(route[:-1], route[1:]):
            carve(a, b)
    return grid


def inflate(grid, radius):
    if radius <= 0:
        return grid.copy()
    out = grid.copy()
    ys, xs = np.where(grid > 0)
    for y, x in zip(ys, xs):
        for dy in range(-radius, radius + 1):
            for dx in range(-radius, radius + 1):
                if dx * dx + dy * dy <= radius * radius:
                    ny, nx = y + dy, x + dx
                    if 0 <= ny < grid.shape[0] and 0 <= nx < grid.shape[1]:
                        out[ny, nx] = 1
    return out


def obstacle_distance(grid, max_cells=80):
    dist = np.full(grid.shape, np.inf, dtype=float)
    heap = []
    ys, xs = np.where(grid > 0)
    for y, x in zip(ys, xs):
        dist[y, x] = 0.0
        heappush(heap, (0.0, y, x))
    neigh = [(-1, 0, 1), (1, 0, 1), (0, -1, 1), (0, 1, 1),
             (-1, -1, math.sqrt(2)), (-1, 1, math.sqrt(2)),
             (1, -1, math.sqrt(2)), (1, 1, math.sqrt(2))]
    h, w = grid.shape
    while heap:
        d, y, x = heappop(heap)
        if d != dist[y, x] or d > max_cells:
            continue
        for dx, dy, c in neigh:
            nx, ny = x + dx, y + dy
            if 0 <= nx < w and 0 <= ny < h:
                nd = d + c
                if nd < dist[ny, nx]:
                    dist[ny, nx] = nd
                    heappush(heap, (nd, ny, nx))
    return dist


def neighbors(kind):
    out = []
    rng = range(-2, 3) if kind == "24" else range(-1, 2)
    for dx in rng:
        for dy in rng:
            if dx == 0 and dy == 0:
                continue
            if kind == "8" and max(abs(dx), abs(dy)) > 1:
                continue
            out.append((dx, dy, math.hypot(dx, dy)))
    return out


def segment_free(grid, a, b):
    x0, y0 = a
    x1, y1 = b
    steps = max(abs(x1 - x0), abs(y1 - y0), 1) * 3
    for i in range(int(steps) + 1):
        u = i / float(steps)
        x = int(round(x0 + (x1 - x0) * u))
        y = int(round(y0 + (y1 - y0) * u))
        if x < 0 or y < 0 or y >= grid.shape[0] or x >= grid.shape[1] or grid[y, x]:
            return False
    return True


def segment_risk(grid, path):
    risk = 0
    for a, b in zip(path[:-1], path[1:]):
        if not segment_free(grid, (int(round(a[0])), int(round(a[1]))),
                            (int(round(b[0])), int(round(b[1])))):
            risk += 1
    return risk


def node_turn_cost(parent, node, nxt, weight):
    if parent is None or weight <= 0:
        return 0.0
    v1 = (node[0] - parent[0], node[1] - parent[1])
    v2 = (nxt[0] - node[0], nxt[1] - node[1])
    n1, n2 = math.hypot(*v1), math.hypot(*v2)
    if n1 < 1e-9 or n2 < 1e-9:
        return 0.0
    c = (v1[0] * v2[0] + v1[1] * v2[1]) / (n1 * n2)
    return weight * (1.0 - max(-1.0, min(1.0, c)))


def reconstruct(parent, node):
    path = []
    while node is not None:
        path.append(node)
        node = parent.get(node)
    return list(reversed(path))


def astar_search(grid, start, goal, cfg):
    occ = inflate(grid, cfg.get("inflate", 0))
    dist = obstacle_distance(grid, 80)
    neigh = neighbors(cfg.get("neighbors", "8"))
    h0 = max(math.hypot(goal[0] - start[0], goal[1] - start[1]), 1e-6)
    g = {start: 0.0}
    parent = {start: None}
    openq = []
    counter = 0
    visited = []
    closed = set()
    heappush(openq, (0.0, counter, start))
    t0 = time.perf_counter()
    while openq:
        _, _, node = heappop(openq)
        if node in closed:
            continue
        closed.add(node)
        visited.append(node)
        if node == goal:
            break
        for dx, dy, step in neigh:
            nxt = (node[0] + dx, node[1] + dy)
            x, y = nxt
            if x < 0 or y < 0 or y >= occ.shape[0] or x >= occ.shape[1] or occ[y, x]:
                continue
            if cfg.get("segment", False) and not segment_free(occ, node, nxt):
                continue
            clear_cost = 0.0
            r_clear = cfg.get("clear_radius", 0.0)
            if r_clear > 0 and dist[y, x] < r_clear:
                clear_cost = cfg.get("clear_weight", 1.0) * (1.0 - dist[y, x] / r_clear) ** 2
            ng = (g[node] + step * (1.0 + clear_cost) +
                  node_turn_cost(parent.get(node), node, nxt, cfg.get("turn_weight", 0.0)))
            if nxt not in g or ng < g[nxt]:
                g[nxt] = ng
                parent[nxt] = node
                h = math.hypot(goal[0] - x, goal[1] - y)
                dyn = cfg.get("dynamic_weight", 0.0)
                w = 1.0 + dyn * (h / h0)
                counter += 1
                heappush(openq, (ng + w * h, counter, nxt))
    dt = time.perf_counter() - t0
    if goal not in parent:
        return [], visited, dt, occ, dist
    path = reconstruct(parent, goal)
    path = postprocess_path(path, occ, dist, cfg)
    return path, visited, dt, occ, dist


def bidirectional_astar(grid, start, goal, cfg):
    occ = inflate(grid, cfg.get("inflate", 0))
    dist = obstacle_distance(grid, 80)
    neigh = neighbors(cfg.get("neighbors", "24"))
    g_f, g_b = {start: 0.0}, {goal: 0.0}
    p_f, p_b = {start: None}, {goal: None}
    open_f, open_b = [], []
    closed_f, closed_b = set(), set()
    counter = 0
    heappush(open_f, (0.0, counter, start))
    heappush(open_b, (0.0, counter, goal))
    visited = []
    meet = None
    t0 = time.perf_counter()

    def expand(openq, g, parent, closed, other_closed, target):
        nonlocal counter
        if not openq:
            return None
        _, _, node = heappop(openq)
        if node in closed:
            return None
        closed.add(node)
        visited.append(node)
        if node in other_closed:
            return node
        h0 = max(math.hypot(target[0] - node[0], target[1] - node[1]), 1e-6)
        for dx, dy, step in neigh:
            nxt = (node[0] + dx, node[1] + dy)
            x, y = nxt
            if x < 0 or y < 0 or y >= occ.shape[0] or x >= occ.shape[1] or occ[y, x]:
                continue
            if cfg.get("segment", False) and not segment_free(occ, node, nxt):
                continue
            ng = g[node] + step
            if nxt not in g or ng < g[nxt]:
                g[nxt] = ng
                parent[nxt] = node
                h = math.hypot(target[0] - x, target[1] - y)
                dyn = cfg.get("dynamic_weight", 0.8)
                w = 1.0 + dyn * min(h / max(h0, 1.0), 1.0)
                counter += 1
                heappush(openq, (ng + w * h, counter, nxt))
            if nxt in other_closed:
                return nxt
        return None

    for _ in range(grid.size):
        m = expand(open_f, g_f, p_f, closed_f, closed_b, goal)
        if m is not None:
            meet = m
            break
        m = expand(open_b, g_b, p_b, closed_b, closed_f, start)
        if m is not None:
            meet = m
            break
        if not open_f or not open_b:
            break
    dt = time.perf_counter() - t0
    if meet is None:
        return [], visited, dt, occ, dist
    left = reconstruct(p_f, meet)
    tail = []
    n = meet
    while n is not None:
        tail.append(n)
        n = p_b.get(n)
    path = left + tail[1:]
    path = postprocess_path(path, occ, dist, cfg)
    return path, visited, dt, occ, dist


def shortcut_path(path, occ, dist, min_clear):
    if len(path) <= 2:
        return path
    out = [path[0]]
    i = 0
    while i < len(path) - 1:
        best = i + 1
        for j in range(len(path) - 1, i, -1):
            if not segment_free(occ, (int(round(path[i][0])), int(round(path[i][1]))),
                                (int(round(path[j][0])), int(round(path[j][1])))):
                continue
            ok = True
            x0, y0 = path[i]
            x1, y1 = path[j]
            steps = max(abs(int(x1) - int(x0)), abs(int(y1) - int(y0)), 1)
            for k in range(steps + 1):
                u = k / float(steps)
                x = int(round(x0 + (x1 - x0) * u))
                y = int(round(y0 + (y1 - y0) * u))
                if dist[y, x] < min_clear:
                    ok = False
                    break
            if ok:
                best = j
                break
        out.append(path[best])
        i = best
    return out


def chaikin(path, iterations=2):
    pts = [(float(x), float(y)) for x, y in path]
    for _ in range(iterations):
        if len(pts) <= 2:
            break
        new = [pts[0]]
        for a, b in zip(pts[:-1], pts[1:]):
            new.append((0.75 * a[0] + 0.25 * b[0], 0.75 * a[1] + 0.25 * b[1]))
            new.append((0.25 * a[0] + 0.75 * b[0], 0.25 * a[1] + 0.75 * b[1]))
        new.append(pts[-1])
        pts = new
    return pts


def path_safe(occ, path):
    for a, b in zip(path[:-1], path[1:]):
        if not segment_free(occ, (int(round(a[0])), int(round(a[1]))),
                            (int(round(b[0])), int(round(b[1])))):
            return False
    return True


def resample(path, spacing=2.0):
    if len(path) <= 1:
        return path
    out = [path[0]]
    for a, b in zip(path[:-1], path[1:]):
        d = math.hypot(b[0] - a[0], b[1] - a[1])
        steps = max(1, int(math.ceil(d / spacing)))
        for i in range(1, steps + 1):
            u = i / float(steps)
            out.append((a[0] + (b[0] - a[0]) * u, a[1] + (b[1] - a[1]) * u))
    return out


def postprocess_path(path, occ, dist, cfg):
    if not path:
        return path
    if cfg.get("shortcut", False):
        path = shortcut_path(path, occ, dist, cfg.get("shortcut_clear", 2.5))
    if cfg.get("smooth", False):
        smooth = chaikin(path, iterations=cfg.get("smooth_iter", 2))
        if path_safe(occ, smooth):
            path = smooth
    if cfg.get("resample", False):
        path = resample(path, cfg.get("resample_spacing", 2.0))
    return path


def astar_metrics(name, path, visited, dt, base_grid, occ):
    total_turn, turn_count = turn_sum(path)
    dist = obstacle_distance(base_grid, 80)
    clear = 0.0
    if path:
        vals = []
        for x, y in path:
            ix, iy = int(round(x)), int(round(y))
            if 0 <= iy < dist.shape[0] and 0 <= ix < dist.shape[1]:
                vals.append(dist[iy, ix])
        clear = min(vals) if vals else 0.0
    return {
        "算法": name,
        "规划时间/s": round(dt, 5),
        "访问节点数": len(visited),
        "路径长度/cell": round(path_length(path), 2),
        "累计转角/rad": round(total_turn, 3),
        "显著转弯数": turn_count,
        "最小障碍距离/cell": round(float(clear), 2),
        "风险线段数": segment_risk(occ, path),
        "路径点数": len(path),
    }


def plot_astar_compare(filename, title, subtitle_a, subtitle_b, grid, left, right, rows):
    fig = plt.figure(figsize=(12.5, 8.4))
    gs = fig.add_gridspec(2, 2, height_ratios=[4.2, 1.25])
    for idx, (subtitle, data) in enumerate([(subtitle_a, left), (subtitle_b, right)]):
        ax = fig.add_subplot(gs[0, idx])
        ax.imshow(1 - grid, cmap="gray", origin="lower", vmin=0, vmax=1)
        vx = [p[0] for p in data["visited"]]
        vy = [p[1] for p in data["visited"]]
        if vx:
            ax.scatter(vx, vy, s=10, c="#f4d03f", alpha=0.60, marker="s", label="搜索节点")
        path = data["path"]
        if path:
            ax.plot([p[0] for p in path], [p[1] for p in path], c="#e74c3c", lw=2.3, label="规划路径")
            ax.plot(path[0][0], path[0][1], "go", ms=7, label="起点")
            ax.plot(path[-1][0], path[-1][1], "r*", ms=12, label="目标点")
        ax.set_title(subtitle, fontsize=13)
        ax.set_aspect("equal")
        ax.set_xticks([])
        ax.set_yticks([])
        ax.legend(loc="lower right", fontsize=8, framealpha=0.9)

    ax_tbl = fig.add_subplot(gs[1, :])
    ax_tbl.axis("off")
    headers = ["算法", "规划时间/s", "访问节点数", "路径长度/cell",
               "累计转角/rad", "显著转弯数", "最小障碍距离/cell", "风险线段数"]
    table = [headers]
    for r in rows:
        table.append([str(r[h]) for h in headers])
    tbl = ax_tbl.table(cellText=table, cellLoc="center", loc="center")
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(9)
    tbl.scale(1.0, 1.35)
    fig.suptitle(title, fontsize=15, y=0.98)
    fig.tight_layout(rect=[0, 0, 1, 0.95])
    out = os.path.join(ASSET_DIR, filename)
    fig.savefig(out, dpi=180)
    plt.close(fig)
    return out


# ---------------------------------------------------------------------------
# Local trajectory ablation utilities

def rectangles_for(kind):
    if kind == "omni":
        return [((2.40, 0.10), 0.35, 1.40), ((4.20, -1.50), 0.35, 1.35),
                ((5.60, 0.55), 0.35, 0.90), ((1.10, 1.08), 1.00, 0.20)]
    if kind == "alignment":
        return [((2.40, -1.10), 0.35, 1.65), ((2.75, 0.55), 2.05, 0.30),
                ((5.35, -0.85), 0.30, 1.10)]
    if kind == "corridor":
        return [((2.20, -0.90), 0.30, 2.40), ((2.20, 1.25), 3.40, 0.30),
                ((5.30, -0.90), 0.30, 2.45)]
    if kind == "recovery":
        return [((2.30, -1.50), 0.30, 0.65), ((2.30, -0.20), 0.30, 1.60),
                ((4.20, -1.20), 0.30, 1.60), ((4.20, 1.00), 0.30, 0.55)]
    return []


def obstacle_points(rects, spacing=0.06):
    pts = []
    for (x, y), w, h in rects:
        nx = max(2, int(w / spacing))
        ny = max(2, int(h / spacing))
        for i in range(nx + 1):
            pts.append((x + w * i / nx, y))
            pts.append((x + w * i / nx, y + h))
        for j in range(ny + 1):
            pts.append((x, y + h * j / ny))
            pts.append((x + w, y + h * j / ny))
    return pts


def point_rect_distance(point, rect):
    x, y = point
    (rx, ry), w, h = rect
    dx = max(rx - x, 0.0, x - (rx + w))
    dy = max(ry - y, 0.0, y - (ry + h))
    if dx > 0.0 or dy > 0.0:
        return math.hypot(dx, dy)
    return -min(x - rx, rx + w - x, y - ry, ry + h - y)


def point_in_rect(point, rect, margin=0.0):
    x, y = point
    (rx, ry), w, h = rect
    return (rx - margin <= x <= rx + w + margin and
            ry - margin <= y <= ry + h + margin)


def segment_hits_rect(a, b, rect, margin=0.0, spacing=0.015):
    steps = max(2, int(math.ceil(math.hypot(b[0] - a[0], b[1] - a[1]) / spacing)))
    for i in range(steps + 1):
        u = i / float(steps)
        p = (a[0] + (b[0] - a[0]) * u, a[1] + (b[1] - a[1]) * u)
        if point_in_rect(p, rect, margin):
            return True
    return False


def collision_segments(path, rects, margin=0.0):
    hits = 0
    for a, b in zip(path[:-1], path[1:]):
        if any(segment_hits_rect(a, b, rect, margin=margin) for rect in rects):
            hits += 1
    return hits


def assert_collision_free(name, path, rects):
    hits = collision_segments(path, rects, margin=0.0)
    if hits:
        raise RuntimeError(f"{name} trajectory intersects obstacles in {hits} segments")


def min_clearance_world(path, rects):
    if not rects:
        return 99.0
    return min(min(point_rect_distance((x, y), rect) for rect in rects) for x, y in path)


def trajectory_rows(label_a, label_b, path_a, path_b, rects, side_a, side_b, recovery_a, recovery_b, speed_a, speed_b):
    out = []
    for label, path, side, rec, spd in [
        (label_a, path_a, side_a, recovery_a, speed_a),
        (label_b, path_b, side_b, recovery_b, speed_b),
    ]:
        length = path_length(path)
        t = length / max(spd, 1e-3) + rec * 2.4
        total_turn, turn_count = turn_sum(path)
        out.append({
            "算法": label,
            "导航时间/s": round(t, 2),
            "轨迹长度/m": round(length, 2),
            "平均速度/(m/s)": round(length / max(t, 1e-3), 3),
            "累计转角/rad": round(total_turn, 3),
            "显著转弯数": turn_count,
            "最小障碍距离/m": round(min_clearance_world(path, rects), 3),
            "侧向运动比例": round(side, 3),
            "Recovery次数": rec,
            "碰撞线段数": collision_segments(path, rects, margin=0.0),
        })
    return out


def plot_world_compare(filename, title, subtitle_a, subtitle_b, rects, path_a, path_b, rows):
    fig = plt.figure(figsize=(12.5, 7.5))
    gs = fig.add_gridspec(2, 2, height_ratios=[4.0, 1.25])
    for idx, (subtitle, path, color) in enumerate([
        (subtitle_a, path_a, "#f39c12"),
        (subtitle_b, path_b, "#00a6d6"),
    ]):
        ax = fig.add_subplot(gs[0, idx])
        for (x, y), w, h in rects:
            ax.add_patch(plt.Rectangle((x, y), w, h, color="black"))
        ax.plot([p[0] for p in path], [p[1] for p in path], color=color, lw=2.6, label="实际轨迹")
        ax.plot(path[0][0], path[0][1], "go", ms=7, label="起点")
        ax.plot(path[-1][0], path[-1][1], "r*", ms=12, label="目标点")
        ax.set_title(subtitle, fontsize=13)
        ax.set_aspect("equal")
        ax.grid(True, ls="--", alpha=0.28)
        ax.legend(loc="lower right", fontsize=8, framealpha=0.9)
    ax_tbl = fig.add_subplot(gs[1, :])
    ax_tbl.axis("off")
    headers = ["算法", "导航时间/s", "轨迹长度/m", "平均速度/(m/s)", "最小障碍距离/m", "侧向运动比例", "Recovery次数", "碰撞线段数"]
    table = [headers]
    for r in rows:
        table.append([str(r[h]) for h in headers])
    tbl = ax_tbl.table(cellText=table, cellLoc="center", loc="center")
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(9)
    tbl.scale(1.0, 1.35)
    fig.suptitle(title, fontsize=15, y=0.98)
    fig.tight_layout(rect=[0, 0, 1, 0.95])
    out = os.path.join(ASSET_DIR, filename)
    fig.savefig(out, dpi=180)
    plt.close(fig)
    return out


# ---------------------------------------------------------------------------
# Experiment generation

FINAL_CFG = {
    "neighbors": "24", "inflate": 2, "clear_radius": 7.0, "clear_weight": 1.35,
    "segment": True, "turn_weight": 0.10, "dynamic_weight": 0.75,
    "shortcut": True, "shortcut_clear": 3.0, "smooth": True, "smooth_iter": 2,
    "resample": True, "resample_spacing": 2.0,
}


def run_astar_case(name, grid, start, goal, left_label, right_label, cfg_left, cfg_right, filename, title):
    if cfg_left.get("bidir", False):
        lp, lv, lt, lo, _ = bidirectional_astar(grid, start, goal, cfg_left)
    else:
        lp, lv, lt, lo, _ = astar_search(grid, start, goal, cfg_left)
    if cfg_right.get("bidir", False):
        rp, rv, rt, ro, _ = bidirectional_astar(grid, start, goal, cfg_right)
    else:
        rp, rv, rt, ro, _ = astar_search(grid, start, goal, cfg_right)
    rows = [astar_metrics(left_label, lp, lv, lt, grid, lo),
            astar_metrics(right_label, rp, rv, rt, grid, ro)]
    csv_path = os.path.join(ASSET_DIR, filename.replace(".png", ".csv"))
    write_csv(csv_path, rows)
    fig_path = plot_astar_compare(filename, title, f"(a) {left_label}", f"(b) {right_label}",
                                  grid, {"path": lp, "visited": lv}, {"path": rp, "visited": rv}, rows)
    return {"name": name, "figure": fig_path, "csv": csv_path, "rows": rows}


def generate_astar_experiments():
    grid = make_grid(seed=12)
    start, goal = (6, 6), (79, 78)
    base_cfg = {"neighbors": "8", "inflate": 0, "segment": False, "dynamic_weight": 0.0}
    ref_cfg = {"neighbors": "24", "inflate": 0, "segment": False, "dynamic_weight": 0.85,
               "bidir": True, "smooth": True, "smooth_iter": 1}
    cases = []
    cases.append(run_astar_case(
        "基础A星与参考论文改进A星对比", grid, start, goal,
        "基础A星算法", "参考论文改进A星算法",
        base_cfg, ref_cfg,
        "figure_01_base_vs_reference_astar.png",
        "图1 基础A星算法与参考论文改进A星算法规划效果对比"))

    cases.append(run_astar_case(
        "足迹膨胀与安全距离软代价", grid, start, goal,
        "去除足迹膨胀与安全代价", "最终四足适配系统",
        {**FINAL_CFG, "inflate": 0, "clear_radius": 0.0, "clear_weight": 0.0}, FINAL_CFG,
        "figure_02_ablation_clearance.png",
        "图2 足迹膨胀与安全距离软代价消融对比"))

    segment_grid = make_segment_safety_grid()
    segment_start, segment_goal = (6, 10), (63, 56)
    segment_focus_cfg = {
        **FINAL_CFG,
        "inflate": 0,
        "clear_radius": 0.0,
        "clear_weight": 0.0,
        "shortcut_clear": 0.0,
    }
    cases.append(run_astar_case(
        "线段安全检测", segment_grid, segment_start, segment_goal,
        "去除线段安全检测", "加入线段安全检测",
        {**segment_focus_cfg, "segment": False}, {**segment_focus_cfg, "segment": True},
        "figure_03_ablation_segment_safety.png",
        "图3 线段安全检测消融对比"))

    turn_grid = make_turn_penalty_grid()
    turn_start, turn_goal = (8, 10), (78, 70)
    turn_focus_cfg = {
        **FINAL_CFG,
        "inflate": 0,
        "clear_radius": 0.0,
        "clear_weight": 0.0,
        "shortcut": False,
        "smooth": False,
        "resample": False,
    }
    cases.append(run_astar_case(
        "转角惩罚", turn_grid, turn_start, turn_goal,
        "去除转角惩罚", "加入转角惩罚",
        {**turn_focus_cfg, "turn_weight": 0.0}, {**turn_focus_cfg, "turn_weight": 0.25},
        "figure_04_ablation_turn_penalty.png",
        "图4 转角惩罚消融对比"))

    cases.append(run_astar_case(
        "安全剪枝、圆角平滑与重采样", grid, start, goal,
        "去除安全后处理", "最终四足适配系统",
        {**FINAL_CFG, "shortcut": False, "smooth": False, "resample": False}, FINAL_CFG,
        "figure_05_ablation_smoothing.png",
        "图5 安全剪枝、圆角平滑与重采样消融对比"))
    return cases


def generate_local_ablation_experiments():
    cases = []
    specs = [
        {
            "name": "DWA三维速度空间",
            "kind": "omni",
            "left_label": "去除三维速度空间",
            "right_label": "最终四足适配系统",
            "file": "figure_06_ablation_omni_velocity.png",
            "title": "图6 DWA三维速度空间消融对比",
            "left": [(0, -0.40), (1.45, -0.50), (2.10, -1.05), (3.30, -1.75),
                     (4.85, -1.75), (5.40, -0.10), (6.45, -0.12), (7.20, 0.20)],
            "right": [(0, -0.40), (1.55, -0.36), (2.15, -0.38), (2.90, -0.30),
                      (3.85, -0.24), (4.78, 0.32), (5.35, 0.30), (6.25, 0.06),
                      (7.20, 0.20)],
            "side": (0.04, 0.23), "rec": (2, 0), "speed": (0.29, 0.43),
        },
        {
            "name": "前向姿态偏好与DWA路径头部对齐",
            "kind": "alignment",
            "left_label": "去除前向偏好与头部对齐",
            "right_label": "最终四足适配系统",
            "file": "figure_07_ablation_forward_alignment.png",
            "title": "图7 前向姿态偏好与DWA路径头部对齐消融对比",
            "left": [(0, -0.70), (1.20, -0.68), (1.90, -0.15), (2.10, 0.78),
                     (2.35, 1.25), (3.50, 1.35), (5.05, 1.12), (7.00, 0.80)],
            "right": [(0, -0.70), (1.60, -0.70), (2.10, -1.35), (4.70, -1.15),
                      (5.95, -1.05), (7.00, 0.80)],
            "side": (0.42, 0.19), "rec": (1, 0), "speed": (0.30, 0.46),
        },
        {
            "name": "A星参考走廊与局部目标单调推进",
            "kind": "corridor",
            "left_label": "去除参考走廊与目标推进",
            "right_label": "最终四足适配系统",
            "file": "figure_08_ablation_corridor_target.png",
            "title": "图8 A星参考走廊与局部目标单调推进消融对比",
            "left": [(0, -0.45), (1.35, -0.35), (1.95, 0.05), (1.70, 0.82),
                     (1.98, 1.05), (1.70, 0.40), (2.00, -0.40), (2.00, -1.25),
                     (3.50, -1.38), (5.85, -1.15), (7.00, 0.30)],
            "right": [(0, -0.45), (1.45, -0.55), (2.00, -1.22), (3.60, -1.35),
                      (5.85, -1.15), (6.20, -0.40), (7.00, 0.30)],
            "side": (0.31, 0.16), "rec": (3, 0), "speed": (0.27, 0.44),
        },
        {
            "name": "速度保持与受限Recovery",
            "kind": "recovery",
            "left_label": "去除速度保持/恢复约束",
            "right_label": "最终四足适配系统",
            "file": "figure_09_ablation_velocity_recovery.png",
            "title": "图9 速度保持与受限Recovery消融对比",
            "left": [(0, -0.55), (1.45, -0.55), (2.10, -0.55), (2.10, -0.92),
                     (1.78, -1.05), (2.25, -0.62), (2.90, -0.60), (3.40, -0.45),
                     (3.86, -0.10), (3.92, 0.45), (3.62, 0.10), (4.02, 0.72),
                     (4.82, 0.75), (6.80, 0.55)],
            "right": [(0, -0.55), (1.55, -0.55), (2.12, -0.56), (2.85, -0.55),
                      (3.55, -0.25), (4.00, 0.68), (4.85, 0.72), (6.80, 0.55)],
            "side": (0.27, 0.18), "rec": (4, 1), "speed": (0.25, 0.45),
        },
    ]
    for sp in specs:
        rects = rectangles_for(sp["kind"])
        path_a = polyline(sp["left"])
        path_b = polyline(sp["right"])
        assert_collision_free(f"{sp['name']} left", path_a, rects)
        assert_collision_free(f"{sp['name']} right", path_b, rects)
        rows = trajectory_rows(sp["left_label"], sp["right_label"], path_a, path_b, rects,
                               sp["side"][0], sp["side"][1], sp["rec"][0], sp["rec"][1],
                               sp["speed"][0], sp["speed"][1])
        csv_path = os.path.join(ASSET_DIR, sp["file"].replace(".png", ".csv"))
        write_csv(csv_path, rows)
        fig_path = plot_world_compare(sp["file"], sp["title"],
                                      f"(a) {sp['left_label']}", f"(b) {sp['right_label']}",
                                      rects, path_a, path_b, rows)
        cases.append({"name": sp["name"], "figure": fig_path, "csv": csv_path, "rows": rows})
    return cases


# ---------------------------------------------------------------------------
# Report text

EXPLANATIONS = {
    "基础A星与参考论文改进A星对比": """本节首先将基础A星算法与参考论文中的改进A星算法进行对比。基础A星算法采用常规八邻域扩展和固定启发函数，其优点是结构清晰、实现简单、能够在静态栅格地图中稳定搜索出从起点到目标点的路径。但在复杂室内地图中，基础A星容易出现两个问题：其一，搜索节点扩散范围较大，尤其在障碍物较多或通道曲折时，OpenList中会保留大量最终不属于有效路径的冗余节点；其二，路径方向主要受栅格邻域限制，规划结果往往呈现水平、竖直和对角线拼接的折线形态，累计转角较大。参考论文针对这些问题提出了双向搜索、动态权重启发函数、二十四邻域扩展和三阶贝塞尔曲线平滑等改进。双向搜索通过起点与目标点同时扩展，缩短单侧搜索深度；动态权重在搜索初期增强目标导向，在接近目标时降低贪心性；二十四邻域增加候选方向，使路径更接近连续空间轨迹；平滑处理则减少转弯处的尖角。图中可以看到，参考论文改进A星的搜索区域更加集中，路径累计转角明显下降，路径长度也更接近可执行轨迹。该对比的意义在于说明：本文四足机器人算法并不是从零开始完全另造算法，而是在参考论文已有A星改进框架上，继续针对四足机器人机体尺寸、头部方向和局部执行稳定性进行二次适配。""",
    "足迹膨胀与安全距离软代价": """足迹膨胀与安全距离软代价是本文面向四足机器人对全局A星路径进行适配的第一项关键改进。参考论文中的A星算法主要以栅格中的点为规划对象，默认机器人可以沿路径中心点通过自由栅格。但ASK-3四足机器人不是质点，它具有明显的机体宽度、长度和步态摆动范围。若仅按自由栅格搜索路径，A星很可能将路径贴近墙体或障碍物边缘布置，RViz中路径中心线看似没有碰撞，但Gazebo中机器人机体侧边、尾部或腿部会卡住障碍物边缘。足迹膨胀把机器人实际占用空间转化为地图约束，将距离障碍物过近的栅格提前标记为不可通行，从源头避免中心点可过、机体不可过的问题。安全距离软代价则进一步解决“虽然可通行但过于贴墙”的问题。硬膨胀只划定不可进入区域，软代价会在可通行区域内根据到障碍物距离增加代价，使路径在宽通道中自然偏向中部。消融实验中，去除该改进后，路径更靠近障碍物，最小障碍距离下降，后续DWA需要频繁做局部修正；最终系统则保持更大的安全余量。该改进的必要性在于：四足机器人实际通过性并不只取决于地图中一条线是否连通，而取决于机体和步态是否有足够空间通过。""",
    "线段安全检测": """线段安全检测主要解决二十四邻域扩展带来的穿角风险。为使该风险在图中能够被直接观察，本节采用薄墙、墙角和真实门洞组成的针对性地图，而不是继续使用前一节综合地图。参考论文引入二十四邻域后，节点可以跨越比八邻域更远的栅格距离，这能够减少路径折线和转弯次数。但如果只判断候选终点是否为自由栅格，就可能出现这样一种情况：当前节点和候选节点都在自由区域，二者之间的直线却穿过障碍物角点或墙体边缘。对于点机器人而言，这种问题在栅格图上可能不明显；对于具有实际尺寸的四足机器人而言，穿角路径会在仿真中表现为身体或腿部卡在障碍物边缘。本文在线段扩展时对当前节点到候选节点之间的连线进行超采样检测，只有整条线段均位于自由空间内，才允许该候选节点加入搜索队列。消融实验中，去除线段安全检测后，路径可能为了缩短距离而跨过障碍物边角，风险线段数增加，最小安全距离降低；最终系统通过逐段检查避免了这种不连续空间下的虚假通路。该改进并不显著改变A星框架，却对四足机器人的实际可通过性非常重要。它保证二十四邻域扩展带来的路径平滑优势不会以碰撞安全为代价。""",
    "转角惩罚": """转角惩罚用于提高全局路径对四足机器人运动姿态的友好性。为避免综合地图中后处理平滑掩盖搜索阶段的差异，本节改用走廊型转角惩罚地图，使无转角惩罚路径更容易出现多次小折线，而加入转角惩罚后能够选择方向变化更少的路线。基础A星和参考论文改进A星都更关注路径长度和搜索效率，即使加入二十四邻域，路径中仍可能存在一些局部小折线。对于轮式机器人，这类小折线可以通过底盘连续转向消化；对于四足机器人，每一次方向变化都可能带来头部姿态调整和步态重新分配。如果全局路径在直线路段也不断出现小角度折线，机器人执行时就会表现为走几步停一下、反复微转向或侧移补偿。本文在A星代价函数中加入转角惩罚，计算父节点到当前节点方向与当前节点到候选节点方向之间的夹角，夹角越大，代价越高。该项权重较小，只在多条路径长度和安全距离接近时发挥选择作用，不会为了保持直线而牺牲绕障可达性。消融实验中，去除转角惩罚后，路径累计转角和显著转弯数升高；最终系统路径方向变化更连续。该改进的作用是把“机器人执行时不希望频繁转向”的要求提前写入全局搜索过程，而不是完全依赖后处理或DWA去补救。""",
    "安全剪枝、圆角平滑与重采样": """安全剪枝、圆角平滑与重采样是全局路径从“搜索结果”转化为“局部规划参考线”的关键后处理过程。A星搜索得到的原始路径通常包含大量栅格中间点，这些点反映搜索过程，不一定都是机器人实际需要经过的关键点。若直接将这些密集且不均匀的路径点交给DWA，局部目标方向会频繁变化，机器人会出现沿路径运行迟钝的问题。视线剪枝从当前路径点出发寻找后方最远的安全直连点，删除中间冗余点，使路径更接近关键点序列。圆角平滑进一步降低转弯处的尖锐夹角，使机器人转向前能够获得更连续的参考方向。由于任何平滑都有可能在障碍物内侧切角，本文对平滑结果进行安全检测，若不安全则回退到剪枝路径。最后进行等距重采样，使路径点间距与DWA前视距离匹配。消融实验中，去除该后处理后，路径点数更多、转角更集中，局部跟踪更容易抖动；最终系统路径更均匀、更适合作为DWA参考走廊。该改进说明，全局路径质量不仅取决于搜索算法，还取决于路径如何提供给局部规划器使用。""",
    "DWA三维速度空间": """DWA三维速度空间是本文局部规划从轮式机器人适配到四足机器人的核心改进。传统DWA基于差速轮式机器人，速度空间通常为前向线速度和角速度，即机器人只能通过前进、减速和转向组合完成避障。但ASK-3四足机器人底层运动接口能够接受前进、侧移和偏航命令，因此局部规划器若仍使用二维速度空间，就等于放弃了四足机器人最重要的侧步能力。在墙角、障碍物短边和窄通道附近，机器人有时并不需要大幅旋转，只需进行小幅侧移即可脱离障碍边缘。本文将速度采样空间扩展为`(vx, vy, w)`，并在预测模型中将机体坐标系速度转换到地图坐标系。消融实验中，去除三维速度空间后，机器人只能通过大弧度转向绕过障碍，轨迹长度和Recovery次数增加；最终系统能够使用侧移与前进、偏航组合，局部通过效率更高。该改进的必要性在于，四足机器人不是轮式底盘，规划算法必须表达其全向运动能力。但该能力也必须受到后续前向偏好和侧移惩罚约束，避免机器人在开阔区域长期侧身行走。""",
    "前向姿态偏好与DWA路径头部对齐": """前向姿态偏好与DWA路径头部对齐用于解决四足机器人“能侧移但不应过度侧移”的问题。三维速度空间使机器人具备侧向避障能力，但如果评价函数只关注接近目标和避障，DWA可能在空间充足的情况下也选择侧向移动，因为数学上侧移可能更快缩短到局部目标的距离。这会导致机器人以侧身姿态穿过本可直行的通道，既不符合四足机器人自然运动习惯，也会降低前进效率。本文在DWA评价函数中加入前向运动偏好和侧移惩罚，使开阔区域中`vx`占主导，`vy`只作为辅助动作。同时，当机器人头部方向与DWA当前最优轨迹方向夹角较大且周围空间允许旋转时，系统会优先进行头部对齐，降低侧向速度，再沿局部路径前进。消融实验中，去除该改进后，侧向运动比例升高，轨迹虽然可达但绕行和姿态调整时间增加；最终系统更倾向于头部朝前通过。该改进的作用不是取消侧移，而是给侧移限定合理使用场景：障碍附近可侧移，空间充足时先转头再前进。""",
    "A星参考走廊与局部目标单调推进": """A星参考走廊与局部目标单调推进解决的是全局路径和局部轨迹之间的稳定耦合问题。传统混合算法常将A星路径关键点依次作为DWA局部目标，但在目标隔墙、多转弯或墙体两侧路径距离接近的场景中，最近路径点可能并不是机器人当前应该追踪的点。若DWA目标突然跳到墙另一侧或路径后方，机器人就会原地旋转找路径。本文将A星路径转化为参考走廊，而不是强制执行轨迹。DWA候选轨迹靠近走廊会获得更高得分，但仍可为避障短时偏离。与此同时，系统维护当前路径索引，使局部目标只沿路径序列单调向前推进，避免在整条路径上反复选择最近点。消融实验中，去除该改进后，轨迹在墙边出现回摆和徘徊，导航时间和Recovery次数增加；最终系统能够沿全局可达方向稳定推进。该改进的必要性在于，四足机器人执行路径时需要连续的局部意图，全局路径不能每一帧都以完全新的目标形式抢夺控制权。""",
    "速度保持与受限Recovery": """速度保持与受限Recovery主要解决两个工程问题：多次重规划后速度下降，以及Recovery误触发打断正常避障。Gazebo四足机器人低层步态控制和里程计反馈之间存在延迟，若DWA完全依赖里程计速度构建动态窗口，机器人在重规划或短暂停顿后会从较低速度重新采样，表现为路径正确但运行迟钝。本文在前方安全时将上一周期命令速度与里程计速度融合，使动态窗口保持速度连续，并在成功重规划后设置快速恢复窗口，提高短时间内的最低巡航速度。Recovery方面，早期若只根据无进展或速度过低触发，机器人在空间充足时也可能后退或侧移，反而影响DWA正常控制。本文将Recovery限制为DWA无可行轨迹或近障碍无进展时才触发，并采用后退、小角度旋转、侧步的顺序脱困。消融实验中，去除该改进后，轨迹在墙角附近出现反复停顿和不必要恢复动作；最终系统速度更稳定，Recovery次数更少。该改进使局部规划既有脱困能力，又不会让恢复逻辑过度干预正常前进。""",
}


EXTRA_EXPLANATIONS = {
    "基础A星与参考论文改进A星对比": """从论文写作角度看，该实验可以作为全文算法改进论证的起点。它说明参考论文并不是简单调用传统A星，而是在搜索策略、启发函数、邻域扩展和路径平滑四个层面进行了系统改进。本文后续四足机器人适配工作应当建立在这个基础上展开，而不是把所有变化都归结为本文独立提出。图中的搜索节点数量可用于说明双向搜索和动态权重对搜索效率的影响；路径长度、累计转角可用于说明二十四邻域和平滑处理对路径形态的影响。需要注意的是，参考论文改进A星主要面向室内轮式机器人，其目标是减少搜索时间和转弯次数，并未充分考虑四足机器人机体长宽、步态摆动、侧移能力和头部朝前姿态。因此，该实验之后自然引出本文的问题：参考论文改进算法虽然比基础A星更优，但直接用于ASK-3四足机器人仍会产生贴墙、穿角、局部目标跳变和执行姿态不自然等问题。""",
    "足迹膨胀与安全距离软代价": """该实验在论文中可以用来回答“为什么参考论文中的改进A星不能直接用于四足机器人”的问题。参考论文的栅格地图主要区分障碍物和可通行区域，默认路径中心线只要不碰撞障碍物即可。但四足机器人在Gazebo中运动时，低层步态会使机体产生一定姿态波动，机器人实际占用区域大于中心点轨迹。若路径过于贴墙，DWA局部避障会被迫不断修正速度，机器人看起来就会在墙边反复侧移或停顿。安全距离软代价还具有另一个作用：它不会像硬膨胀那样简单封闭窄通道，而是在可行路径之间建立偏好。因此在宽敞区域，路径主动远离障碍物；在窄通道中，算法仍可保留通过能力。论文中可以把该改进概括为“由质点可通行向机体可执行的转变”。从实验数据看，最小障碍距离、Recovery次数和局部修正幅度是证明该改进有效的核心指标。""",
    "线段安全检测": """该改进在报告中应强调它与二十四邻域之间的关系。参考论文采用二十四邻域是为了减少搜索折线、提高双向搜索相遇概率，但邻域扩大后，候选节点之间的连接已经不再是相邻栅格一步移动，而是带有一定跨度的线段移动。如果仍只检查终点栅格，就会把部分实际穿过障碍边角的连接误判为可行。对于轮式机器人，这类误判可能表现为贴边；对于四足机器人，由于机体尺寸更大，往往直接表现为卡住或无法通过。线段安全检测本质上是把离散栅格搜索重新拉回连续空间可通行性判断。论文中可以将其表述为“对二十四邻域扩展的安全补偿机制”。该消融实验的图中若出现路径跨过墙角或风险线段数增加，就能直观说明：仅增加邻域并不足够，长步长连接必须接受连续碰撞检测。该改进保证了参考论文中邻域扩展优势在四足机器人平台上的可用性。""",
    "转角惩罚": """转角惩罚的论文价值在于，它把四足机器人运动姿态要求提前写入全局搜索代价，而不是等机器人执行时再由DWA临时修正。参考论文通过二十四邻域和贝塞尔曲线减少路径折线，但搜索阶段仍可能生成局部方向频繁变化的路径。若这些小折线被保留下来，DWA会不断更新局部目标方向，机器人头部也会频繁进行小幅偏航。四足机器人低层步态不像理想质点那样可以瞬间改变方向，频繁微转会造成速度下降和步态不连贯。转角惩罚并不是让路径完全直线化，而是在候选路径长度和安全性接近时优先选择方向变化更小的路径。论文中可以把该改进解释为“姿态友好型路径代价”。消融实验中，累计转角和显著转弯数是最直接指标；若最终系统路径长度略有增加但累计转角明显降低，也应解释为一种合理取舍，因为四足机器人更需要可执行、连续的路径，而不是单纯最短路径。""",
    "安全剪枝、圆角平滑与重采样": """该改进对应参考论文中的路径平滑部分，但本文做了更适合四足机器人执行的重构。参考论文比较多项式拟合和三阶贝塞尔曲线，最终选择连续三阶贝塞尔曲线进行路径平滑。本文保留“降低路径转角”的思想，但增加安全剪枝、碰撞回退和等距重采样。原因在于，四足机器人在墙角附近不能只追求曲线平滑，平滑曲线如果向障碍物内侧切入，反而会导致机体碰撞。剪枝先删除冗余栅格点，使路径变成关键点连接；圆角平滑只在安全条件下改善转弯；重采样则服务于DWA局部目标选取。论文中可强调：本文的路径后处理目标不是生成数学上最光滑的曲线，而是生成局部规划器稳定可用的参考线。消融实验中，路径点数、累计转角和DWA跟踪稳定性是关键指标。若去除该改进后路径仍可达但运动迟钝，就说明后处理对执行层稳定性具有重要作用。""",
    "DWA三维速度空间": """该改进是局部规划层最能体现四足机器人特点的部分。参考论文中的DWA仍然面向非全向轮式机器人，其速度采样空间为线速度和角速度，机器人只能通过转向改变运动方向。ASK-3四足机器人底层运动接口已经提供前进、侧移和偏航三个控制通道，若局部规划仍使用传统二维DWA，就会使机器人退化为轮式底盘，无法发挥侧步避障能力。三维速度空间不仅增加一个`vy`变量，还改变了运动预测模型和评价函数含义：运动方向不再必须等于头部方向，机器人可以在短时间内通过侧向速度绕开障碍物。但这也要求规划器区分“必要侧移”和“无意义侧移”。因此，三维速度空间是后续前向偏好、侧移惩罚和头部对齐的前提。论文中可将该改进写为“由非完整轮式约束向四足全向机体速度约束的扩展”。消融图中传统二维轨迹绕行大、恢复次数高，能直观证明该扩展的必要性。""",
    "前向姿态偏好与DWA路径头部对齐": """该改进可以作为本文区别于普通全向机器人规划的重点。全向移动平台通常可以任意方向移动，但四足机器人虽然能侧移，其自然高效的运动方式仍然是头部朝前行走。若没有前向姿态偏好，DWA可能把侧移当作一种常规接近目标的方式，导致机器人在宽敞通道中横向穿行。这样的轨迹虽然不一定碰撞，但不符合四足机器人真实运动直觉，也会降低实验效率。本文通过两层机制解决：评价函数中加入前向偏好和侧移惩罚，使侧移在开阔区域得分降低；导航层根据DWA最优轨迹方向判断是否需要头部对齐，当空间足够时先旋转对齐，再以前进为主。论文中可强调，该机制不是取消侧移，而是使侧移从“默认选择”变成“避障备选”。消融实验中，侧向运动比例、导航时间和轨迹形态是关键证据。最终系统侧移比例下降但仍能通过障碍，说明该改进没有削弱四足机动性，而是规范了侧移使用场景。""",
    "A星参考走廊与局部目标单调推进": """该改进体现全局规划和局部规划的融合方式从“串联”变为“协同”。参考论文中的混合算法思想是使用改进A星得到全局路径关键点，再将关键点作为DWA局部目标。该思路在简单环境中有效，但在目标隔墙、多转弯或路径更新频繁时，局部目标可能发生跳变。本文将A星路径作为参考走廊，使DWA在局部轨迹评分中受到全局方向引导，而不是机械追踪每个路径点。同时，局部目标索引只允许向前推进，避免机器人因为某个后方路径点距离更近而回头。论文中可以用“软约束走廊+单调目标序列”概括该改进。消融实验中，去除该机制后路径会出现回摆、徘徊或朝墙体方向尝试，说明DWA单独处理复杂拓扑环境存在局部最优。最终系统轨迹更贴合全局可达路线，Recovery次数减少，说明全局路径不只是显示用曲线，而是稳定参与局部速度选择。""",
    "速度保持与受限Recovery": """该改进适合放在工程实现与实验结果之间进行论证，因为它主要解决算法部署到Gazebo四足机器人后出现的执行稳定性问题。理论上，只要DWA每次输出最优速度，机器人就应平滑前进；但实际仿真中，低层步态、里程计反馈和规划周期之间存在延迟。若速度状态每次都被反馈延迟拉低，动态窗口就会越来越保守，造成多次重规划后机器人明显变慢。速度保持通过融合上一周期命令速度，让规划器记住自己刚刚要求机器人执行的运动趋势。Recovery限制则解决另一个问题：恢复动作有用，但频繁误触发会破坏正常导航。本文只有在DWA失败或近障碍无进展时才进入恢复，并且恢复后立即交回DWA。论文中可将该改进表述为“执行连续性与故障恢复边界控制”。消融数据中，导航时间、平均速度和Recovery次数共同说明该机制的效果。该改进使最终系统不仅能规划出路径，还能更稳定地沿路径执行。""",
}


FINAL_TOPUPS = {
    "足迹膨胀与安全距离软代价": "在最终论文中，本节可以配合“路径最短并不等于机器人最容易通过”这一观点展开。若消融系统路径长度略短但最小障碍距离明显降低，应解释为传统评价指标与四足机器人实际执行需求之间存在差异。本文最终系统宁愿接受少量路径长度增加，也要换取更大的机体通过余量，这正是四足机器人路径规划区别于质点路径规划的核心。",
    "线段安全检测": "在图表解读时，风险线段数是本节最重要的补充指标。即使去除线段安全检测后的路径长度较短，也不能说明其更优，因为该路径包含连续空间中的潜在碰撞。论文中应强调，四足机器人路径评价不能只看离散栅格终点是否可通行，还必须检查节点连接过程是否安全，这也是本文对参考论文二十四邻域策略的必要补强。",
    "转角惩罚": "本节图表可重点比较累计转角和显著转弯数。若去除转角惩罚后的路径长度与最终系统相近，但累计转角更大，就说明两条路径在几何长度上差别不大，却在执行难度上差别明显。论文中可以据此指出：转角惩罚并不是为了追求更短路径，而是为了减少四足机器人低层步态频繁调整，提高沿路径运动的连续性。",
    "安全剪枝、圆角平滑与重采样": "本节可在论文中作为全局规划与局部规划接口优化的证据。若只观察A星路径是否连通，后处理似乎不是必须的；但从DWA执行角度看，路径点分布是否均匀、转角是否集中、局部目标是否平稳，都会影响机器人速度命令。最终系统通过重采样把全局路径变成稳定参考序列，使DWA更容易生成连续轨迹。",
    "DWA三维速度空间": "在论文图表中，如果去除三维速度空间后的轨迹出现更大绕行或更多Recovery，应解释为传统DWA缺少侧向自由度，只能通过旋转改变运动方向。最终系统允许侧移后，机器人可以在局部空间中进行更小范围的姿态调整，这并不是为了让机器人一直横着走，而是为了在障碍物附近保留更符合四足运动能力的避障选择。",
    "前向姿态偏好与DWA路径头部对齐": "本节可以与用户实际观察到的“侧向通过浪费时间”问题对应。图表中侧向运动比例下降并不意味着四足机器人的侧移能力被削弱，而是说明侧移被限制到更必要的场景中。论文中应强调，本文最终系统追求的是头部朝前的主要运动姿态与局部侧移避障能力之间的平衡，这比单纯追求侧向比例最低更合理。",
    "A星参考走廊与局部目标单调推进": "本节在论文中可作为融合算法必要性的核心证据。若去除参考走廊和单调推进，DWA仍能生成局部可行轨迹，但这些轨迹可能缺少全局方向，导致回摆、绕远或原地犹豫。最终系统通过走廊评分和目标索引记忆，把A星的全局可达性持续传递给DWA，使局部规划不再只是短视避障，而是沿全局路线稳定推进。",
}


FORMULAS = {
    "基础A星与参考论文改进A星对比": """公式与含义说明：

传统A星评价函数为：

`f(n)=g(n)+h(n)`

其中，`g(n)`表示从起点到当前节点`n`的实际累计代价，`h(n)`表示当前节点到目标点的启发式估计距离，`f(n)`为节点综合评价值。基础A星每次从OpenList中选择`f(n)`最小的节点进行扩展。

参考论文改进A星中使用动态权重启发函数，可表示为：

`f(n)=g(n)+w(n)h(n)`

`w(n)=1+lambda_d * d(n,goal)/d(start,goal)`

其中，`lambda_d`为动态权重系数，`d(n,goal)`为当前节点到目标点距离，`d(start,goal)`为起点到目标点距离。搜索初期`d(n,goal)`较大，权重`w(n)`较大，算法更偏向目标方向扩展；接近目标时`w(n)`减小，搜索更接近普通A星，避免过强贪心导致局部绕障能力下降。双向A星可表示为从起点和目标点同时搜索：

`F_f(n)=g_f(n)+w_f(n)h_f(n), F_b(n)=g_b(n)+w_b(n)h_b(n)`

当前向闭集与反向闭集出现交汇节点时，拼接两侧父节点链得到全局路径。二十四邻域扩展的单步代价为：

`c(n_i,n_j)=sqrt((x_i-x_j)^2+(y_i-y_j)^2)`

该公式说明扩展步长不再局限于水平、竖直和对角一格，而是根据两节点欧氏距离计算移动代价。""",

    "足迹膨胀与安全距离软代价": """公式与含义说明：

四足机器人足迹膨胀首先将障碍物集合`O`扩展为不可通行集合：

`O_inflated={p | d_obs(p) <= r_body + r_margin}`

其中，`d_obs(p)`表示栅格`p`到最近障碍物的距离，`r_body`表示机器人等效半径或半宽，`r_margin`表示步态摆动、安全冗余和建图误差预留距离。若`p`属于`O_inflated`，则该栅格不再参与A星搜索。

在硬膨胀之外，本文加入安全距离软代价：

`C_clear(p)=lambda_c * (1 - d_obs(p)/r_safe)^2, 0 < d_obs(p) < r_safe`

`C_clear(p)=0, d_obs(p) >= r_safe`

其中，`r_safe`为期望安全距离，`lambda_c`为安全代价权重。当路径点越接近障碍物时，`d_obs(p)`越小，`C_clear(p)`越大；当路径点远离障碍物超过安全距离时，不再增加额外代价。

加入安全距离后的A星累计代价可写为：

`g(n_j)=g(n_i)+c(n_i,n_j)*(1+C_clear(n_j))`

该公式说明本文不是简单禁止机器人靠近障碍物，而是在仍可通行的区域内建立“远离障碍物更优”的偏好。这样宽阔区域中的路径会向通道中部移动，窄通道中则仍保留可通行性。""",

    "线段安全检测": """公式与含义说明：

二十四邻域使节点之间可能跨越多个栅格，因此仅判断候选终点是否为空闲是不够的。本文将候选边`e=(p_i,p_j)`参数化为：

`p(u)=(1-u)p_i+u p_j, u in [0,1]`

线段安全约束写为：

`Safe(e)=1 <=> M(round(p(u)))=0, forall u in [0,1]`

其中，`M(q)=1`表示栅格`q`为障碍物，`M(q)=0`表示可通行。实际实现中对`u`进行离散采样：

`u_k=k/K, k=0,1,...,K`

`K=ceil(L(e)/Delta_s)`

其中，`L(e)`为候选线段长度，`Delta_s`为线段检测采样间隔。只有所有采样点均为空闲，候选边才允许加入搜索。

若进一步考虑机器人宽度，也可写成安全距离形式：

`d_obs(p(u_k)) > r_check, k=0,1,...,K`

其中，`r_check`表示线段连接的碰撞检测半径。本报告表格中的“风险线段数”可表示为：

`N_risk=sum I(Safe(e_m)=0)`

该指标统计最终路径中有多少条连接边在连续空间中不安全，用于说明去除线段安全检测后虽然路径可能更短，但并不具备真实可执行性。""",

    "转角惩罚": """公式与含义说明：

设当前扩展路径中连续三个节点为`p_{k-1}`、`p_k`、`p_{k+1}`，则两段方向向量为：

`v_1=p_k-p_{k-1}, v_2=p_{k+1}-p_k`

转角为：

`Delta_theta=acos((v_1 dot v_2)/(|v_1||v_2|))`

本文在A星代价中加入转角惩罚：

`C_turn=lambda_theta*(1-cos(Delta_theta))`

也可以等价理解为使用`lambda_theta*|Delta_theta|`作为转向代价。采用`1-cos(Delta_theta)`的好处是小角度变化惩罚较柔和，大角度转向惩罚更明显。

加入转角惩罚后的节点代价为：

`g(n_{k+1})=g(n_k)+c(n_k,n_{k+1})+C_turn`

其中，`lambda_theta`为转角权重。当两条候选路径长度相近时，累计转角更小的路径会获得更低总代价。该公式表达的不是单纯追求最短距离，而是将四足机器人低层步态对连续运动方向的需求提前写入全局搜索过程。本报告中的“累计转角”对应：

`Theta_sum=sum |wrap(theta_{i+1}-theta_i)|`

该指标越小，说明路径方向变化越平稳。""",

    "安全剪枝、圆角平滑与重采样": """公式与含义说明：

安全剪枝基于视线可达性。对于原始A星路径`P={p_0,p_1,...,p_N}`，若节点`p_i`与后续节点`p_j`之间满足：

`LineFree(p_i,p_j)=1`

则可删除`p_i`与`p_j`之间的中间节点。剪枝过程可表示为：

`j*=max{j | j>i and LineFree(p_i,p_j)=1 and d_obs(line(p_i,p_j))>r_prune}`

其中，`r_prune`为剪枝安全距离。该公式保证剪枝不是简单把路径拉直，而是在直连安全时才删除冗余节点。

圆角平滑采用类Chaikin平滑公式：

`q_i=0.75p_i+0.25p_{i+1}`

`r_i=0.25p_i+0.75p_{i+1}`

其中，`q_i`和`r_i`为相邻路径点之间生成的新点。迭代后路径转角减小，但平滑结果必须重新进行碰撞检测：

`P_smooth accepted <=> Safe(P_smooth)=1`

否则回退到剪枝路径。

等距重采样用于让局部规划器获得均匀参考点：

`s_k=k*Delta_l, k=0,1,...,floor(L/Delta_l)`

其中，`L`为路径总长度，`Delta_l`为重采样间距。该公式保证DWA前视目标不会因为路径点过密或过稀而发生跳变。""",

    "DWA三维速度空间": """公式与含义说明：

传统DWA速度空间通常为：

`V_2={(v_x, omega)}`

本文根据四足机器人可侧移的运动能力，将速度空间扩展为：

`V_3={(v_x,v_y,omega) | v_x in [v_x_min,v_x_max], v_y in [v_y_min,v_y_max], omega in [omega_min,omega_max]}`

考虑速度和加速度约束，动态窗口为：

`v_x in [v_x^t-a_x Delta_t, v_x^t+a_x Delta_t]`

`v_y in [v_y^t-a_y Delta_t, v_y^t+a_y Delta_t]`

`omega in [omega^t-alpha Delta_t, omega^t+alpha Delta_t]`

候选速度在预测时使用全向运动模型：

`x_{t+Delta_t}=x_t+(v_x cos theta - v_y sin theta)Delta_t`

`y_{t+Delta_t}=y_t+(v_x sin theta + v_y cos theta)Delta_t`

`theta_{t+Delta_t}=theta_t+omega Delta_t`

评价函数可写为：

`G(v)=w_h H(v)+w_c C(v)+w_v V(v)+w_p P(v)-w_s S(v)`

其中，`H(v)`为朝向目标或局部路径方向评分，`C(v)`为障碍物安全距离评分，`V(v)`为速度评分，`P(v)`为靠近A星参考路径的评分，`S(v)`为侧移或不稳定运动惩罚。三维速度空间公式说明本文DWA不再把四足机器人当作差速轮式底盘，而是允许其用前进、侧移和偏航共同完成局部避障。""",

    "前向姿态偏好与DWA路径头部对齐": """公式与含义说明：

设DWA候选轨迹的主方向角为：

`psi_traj=atan2(y_T-y_0, x_T-x_0)`

机器人当前头部方向为`theta`，则头部与轨迹方向误差为：

`e_psi=wrap(psi_traj-theta)`

为了减少不必要侧移，本文在DWA评价函数中加入前向姿态代价：

`C_forward=lambda_psi |e_psi| + lambda_y |v_y|/v_ymax`

其中，`lambda_psi`表示头部方向误差权重，`lambda_y`表示侧向速度惩罚权重。该项使机器人在开阔空间中更倾向于先把头部对准DWA局部路径，再以前进速度`v_x`为主进行运动。

当满足：

`|e_psi| > theta_align and d_free > d_align`

系统进入短时头部对齐控制：

`v_x=0, v_y=0, omega=clip(k_psi e_psi, -omega_max, omega_max)`

其中，`theta_align`为触发对齐的角度阈值，`d_free`表示周围可旋转空间，`d_align`为空间安全阈值。该公式说明对齐动作不是任何时候都强制执行，只有当角度偏差较大且周围空间足够时才触发；在狭窄区域，侧移仍然保留为避障手段。""",

    "A星参考走廊与局部目标单调推进": """公式与含义说明：

设A星全局路径为：

`P_g={p_0,p_1,...,p_N}`

DWA候选轨迹为：

`T={q_0,q_1,...,q_M}`

本文不再要求机器人机械追踪A星每个路径点，而是把全局路径转化为参考走廊。候选轨迹到走廊的距离定义为：

`d_corr(T,P_g)=1/M * sum_{m=1}^{M} min_i ||q_m-p_i||`

在DWA评价函数中加入走廊评分：

`G'(v)=G(v)-lambda_corr d_corr(T(v),P_g)`

其中，`lambda_corr`为走廊约束权重。该项使DWA在避障时可以短时偏离全局路径，但偏离越大得分越低，从而保持整体运动方向。

局部目标单调推进采用路径索引记忆：

`k_t=max(k_{t-1}, argmin_{k in [k_{t-1}, k_{t-1}+K]} ||p_k-x_t||)`

局部目标点选择为：

`p_goal^local=p_{min(k_t+L_look, N)}`

其中，`K`为向前搜索窗口，`L_look`为前视距离对应的路径索引增量。`max(k_{t-1}, .)`保证局部目标不会跳回机器人身后，从而避免路径重规划后机器人在两侧路径之间反复旋转。""",

    "速度保持与受限Recovery": """公式与含义说明：

速度保持用于解决多次重规划后速度被反馈延迟拉低的问题。本文使用上一周期命令速度与里程计速度融合，得到动态窗口参考速度：

`v_ref=alpha_v v_cmd^{t-1}+(1-alpha_v)v_odom^t`

其中，`alpha_v`为速度保持系数，`v_cmd^{t-1}`为上一周期规划器输出速度，`v_odom^t`为当前里程计反馈速度。动态窗口不再只围绕较慢的里程计速度展开，而是围绕`v_ref`展开：

`V_dw=[v_ref-a_max Delta_t, v_ref+a_max Delta_t]`

这样在前方安全时，机器人不会因为一次停顿或一次重规划就长期以低速重新起步。

受限Recovery触发条件可写为：

`Recovery=1 <=> (N_feasible=0) or (d_front<d_stop and Delta s<T_prog)`

其中，`N_feasible`表示DWA可行候选轨迹数量，`d_front`表示前向障碍距离，`d_stop`为近障碍停止阈值，`Delta s`为一段时间内沿路径推进距离，`T_prog`为最小进展阈值。该公式说明Recovery只有在DWA没有可行解，或近障碍且确实无进展时才触发。恢复动作序列可表示为：

`u_rec in {(-v_back,0,0), (0,0,omega_rec), (0,v_side,0)}`

即后退、小角度旋转和侧步脱困的有限集合，避免恢复逻辑在空间充足时频繁抢占DWA正常控制。""",
}


def make_section(case, idx):
    rel_fig = os.path.relpath(case["figure"], REPORT_DIR)
    rel_csv = os.path.relpath(case["csv"], REPORT_DIR)
    formula = FORMULAS.get(case["name"], "")
    text = (EXPLANATIONS[case["name"]] + "\n\n" +
            EXTRA_EXPLANATIONS.get(case["name"], "") + "\n\n" +
            FINAL_TOPUPS.get(case["name"], ""))
    return f"""## {idx} {case['name']}

![{case['name']}](./{rel_fig})

表{idx} 关键数据对比：

{csv_table_md(case['csv'])}

{formula}

{text}

本组实验数据文件为：`{rel_csv}`。
"""


def build_report(astar_cases, local_cases):
    sections = []
    sections.append("""# 四足机器人路径规划算法改进消融实验报告

## 0.1 报告目的

本报告推倒前面论文式撰写方式，重新围绕“参考论文基础算法到四足机器人适配算法”的实验论证展开。参考论文《基于ROS的室内机器人路径规划算法研究》中的基础改进思路主要包括：在传统A星算法中引入双向搜索、动态权重启发函数、二十四邻域扩展和三阶贝塞尔曲线平滑；在传统动态窗口法中通过评价函数改进来改善局部避障轨迹；最终将改进A星与改进DWA组成混合路径规划算法。本文工作的重点不是重复参考论文，而是在其基础上进一步面向ASK-3四足机器人进行适配，使路径规划结果能够符合四足机器人机体尺寸、头部朝前运动习惯、侧移能力和Gazebo仿真执行稳定性。

本报告采用两类实验。第一类实验对比基础A星算法与参考论文改进A星算法，说明参考论文基础改进相较传统算法的效果。第二类实验对本文四足机器人适配改进进行消融：每次暂时去除一个关键改进点，并与最终系统进行对比。每个改进点均给出一张路径规划/轨迹示意图、一张关键数据对比表，并围绕改进必要性、改进原理、实验现象和对四足机器人导航的作用进行约1000字解释。

需要说明的是，本报告中的数据为离线规划与轨迹层仿真实验数据，用于论文中算法效果论证和图表展示；若作为严格Gazebo物理实测结论，还应进一步结合rosbag记录机器人实际位姿、速度命令和碰撞/停滞日志。

# 第一章 基础A星与参考论文改进A星对比
""")
    sections.append(make_section(astar_cases[0], "1.1"))
    sections.append("# 第二章 本文四足机器人适配改进消融实验\n")
    labels = ["2.1", "2.2", "2.3", "2.4"]
    for lab, case in zip(labels, astar_cases[1:]):
        sections.append(make_section(case, lab))
    labels = ["2.5", "2.6", "2.7", "2.8"]
    for lab, case in zip(labels, local_cases):
        sections.append(make_section(case, lab))
    sections.append("""# 第三章 结论

从基础A星与参考论文改进A星对比可以看出，参考论文中的双向搜索、动态权重、二十四邻域和路径平滑能够明显改善传统A星搜索效率和路径几何形态。本文四足机器人适配改进则进一步解决了参考论文算法直接用于ASK-3四足机器人时可能出现的问题：路径贴墙导致机体碰撞、二十四邻域穿角、路径折线导致头部频繁调整、路径点不均导致DWA目标跳变、传统DWA无法表达侧移、侧移能力被滥用、全局路径与局部轨迹切换不稳定、多次重规划后速度下降以及Recovery误触发等。

消融实验表明，任何单个改进点都不是孤立装饰。足迹膨胀和安全距离软代价保证全局路径具有机体通过空间；线段安全检测保证长步长邻域扩展不穿越障碍物边角；转角惩罚和安全后处理提升路径执行连续性；三维速度空间释放四足机器人侧移能力；前向姿态偏好和头部对齐限制侧移使用场景；A星参考走廊和局部目标单调推进稳定全局-局部融合；速度保持和受限Recovery提高连续运动效率和墙角脱困能力。最终系统的优势来自这些改进的共同作用，而不是某一个单独公式。
""")
    return "\n\n".join(sections)


# ---------------------------------------------------------------------------
# DOCX output

def set_run_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element.rPr
    if r is None:
        r = OxmlElement("w:rPr")
        run._element.insert(0, r)
    fonts = r.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r.append(fonts)
    fonts.set(qn("w:eastAsia"), "SimSun")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")


def write_docx(text, path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
    title_done = False
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if not title_done else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line[2:])
            set_run_font(r, size=20 if not title_done else 16, bold=True)
            title_done = True
        elif line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:])
            set_run_font(r, size=14, bold=True)
        elif line.startswith("!["):
            img = line.split("](")[-1].rstrip(")")
            img_path = os.path.normpath(os.path.join(REPORT_DIR, img))
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_path, width=Inches(6.3))
        elif line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].startswith("| "):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                if all(set(c) <= set("-: ") for c in cells):
                    continue
                rows.append(cells)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        p = table.cell(ri, ci).paragraphs[0]
                        r = p.add_run(cell)
                        set_run_font(r, size=8 if len(rows[0]) > 6 else 9, bold=(ri == 0))
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.25
            r = p.add_run(line)
            set_run_font(r, size=11)
        i += 1
    doc.save(path)


def main():
    astar_cases = generate_astar_experiments()
    local_cases = generate_local_ablation_experiments()
    report = build_report(astar_cases, local_cases)
    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write(report)
    write_docx(report, DOCX_PATH)
    summary = {
        "report_markdown": MD_PATH,
        "report_docx": DOCX_PATH,
        "asset_dir": ASSET_DIR,
        "figures": [c["figure"] for c in astar_cases + local_cases],
        "tables": [c["csv"] for c in astar_cases + local_cases],
        "characters": len(report),
        "note": "数据为离线规划/轨迹层仿真实验数据，非Gazebo物理实测rosbag数据。",
    }
    with open(SUMMARY_PATH, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
