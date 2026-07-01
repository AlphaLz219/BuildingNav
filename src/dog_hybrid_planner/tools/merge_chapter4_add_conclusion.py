#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Merge thesis chapter 4 into chapter 3 and add a conclusion chapter.

The script edits a working copy of the user's docx. It keeps figures, tables
and existing experiment content in place, while changing the chapter structure:

  Chapter 3: DWA local planning + A*/DWA fusion
  Chapter 4: ROS/Gazebo simulation and result analysis
  Chapter 5: conclusion and outlook
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


INPUT = Path("/tmp/四足机器人导航路径规划_work.docx")
OUTPUT = Path("/tmp/四足机器人导航路径规划_章节合并结论版.docx")


def set_text(paragraph, text: str) -> None:
    paragraph.text = text


def find_para(doc: Document, startswith: str):
    for paragraph in doc.paragraphs:
        if "toc" in paragraph.style.name.lower():
            continue
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    raise RuntimeError(f"paragraph not found: {startswith}")


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def paragraph_index(doc: Document, paragraph) -> int:
    for idx, candidate in enumerate(doc.paragraphs):
        if candidate._element is paragraph._element:
            return idx
    raise RuntimeError("paragraph index not found")


def body_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.25


def insert_before(anchor, text: str, style: str = "Normal"):
    p = anchor.insert_paragraph_before(text)
    p.style = style
    if style == "Normal":
        body_format(p)
    return p


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


def replace_intro_and_structure(doc: Document) -> None:
    # Chapter structure description in Chapter 1.
    set_text(
        find_para(doc, "第三章研究适用于四足机器人运动特性的改进DWA局部路径规划算法"),
        "第三章研究适用于四足机器人运动特性的改进DWA局部路径规划算法及其与A*全局路径的融合方法。首先，将DWA速度空间扩展为前进速度、侧向速度和偏航角速度组成的三维动态窗口，使局部规划器能够表达四足机器人前进、侧移和转向的复合运动能力；其次，通过前向姿态偏好与路径头部对齐机制限制无意义侧移，使机器人在空间充足时优先头部朝前沿路径运动；最后，将A*全局路径转化为DWA参考走廊，并设计局部目标单调推进机制，使局部轨迹既具有实时避障能力，又能够稳定沿全局方向推进。",
    )
    set_text(
        find_para(doc, "第四章研究A*算法与DWA算法的融合改进方法"),
        "第四章进行ROS与Gazebo平台下的系统仿真实验。该章首先介绍ASK-3四足机器人仿真模型、激光雷达、AMCL定位、TF坐标变换关系、RViz可视化和底层步态控制节点之间的关系；随后补充速度保持与受限Recovery机制，解决多次重规划后速度下降以及恢复行为过度触发的问题；最后结合各项消融实验数据，对本文改进算法在导航时间、轨迹长度、平均速度、累计转角、侧向运动比例、Recovery次数和碰撞风险等指标上的效果进行综合分析。",
    )
    set_text(
        find_para(doc, "第五章进行ROS与Gazebo平台下的系统仿真实验"),
        "第五章为结论与展望。该章在理论分析与仿真实验结果基础上，对本文围绕四足机器人路径规划所完成的全局A*改进、局部DWA改进、全局—局部融合改进以及ROS/Gazebo仿真验证进行归纳总结，分析当前研究中仍存在的模型简化、实机验证不足和复杂动态场景适应性不足等问题，并对后续面向真实四足机器人平台、多传感器融合、复杂地形导航和学习型规划控制结合等方向进行展望。",
    )

    # Chapter 3 title and introduction.
    ch3 = find_para(doc, "3  基于DWA算法的局部路径规划算法的研究")
    ch3.text = "3  基于DWA算法的局部路径规划与混合路径规划算法研究"
    try:
        ch3.style = "Heading 1"
    except Exception:
        pass

    intro = [
        "上一章研究了基于已有改进A*算法的全局路径规划方法，并针对ASK-3四足机器人机体尺寸、通行安全距离和路径执行连续性进行了适配。全局路径能够给出从起点到目标点的整体通行方向，但它本质上仍是基于静态地图的宏观规划结果，无法直接处理机器人周围实时障碍物、定位误差、模型惯性和步态控制延迟等问题。对于四足机器人而言，即使全局路径在栅格地图中可行，机器人在墙角、障碍物短边或狭窄通道处仍可能因为姿态调整不当、侧向运动过多或局部目标跳变而出现停顿。因此，有必要在全局路径基础上进一步研究局部实时规划与全局—局部融合方法。",
        "动态窗口法（Dynamic Window Approach，DWA）能够根据机器人当前速度、加速度约束和预测时间，在速度空间中采样候选速度，并通过短时轨迹预测评价每组速度的安全性、目标趋近性和运动效率。该方法的优势在于能够直接输出速度命令，适合与四足机器人底层步态控制接口连接。但是，传统DWA主要面向差速轮式机器人，速度空间通常只包含前向速度和角速度，无法完整表达ASK-3四足机器人具备的前进、侧移和偏航复合运动能力。如果直接使用传统二维DWA，机器人在局部避障时容易依赖大幅转向完成绕行，通行效率和姿态自然性均会受到影响。",
        "另一方面，仅将A*全局路径的关键点依次交给DWA跟踪，也难以保证复杂地图中的稳定导航。当目标点与机器人之间间隔墙体，或者全局路径在机器人运动过程中发生明显更新时，最近路径点可能突然跳到机器人另一侧或身后，导致机器人原地旋转、反复寻找路径。由此可见，局部DWA改进与A*—DWA融合机制具有内在联系：前者决定机器人能否在局部空间中生成符合四足运动特性的速度，后者决定这些速度是否始终服务于稳定的全局前进方向。因此，将原本单独讨论的混合路径规划内容并入本章更符合算法逻辑。",
        "基于上述分析，本章首先将DWA速度采样空间扩展为三维速度空间，使局部规划器能够同时考虑前进速度、侧向速度和偏航角速度；随后引入前向姿态偏好与DWA路径头部对齐机制，在空间充足且局部轨迹方向明确时优先保持头部朝前运动；最后研究A*参考走廊与局部目标单调推进机制，使DWA在实时避障的同时受到全局路径软约束引导。通过消融实验对比各项改进前后的轨迹长度、导航时间、侧向运动比例、累计转角和Recovery次数，验证本文局部规划与融合规划改进对四足机器人导航稳定性的作用。",
    ]
    for idx, text in zip(range(304, 308), intro):
        doc.paragraphs[idx].text = text
        body_format(doc.paragraphs[idx])


def merge_chapter_four_into_three(doc: Document) -> None:
    # Remove old chapter-3 summary; the merged chapter will have a single final summary.
    for marker in [
        "3.3  章节小结",
        "本章围绕局部路径规划算法展开研究",
        "针对这一问题，本文首先将DWA速度空间扩展为三维速度空间",
        "其次，本文针对侧移能力可能被滥用的问题",
    ]:
        delete_paragraph(find_para(doc, marker))

    old_ch4 = find_para(doc, "4  混合路径规划算法的研究")
    old_ch4.text = "3.3  A*与DWA融合路径规划改进研究"
    old_ch4.style = "Heading 2"

    h = find_para(doc, "4.1  引言")
    h.text = "3.3.1  混合路径规划融合思路"
    h.style = "Heading 3"

    h = find_para(doc, "4.2  DWA算法适配四足机器人的改进研究")
    h.text = "3.3.2  A*参考走廊与局部目标单调推进机制"
    h.style = "Heading 3"

    delete_paragraph(find_para(doc, "A星参考走廊与局部目标单调推进："))

    replacements = {
        "(4.1)": "(3.4)",
        "(4.2)": "(3.5)",
        "(4.3)": "(3.6)",
        "(4.4)": "(3.7)",
        "图4.1": "图3.3",
        "表4.1": "表3.3",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue
        for old, new in replacements.items():
            text = text.replace(old, new)
        paragraph.text = text

    summary_heading = find_para(doc, "4.3  本章小结")
    summary_heading.text = "3.4  本章小结"
    summary_heading.style = "Heading 2"

    summary_texts = [
        "本章在第二章全局路径规划研究基础上，完成了面向ASK-3四足机器人的局部路径规划与混合路径规划融合改进。与传统轮式机器人导航不同，四足机器人既具备一定侧移能力，又具有头部朝前运动更自然、频繁原地转向代价较高、步态控制存在响应延迟等特点。因此，本文没有简单沿用传统二维DWA模型，而是从速度空间、姿态偏好和全局—局部信息传递三个层面重新组织局部规划框架。",
        "首先，本文将DWA速度空间由传统二维形式扩展为包含前进速度、侧向速度和偏航角速度的三维动态窗口，使候选轨迹能够表达四足机器人真实可执行的复合运动。消融实验表明，去除三维速度空间后，机器人只能通过大弧度转向绕过障碍，轨迹长度和Recovery次数增加；加入该改进后，机器人能够在障碍物短边、墙角和窄通道附近通过小幅侧移完成局部调整，提高了局部通过效率。",
        "其次，针对侧移能力可能被滥用的问题，本文加入前向姿态偏好与DWA路径头部对齐机制。该机制并不否定侧移，而是将侧移限制为必要的局部避障动作：在开阔空间中，机器人优先将头部与局部路径方向对齐，并以前进速度为主完成运动；在空间受限或障碍边缘附近，侧移仍作为辅助脱困和避障手段。实验结果表明，该机制能够降低无意义侧向运动比例，使机器人运动姿态更符合现实四足机器人头部朝前行走的习惯。",
        "最后，本文将原第四章的混合路径规划内容并入本章，进一步研究A*全局路径与DWA局部轨迹之间的稳定耦合。A*参考走廊使全局路径不再是必须机械追踪的折线，而是作为DWA评分函数中的软约束持续影响候选轨迹选择；局部目标单调推进机制则通过路径索引记忆避免目标点跳回身后或突然切换到墙体另一侧。消融实验说明，去除该机制后DWA仍能生成局部可行轨迹，但容易出现回摆、绕远和原地犹豫；加入参考走廊与目标推进后，机器人能够在实时避障的同时稳定沿全局可达方向前进。本章研究表明，四足机器人局部规划与混合规划不应被割裂处理，二者共同决定机器人能否在复杂室内地图中保持连续、自然和高效的运动。",
    ]
    idx = paragraph_index(doc, summary_heading)
    for off, text in enumerate(summary_texts, start=1):
        doc.paragraphs[idx + off].text = text
        body_format(doc.paragraphs[idx + off])


def renumber_simulation_chapter(doc: Document) -> None:
    set_text(
        find_para(doc, "5  基于ROS平台的四足机器人混合路径规划仿真实验和结果分析"),
        "4  基于ROS平台的四足机器人混合路径规划仿真实验和结果分析",
    )
    find_para(doc, "4  基于ROS平台").style = "Heading 1"
    heading_map = {
        "5.1  引言": "4.1  引言",
        "5.2  ROS与Gazebo仿真实验平台": "4.2  ROS与Gazebo仿真实验平台",
        "5.3  速度保持与受限Recovery机制": "4.3  速度保持与受限Recovery机制",
        "5.4  本章小结": "4.4  本章小结",
    }
    for old, new in heading_map.items():
        set_text(find_para(doc, old), new)

    set_text(
        find_para(doc, "前四章分别完成了四足机器人全局路径规划、局部路径规划以及全局—局部融合路径规划算法的研究"),
        "前三章分别完成了四足机器人全局路径规划、局部路径规划以及全局—局部融合路径规划算法的研究。第二章在已有改进A*算法基础上，针对ASK-3四足机器人机体尺寸、通行安全距离和路径执行连续性进行了进一步适配；第三章将DWA局部规划算法由传统二维速度空间扩展到四足机器人可执行的三维速度空间，并通过前向姿态偏好、头部对齐、A*参考走廊和局部目标单调推进机制增强混合规划稳定性。上述算法改进主要从路径搜索、局部避障和融合逻辑三个层面提高了机器人导航效果。",
    )
    set_text(
        find_para(doc, "同时，第五章也承担对全文算法体系进行工程闭环验证的作用"),
        "同时，第四章也承担对全文算法体系进行工程闭环验证的作用。前三章分别从全局搜索、局部避障和算法融合角度证明各模块有效，本章则进一步关注这些模块在真实仿真流程中的协同表现，尤其关注多次重规划、局部停顿和恢复动作对最终导航效率的影响。",
    )

    replacements = {
        "图5.1  RViz中仿真四足机器人路径规划图": "图4.1  RViz中仿真四足机器人路径规划图",
        "图5.1 四足机器人路径规划TF框架图": "图4.2 四足机器人路径规划TF框架图",
        "如图5.1所示，本文仿真系统": "如图4.2所示，本文仿真系统",
        "图5.2 四足机器人路径规划ROS节点图": "图4.3 四足机器人路径规划ROS节点图",
        "如图5.2所示，Gazebo": "如图4.3所示，Gazebo",
        "(5.1)": "(4.1)",
        "图5.3 速度保持与受限Recovery消融实验对比图": "图4.4 速度保持与受限Recovery消融实验对比图",
        "表5.1": "表4.1",
        "由图5.3和表4.1": "由图4.4和表4.1",
        "由图5.3和表5.1": "由图4.4和表4.1",
        "从表5.1": "从表4.1",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue
        for old, new in replacements.items():
            text = text.replace(old, new)
        paragraph.text = text


def insert_conclusion(doc: Document) -> None:
    ref = find_para(doc, "参考文献")
    conclusion_parts = [
        ("5  结论与展望", "Heading 1"),
        ("5.1  研究结论", "Heading 2"),
        (
            "本文围绕ASK-3四足机器人在室内复杂环境中的自主导航问题，基于ROS与Gazebo仿真平台，对全局路径规划、局部路径规划、全局—局部融合以及仿真执行稳定性进行了系统研究。研究的出发点并不是单纯追求地图上最短的几何路径，而是结合四足机器人机体尺寸、步态摆动空间、头部朝前运动习惯、侧向运动能力和底层速度响应延迟等因素，分析传统移动机器人路径规划算法在四足机器人仿真中出现贴墙、穿角、频繁旋转、侧向运动过多、局部目标跳变以及速度逐渐下降等问题的原因，并在此基础上提出具有四足机器人适配性的改进方案。",
            "Normal",
        ),
        (
            "在理论分析方面，第二章以已有改进A*算法为基础，首先保留双向搜索、动态权重、二十四邻域扩展和路径平滑等能够提高搜索效率与路径形态的优点，随后进一步引入足迹膨胀与安全距离软代价、线段安全检测、转角惩罚、安全剪枝、圆角平滑与重采样等方法。经过分析可以得出：四足机器人全局路径规划不能只判断路径中心线是否位于自由栅格，还必须考虑机器人机体和腿部摆动所需的安全空间；同时，搜索阶段得到的路径不能只追求距离短，还应尽量减少急转弯和不均匀路径点，以降低后续局部规划器的目标跳变风险。因此，改进A*算法在本文中承担的不只是“找到一条路”的功能，更是为局部DWA提供一条具有安全余量、方向连续性和跟踪稳定性的全局参考线。",
            "Normal",
        ),
        (
            "第三章将局部DWA规划与A*—DWA融合机制合并研究，体现了局部避障与全局方向约束之间的内在联系。通过将DWA速度空间扩展为由前进速度、侧向速度和偏航角速度组成的三维动态窗口，本文使局部规划器能够表达四足机器人真实可执行的复合运动；通过前向姿态偏好与头部对齐机制，又限制了侧移能力的滥用，使机器人在空间充足时优先头部朝前运动，在障碍物附近再使用侧移完成局部避障；通过A*参考走廊和局部目标单调推进机制，本文进一步避免了全局路径与局部轨迹之间的目标跳变，使DWA不再仅凭短时局部目标做出决策，而是在全局可达方向的软约束下生成局部速度。由此可以归纳出本文混合路径规划的核心观点：A*与DWA的结合不是简单串联，而是需要设计稳定的信息传递接口，使全局规划提供方向和拓扑可达性，局部规划提供避障和速度执行能力。",
            "Normal",
        ),
        (
            "在实验结果方面，本文通过多组消融实验对各项改进进行了验证。A*相关实验表明，足迹膨胀与安全距离软代价能够提高路径与障碍物之间的安全余量，线段安全检测能够避免二十四邻域扩展造成的穿角问题，转角惩罚能够降低路径累计转角，安全剪枝、圆角平滑与重采样能够改善路径点分布并提升DWA跟踪稳定性。DWA相关实验表明，三维速度空间能够减少绕行和Recovery次数，前向姿态偏好能够降低无意义侧向运动比例，A*参考走廊和局部目标单调推进能够显著抑制路径更新后的原地旋转和目标来回切换。ROS/Gazebo仿真实验进一步说明，速度保持与受限Recovery机制能够缓解多次重规划后的速度下降问题，并减少恢复行为对DWA正常避障的干扰。综合理论分析和实验结果可以判断，本文所提出的改进并非彼此孤立的算法装饰，而是从全局安全、局部运动、融合稳定和执行连续四个层面共同提升四足机器人路径规划效果。",
            "Normal",
        ),
        (
            "因此，本文最终形成的总观点是：面向四足机器人的路径规划算法应从“质点或轮式底盘路径规划”转向“考虑机体尺寸、运动姿态和执行稳定性的系统规划”。对于ASK-3这类具备侧移能力的四足机器人，合理的路径规划结果既要保证地图层面的可达性，也要保证机器人运动姿态的自然性；既要允许局部避障时使用侧向运动，又要避免在宽阔空间中长期侧身通过；既要依靠A*提供全局方向，又不能机械追踪A*折线而忽视局部障碍；既要保留Recovery脱困能力，又要防止恢复逻辑过度抢占DWA控制权。只有将这些因素统一考虑，才能使仿真中的四足机器人表现出更接近真实运动习惯的连续导航行为。",
            "Normal",
        ),
        ("5.2  当前不足", "Heading 2"),
        (
            "尽管本文完成了面向四足机器人路径规划的多项改进，并通过仿真实验验证了算法有效性，但研究仍存在一定不足。首先，本文实验主要基于Gazebo仿真环境完成，虽然仿真平台能够反映机器人模型、传感器、TF坐标关系和底层速度接口之间的工程耦合，但与真实四足机器人相比仍存在差距。真实机器人会受到地面摩擦、机体振动、关节误差、电机响应、传感器噪声和通信延迟等因素影响，仿真中得到的路径平滑性、速度稳定性和Recovery触发效果仍需要通过实机实验进一步验证。",
            "Normal",
        ),
        (
            "其次，本文在全局路径规划中主要采用二维栅格地图，将四足机器人机体和步态摆动范围转化为安全膨胀半径和软代价处理。这种方法适合室内平面环境，但对楼梯、斜坡、台阶、松软地面等三维复杂地形描述不足。四足机器人的优势之一在于具备跨越和适应非平整地形的能力，而本文路径规划仍主要面向平面障碍环境，没有进一步引入地形高度、落足稳定性、足端可达区域和机体姿态约束。因此，本文算法更适合验证室内路径规划与局部避障问题，对复杂地形通行能力的研究仍不充分。",
            "Normal",
        ),
        (
            "再次，本文DWA局部规划虽然扩展到三维速度空间，并加入前向姿态偏好与局部目标单调推进机制，但仍属于基于采样和评分函数的局部优化方法。其性能在一定程度上依赖权重参数、预测时间、速度采样分辨率和障碍物距离阈值。当环境结构发生明显变化，或者机器人面对高速动态障碍物时，固定权重评价函数可能难以及时适应。本文虽然通过消融实验说明各项权重和机制具有合理性，但仍未形成能够根据环境拥挤程度、通道宽度和机器人运动状态自动调整参数的自适应策略。",
            "Normal",
        ),
        (
            "最后，本文对低层步态控制的处理仍然较为抽象。路径规划节点通过前进、侧向和偏航三路速度控制机器人运动，但并未深入研究速度命令与具体足端轨迹、支撑相、摆动相和机体稳定裕度之间的关系。也就是说，本文主要解决上层路径规划和局部避障问题，而没有将足端接触规划、机体姿态控制和路径规划进行更紧密的统一建模。这使得算法能够在现有dog_sim运动代码基础上完成仿真验证，但距离完整的四足机器人运动规划系统仍有一定距离。",
            "Normal",
        ),
        ("5.3  未来展望", "Heading 2"),
        (
            "后续研究可以从以下几个方面进一步展开。第一，应在真实四足机器人平台上开展实机验证，记录机器人实际位姿、速度命令、关节状态、激光雷达数据和IMU数据，对比仿真轨迹与真实轨迹之间的差异。通过实机实验可以进一步检验足迹膨胀半径、侧移惩罚权重、头部对齐阈值和Recovery触发条件是否仍然适用，并为后续参数整定提供更可靠的数据基础。",
            "Normal",
        ),
        (
            "第二，可以将二维栅格地图扩展为包含高度、坡度、粗糙度和落足可行性的三维或二点五维地图。在此基础上，全局规划不再只考虑平面障碍物距离，而是同时考虑地形可通过性、机体姿态变化和足端落点安全性；局部规划也可以从单纯速度采样进一步扩展为速度—姿态—落足区域联合评价。这样能够更充分发挥四足机器人区别于轮式机器人的地形适应优势，使路径规划从室内平面导航拓展到楼梯、坡道和复杂地面环境。",
            "Normal",
        ),
        (
            "第三，可以研究自适应DWA评价函数和学习型局部规划方法。传统DWA权重依赖人工经验，难以在不同地图和不同拥挤程度下保持最优。未来可根据通道宽度、障碍物密度、目标方向变化和机器人当前速度状态动态调整安全距离权重、速度权重、走廊约束权重和侧移惩罚权重；也可以利用强化学习或模仿学习方法学习局部避障策略，再通过安全约束和传统规划方法保证其可解释性与可靠性。这样既能提升算法在复杂环境中的适应能力，又能避免纯学习方法在安全性上的不确定风险。",
            "Normal",
        ),
        (
            "第四，可以进一步加强上层路径规划与底层步态控制之间的耦合。本文将四足机器人底层运动接口抽象为前进、侧向和偏航速度，这种方式便于与DWA结合，但仍没有直接考虑足端接触序列和步态稳定性。未来可以在局部轨迹评价中加入步态代价、支撑稳定裕度、机体姿态变化率和足端可达性约束，使局部规划输出不仅在几何空间中可行，而且在步态执行层面更加平稳。通过将路径规划、运动控制和状态估计进一步融合，有望形成更适合真实四足机器人长期自主导航的完整系统。",
            "Normal",
        ),
        (
            "综上所述，本文工作证明了在已有A*与DWA混合路径规划框架基础上，针对四足机器人运动特点进行系统适配具有必要性和有效性。后续若能进一步结合真实机器人实验、三维地形建模、自适应参数优化和步态层约束，将有望使四足机器人在更复杂、更动态、更接近真实应用的场景中实现稳定、高效且符合运动习惯的自主导航。",
            "Normal",
        ),
    ]
    for text, style in conclusion_parts:
        p = insert_before(ref, text, style)
        if style == "Normal":
            body_format(p)


def main() -> None:
    doc = Document(INPUT)
    replace_intro_and_structure(doc)
    merge_chapter_four_into_three(doc)
    renumber_simulation_chapter(doc)
    insert_conclusion(doc)
    set_update_fields(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
