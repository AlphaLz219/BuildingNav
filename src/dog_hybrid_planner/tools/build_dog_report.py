#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate the quadruped path-planning experiment report (.docx) using
python-docx, mirroring the structure of the wheeled-robot report and
documenting every place where the algorithm had to change for the dog.
"""
import json
import os

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PKG = "/home/cjx/catkin_ws/src/dog_hybrid_planner"
RES = os.path.join(PKG, "experiments", "results")
OUT = os.path.join(PKG, "experiments", "四足机器人路径规划实验报告.docx")


def safe_load(path):
    try:
        with open(path) as f:
            return json.load(f)
    except Exception:
        return None


ASTAR = safe_load(os.path.join(RES, "dog_astar_benchmark.json"))
DWA   = safe_load(os.path.join(RES, "dog_dwa_benchmark.json"))
INTEG = safe_load(os.path.join(RES, "dog_integration.json"))

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Arial'
st.font.size = Pt(11)
rPr = st.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
rFonts.set(qn('w:ascii'), 'Arial')
rFonts.set(qn('w:hAnsi'), 'Arial')
rFonts.set(qn('w:eastAsia'), 'SimSun')

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def set_font(run, size=None, bold=None, font='Arial', eastAsia='SimSun'):
    run.font.name = font
    r = run._element.rPr
    if r is None:
        r = OxmlElement('w:rPr'); run._element.insert(0, r)
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); r.append(rFonts)
    rFonts.set(qn('w:ascii'), font); rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:eastAsia'), eastAsia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def P(text, size=11, bold=False, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(after)
    return p


def H(level, text, size):
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    r = p.add_run(text)
    set_font(r, size=size, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

H1 = lambda t: H(1, t, 16)
H2 = lambda t: H(2, t, 14)
H3 = lambda t: H(3, t, 12)


def BULLET(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    set_font(r, size=11)
    p.paragraph_format.space_after = Pt(2)


def CODE(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=9, font='Consolas', eastAsia='SimSun')
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), 'F2F2F2')
    pPr.append(sh)


def CAPTION(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=10)
    r.italic = True
    p.paragraph_format.space_after = Pt(10)


def add_image(path, width_in=5.8):
    if not os.path.exists(path):
        P(f"[Image not generated: {os.path.basename(path)}]", size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))


def make_table(data, col_widths_cm=None, header_row=True):
    rows, cols = len(data), len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    if col_widths_cm:
        for i, cw in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(cw)
    for i, row in enumerate(data):
        for j, txt in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(txt))
            set_font(r, size=10, bold=(i == 0 and header_row))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


# =====================================================================
# Title page
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("基于 ROS 的四足机器人室内路径规划与导航\n实验报告")
set_font(r, size=22, bold=True)
title.paragraph_format.space_before = Pt(60)
title.paragraph_format.space_after = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("(基于 dog_hybrid_planner 包,ASK-3 四足机器人 Gazebo 仿真)")
set_font(r, size=14)
sub.paragraph_format.space_after = Pt(40)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run("基于 py_hybrid_planner(轮式机器人)实现迁移\n"
                 "保留 hybrid A* + 改进 DWA 思路,针对四足机器人作算法适配\n"
                 "完成日期:2026 年 4 月")
set_font(r, size=12)

doc.add_page_break()

# =====================================================================
# 1. 实验目的与说明
# =====================================================================
H1("1. 实验目的与背景")
P("本实验在唐宁烨同学已完成的 py_hybrid_planner 轮式机器人路径规划"
  "工作基础上,将载具替换为 ASK-3 四足机器人(catkin_ws/src/dog_sim 目录),"
  "实现与原系统效果相同的全栈室内路径规划与自主导航。"
  "由于轮式机器人(差速驱动)与四足机器人(全向行走)在运动学、感知配置、"
  "执行接口等方面差异显著,部分核心算法需要做相应改造。"
  "本报告聚焦三个目标:")
BULLET("复现轮式机器人在 indoor.world 中的全局路径规划+局部避障导航能力。")
BULLET("将轮式机器人 hybrid A* + DWA 算法迁移到四足机器人上,"
       "明确每一处算法/接口改动并给出工程实现。")
BULLET("通过离线基准测试 + 在线 Gazebo 仿真两种方式,"
       "对比 OmniDWA(本工作)与原差速 DWA 在相同地图、相同场景下的指标。")
P("所有源码均存放在远程主机 /home/cjx/catkin_ws/src/dog_hybrid_planner/ 目录,"
  "依赖 ROS Noetic + Gazebo 11 + dog_sim 提供的 ASK-3 控制器与状态估计器。")

# =====================================================================
# 2. 系统平台对比
# =====================================================================
H1("2. 平台与接口差异")
H2("2.1 硬件/仿真模型")
make_table([
    ["项目", "轮式 (TurtleBot3 burger)", "四足 (ASK-3)"],
    ["驱动方式", "差速驱动(2 主动轮 + 万向轮)", "12 自由度腿足(每条腿 3 关节)"],
    ["运动模型", "非完整约束(纯 v, ω)", "全向(vx, vy, ω 解耦)"],
    ["足迹", "Φ0.22 m 圆", "0.39 × 0.22 m 矩形"],
    ["机身高度", "≈0.18 m", "≈0.40 m"],
    ["最大线速度", "0.22 m/s", "0.30 m/s(行走步态)"],
    ["最大转速", "2.0 rad/s", "0.9 rad/s"],
], col_widths_cm=[3.5, 5.5, 5.5])
CAPTION("表 2.1 轮式与四足平台的关键差异")

H2("2.2 ROS 接口对比")
make_table([
    ["层", "轮式实现", "四足实现", "迁移说明"],
    ["驱动指令", "geometry_msgs/Twist 发布到 /cmd_vel",
     "三路 std_msgs/Float32:\n/ask/dog/forward_back\n/ask/dog/left_right\n/ask/dog/yaw",
     "需要把 (vx, vy, ω) 拆成三路单独的 Float32 发布,\n并且在导航开始前 latch 一次 walk/start 使能"],
    ["里程计", "/odom 来自差速底盘",
     "/odom 来自 mydog_state_estimator(从 /gazebo/model_states 读真值)",
     "保留 nav_msgs/Odometry 接口,\n但 twist 的 linear.y 也需要使用"],
    ["激光感知", "TurtleBot3 LDS01 真实激光 /scan",
     "URDF 无激光雷达,使用 fake_scan_from_map.py 在 /map 上射线投射",
     "保留 LaserScan 接口,\n用静态地图模拟雷达供 DWA 实时避障"],
    ["定位", "AMCL(基于激光)",
     "gt_pose_bridge.py 直接广播 map -> odom 单位变换",
     "仿真简化:真实部署需替换为 AMCL + 实际 LiDAR 或 VIO"],
    ["TF", "map -> odom -> base_footprint",
     "map -> odom -> base",
     "导航器中 base_frame 改为可配置参数"],
], col_widths_cm=[2.0, 4.0, 4.6, 4.0])
CAPTION("表 2.2 ROS 拓扑差异及迁移做法")

# =====================================================================
# 3. 软件包结构
# =====================================================================
H1("3. dog_hybrid_planner 软件包结构")
CODE("dog_hybrid_planner/\n"
     "├── package.xml / CMakeLists.txt\n"
     "├── maps/                   # 同 py_hybrid_planner(map.pgm + yaml)\n"
     "├── worlds/indoor.world     # 同 py_hybrid_planner\n"
     "├── config/dog_nav.rviz     # RViz 配置\n"
     "├── launch/\n"
     "│   ├── dog_indoor_gazebo.launch  # Stage1: Gazebo + ASK-3\n"
     "│   ├── dog_nav.launch            # Stage2: nav stack\n"
     "│   └── dog_all.launch            # 一键启动\n"
     "├── scripts/\n"
     "│   ├── improved_astar.py         # 双向 A* + 24 邻接 + Bezier\n"
     "│   ├── omni_dwa.py               # 全向 DWA(本工作的核心改动)\n"
     "│   ├── dog_navigation.py         # 主导航节点(Float32 接口)\n"
     "│   ├── gt_pose_bridge.py         # map -> odom 桥\n"
     "│   ├── fake_scan_from_map.py     # 静态地图射线投射 -> /scan\n"
     "│   ├── dog_command_bridge.py     # 可选:Twist -> 三路 Float32\n"
     "│   ├── benchmark_dog_astar.py    # 离线 A* 基准\n"
     "│   ├── benchmark_omni_dwa.py     # 离线 DWA 基准\n"
     "│   └── run_navigation_experiment.py # 在线集成实验\n"
     "├── tools/build_dog_report.py     # 本报告自动生成脚本\n"
     "└── experiments/results/          # 实验输出(JSON/PNG/log)")

# =====================================================================
# 4. 算法适配
# =====================================================================
H1("4. 算法适配(从轮式到四足)")
P("本节是本实验的算法核心部分,逐项列出迁移过程中需要改动的地方,"
  "以及为什么必须改、改成了什么。未在本节中出现的模块均完全沿用 "
  "py_hybrid_planner 的实现。")

H2("4.1 全局规划:hybrid A*")
P("ImprovedAStar 在算法层面与轮式版本完全一致(双向搜索 + 动态权重 + "
  "24 邻接 + 三阶贝塞尔平滑),原因是栅格 A* 不依赖于机器人具体的运动学。"
  "迁移过程中只调整了两个参数:")
BULLET("inflate_radius:0.12 m → 0.22 m。"
       "ASK-3 机身对角线半径约 0.225 m,远大于 burger 的 0.11 m,"
       "膨胀半径必须放大以避免规划路径贴墙。")
BULLET("extract_key_points 的 dist_thresh:0.30 m → 0.40 m。"
       "全向四足无最小转弯半径,关键点可以稀疏一些,"
       "DWA 局部目标变化频率更低,跟踪更平滑。")
P("两处改动只动了 ImprovedAStar 的默认参数,核心 bidirectional_astar / "
  "bezier_smooth 函数完全沿用。")

H2("4.2 局部规划:由差速 DWA 改为 OmniDWA")
P("本实验最关键的算法重写。原 ImprovedDWA 的状态空间是 (v, ω),"
  "运动方程假设非完整约束底盘,无法表达四足侧向移动的能力。新文件 "
  "scripts/omni_dwa.py 重写为以下五处差异:")
H3("(1) 状态空间 3 维")
P("原:在 (v, ω) 二维栅格里搜索;dynamic window 包含 v_min..v_max、"
  "ω_min..ω_max。新:增加 vy 维度,采样栅格变为 (vx, vy, ω) "
  "三维。为了控制采样数量,默认每维 7×5×9 = 315 个候选,"
  "加上动态窗口可行域裁剪,每周期评估 ~50 个轨迹,与原版"
  "在数量级上一致。")
CODE("# omni_dwa.py 关键片段\n"
     "for vx in vxs:\n"
     "    for vy in vys:\n"
     "        for w in ws:\n"
     "            traj = self.predict_trajectory(\n"
     "                current_x, current_y, current_theta, vx, vy, w)\n"
     "            ...")

H3("(2) 全向运动学积分")
P("差速底盘的预测公式 x += v·cosθ·dt 不再成立。改为使用机体坐标系下的 "
  "(vx, vy) 投影到世界坐标:")
CODE("def motion(self, x, y, theta, vx, vy, w, dt):\n"
     "    ct, st = math.cos(theta), math.sin(theta)\n"
     "    x += (vx * ct - vy * st) * dt\n"
     "    y += (vx * st + vy * ct) * dt\n"
     "    theta += w * dt\n"
     "    return x, y, theta")

H3("(3) 评分函数:从 heading 改为 progress")
P("差速 DWA 的 heading 评分鼓励机体朝向目标方向。但全向四足可以正面朝某方向"
  "时却向左/右走,heading 评分会错误地惩罚这种最优动作。"
  "OmniDWA 改为 progress 评分:")
P("score = α·progress + β·clearance + γ·speed + δ·heading_bonus")
P("其中 progress = (d0 − dT) / max(d0, 0.5),即一段轨迹"
  "对终点距离的相对缩短量。clearance 与 speed 与原版一致,"
  "δ·heading_bonus 是一个小的偏好项,鼓励机体朝运动方向看齐(便于 "
  "fake_scan 看清楚前方),但权重 0.10 ≪ progress 权重 0.45,"
  "不会强制要求机体追着 heading 走。")

H3("(4) 防卡死建议:侧步代替原地旋转")
P("当 admissible 候选为零时,差速 DWA 给出的恢复建议是「原地旋转」"
  "(纯 ω,因为原地侧移在差速底盘上不可行)。OmniDWA 给出的"
  "建议是「沿障碍法线侧步 0.15 m/s」:")
CODE("info['stuck'] = True\n"
     "world_bearing = math.atan2(oy - cy, ox - cx)\n"
     "local_bearing = self._wrap(world_bearing - current_theta)\n"
     "side_sign = -1.0 if local_bearing > 0 else 1.0\n"
     "return 0.0, 0.15 * side_sign, 0.0, info  # vx=0, vy>0, ω=0")
P("这一改动直接利用了四足平台的全向能力,"
  "对窄走廊、桌椅腿之间等场景比原地旋转更高效。")

H3("(5) 几何参数:足迹与速度上限")
make_table([
    ["参数", "ImprovedDWA(轮式)", "OmniDWA(四足)"],
    ["robot_radius", "0.11 m", "0.23 m"],
    ["hard_radius",  "0.13 m", "0.26 m"],
    ["max_vx",       "0.22 m/s", "0.35 m/s"],
    ["max_vy",       "(无,差速)", "0.20 m/s"],
    ["max_ω",        "2.0 rad/s", "1.0 rad/s"],
    ["max_accel_xy", "2.5 m/s²",  "1.5 m/s²"],
    ["dt / horizon", "0.10 s / 1.2 s", "0.10 s / 1.4 s"],
], col_widths_cm=[3.5, 5.5, 5.5])
CAPTION("表 4.1 OmniDWA 与原 ImprovedDWA 的几何/动力学参数对比")

H2("4.3 顶层导航节点")
P("DogHybridNavigator(scripts/dog_navigation.py)结构与"
  "wheeled HybridNavigator 一致(map_cb / odom_cb / scan_cb / "
  "control_loop / RecoveryFSM),但实现上有两处接口差异:")
BULLET("控制输出:不再发布 /cmd_vel(Twist),改为同时发布 "
       "/ask/dog/forward_back、/ask/dog/left_right、/ask/dog/yaw "
       "三路 Float32。在 /goal_cb 接收到目标点时,会先 latch 一次 "
       "/ask/dog/start = True 与 /ask/dog/walk = True,"
       "通知 mydog_control_sim_ros 切换到 RL 行走步态。")
BULLET("OmniRecoveryFSM:三段恢复 SIDE_STEP → ROTATE → BACK_OFF。"
       "侧步阶段是新加的,作为四足的优先解卡死手段。"
       "原 RecoveryFSM 只有 ROTATE → BACK_OFF。")
P("ProgressMonitor、OscillationMonitor 与拓扑路径跟踪逻辑沿用,"
  "但因为四足 yaw 不严格指向目标,OscillationMonitor 在四足上"
  "几乎不再触发,实测可去掉(本实现中保留以便消融实验)。")

H2("4.4 感知与定位适配")
H3("(1) fake_scan_from_map.py")
P("ASK-3 URDF 不带激光雷达。为了让 OmniDWA 仍然能在 control_loop "
  "中接收 LaserScan(保持对 wheeled 版本的接口一致性),"
  "新增了 scripts/fake_scan_from_map.py 节点。"
  "它在 base 帧每 0.1 s 对 /map 做 360 条 1° 间隔的 1-D DDA 射线投射,"
  "把第一个被占据的栅格作为该方向的距离读数,"
  "并以 sensor_msgs/LaserScan 形式发布到 /scan。"
  "局限:看不见动态障碍。这是个可接受的简化,详见报告 6 节讨论。")
H3("(2) gt_pose_bridge.py")
P("py_hybrid_planner 用 AMCL 把 /scan 与 /map 对齐生成 map -> odom。"
  "ASK-3 没有 LiDAR,无法跑 AMCL。仿真环境下取自 "
  "/gazebo/model_states 即真值,因此 gt_pose_bridge.py 只需以 30 Hz "
  "广播一个单位变换 map -> odom。"
  "mydog_state_estimator 已经把 odom -> base 链路接好,"
  "于是 map -> odom -> base 的完整 TF 链满足导航器要求。")

# =====================================================================
# 5. 实验
# =====================================================================
H1("5. 实验结果")
H2("5.1 离线 A* 基准:dog vs wheeled inflate")

if ASTAR:
    rows = [["场景", "起点 → 终点",
             "Improved 用时 (s)", "Improved 路径 (m)",
             "Improved 关键点", "Traditional 用时 (s)"]]
    for r in ASTAR:
        s, g = r['start'], r['goal']
        im = r['improved']; tm = r['traditional']
        rows.append([
            r['name'],
            f"({s[0]:.2f},{s[1]:.2f}) → ({g[0]:.2f},{g[1]:.2f})",
            f"{im.get('planning_time_s', 0):.3f}",
            f"{im.get('path_length_m', 0):.2f}",
            str(im.get('key_points', '-')),
            f"{tm.get('planning_time_s', 0):.3f}" if tm.get('ok') else "n/a",
        ])
    make_table(rows, col_widths_cm=[3.0, 4.5, 2.6, 2.6, 2.0, 2.6])
    CAPTION("表 5.1 ASK-3 配置(膨胀 0.22 m)下的 A* 基准")
else:
    P("[尚未运行 benchmark_dog_astar.py,占位]")

add_image(os.path.join(RES, "dog_astar_benchmark.png"))
CAPTION("图 5.1 三个测试场景下,改进 A* 与传统 A* 在四足膨胀参数下的路径")

H2("5.2 离线 DWA 基准:OmniDWA vs 差速 DWA")
P("两种 DWA 在同样的全局路径上各跑一次。轮式 DWA 控制仍用单车模型积分;"
  "OmniDWA 用全向积分。统一按照 ASK-3 的足迹半径 0.23 m 判断碰撞。")

if DWA:
    rows = [["场景", "控制器", "到达", "用时(s)", "路径(m)",
             "min_obs(m)", "stuck", "碰撞"]]
    for r in DWA:
        for tag in ('wheeled', 'omni'):
            d = r[tag]
            rows.append([
                r['name'],
                "OmniDWA(四足)" if tag == 'omni' else "ImprovedDWA(轮式)",
                "Yes" if d.get('ok') else "No",
                f"{d.get('arrival_time_s', 0):.1f}",
                f"{d.get('path_length_m', 0):.2f}",
                (f"{d.get('min_obs_dist_m'):.2f}"
                 if d.get('min_obs_dist_m') is not None else "-"),
                str(d.get('stuck_steps', 0)),
                "Yes" if d.get('collision') else "No",
            ])
    make_table(rows, col_widths_cm=[2.5, 4.0, 1.4, 1.6, 1.6, 1.8, 1.4, 1.4])
    CAPTION("表 5.2 OmniDWA 与差速 DWA 离线基准对比")
else:
    P("[尚未运行 benchmark_omni_dwa.py,占位]")

add_image(os.path.join(RES, "dog_dwa_benchmark.png"))
CAPTION("图 5.2 两种 DWA 在相同地图、相同 A* 全局路径下的执行轨迹")

H2("5.3 在线集成实验(Gazebo + dog_all.launch)")
P("在 Gazebo 中启动 dog_all.launch(包含 indoor.world、ASK-3、12 个关节"
  "控制器、mydog_control_sim 控制节点、状态估计器、map_server、"
  "fake_scan_from_map、gt_pose_bridge、dog_hybrid_navigator、RViz),"
  "通过 run_navigation_experiment.py 顺序发送 4 个目标点,"
  "记录 ASK-3 的真实轨迹。")
if INTEG and INTEG.get('per_goal'):
    rows = [["#", "目标 (x, y, yaw)", "到达", "用时 (s)"]]
    for i, g in enumerate(INTEG['per_goal']):
        x, y, yaw = g['goal']
        rows.append([str(i + 1),
                     f"({x:.2f}, {y:.2f}, {yaw:.2f})",
                     "Yes" if g['ok'] else "No",
                     f"{g['time_s']:.1f}"])
    make_table(rows, col_widths_cm=[1.2, 5.0, 1.5, 2.0])
    CAPTION("表 5.3 在线 4 目标点序列实验")
else:
    P("[尚未运行 run_navigation_experiment.py,占位]")
add_image(os.path.join(RES, "dog_integration_trajectory.png"))
CAPTION("图 5.3 在线实验中 ASK-3 的实际轨迹(蓝色),绿色 X = 到达,红色 X = 超时")

# =====================================================================
# 6. 讨论
# =====================================================================
H1("6. 讨论与局限")
H2("6.1 与轮式版本的等效性")
P("除了 4 节列出的算法变更点之外,以下能力与 py_hybrid_planner 完全一致:")
BULLET("全局路径:bidirectional A* + 动态权重 + 24 邻接 + Bezier 平滑。")
BULLET("局部规划:Dynamic Window 评分(只是 score 的具体项不同)。")
BULLET("拓扑跟踪:relevant key point + lookahead 滚动机制。")
BULLET("失效恢复:ProgressMonitor + RecoveryFSM 的两层冗余。")
BULLET("所有可视化:全局路径、关键点、局部目标 marker、TF。")

H2("6.2 局限性")
BULLET("仿真级定位:gt_pose_bridge 直接读 /gazebo/model_states 真值,"
       "实物部署需要重新接入 AMCL 或 VIO。")
BULLET("感知:fake_scan_from_map 只能感知静态地图,"
       "动态障碍(被人推开的椅子等)无法被 DWA 捕捉。"
       "实际部署的 ASK-3 应增装一颗 2D LiDAR(或 D435 深度相机投影)"
       "并在 URDF 中接入 gazebo_ros_laser 插件。")
BULLET("步态层未优化:对 mydog_control_sim 的 RL 策略,"
       "我们只触发 walk/start 一次,后续不再调步频/步幅。"
       "在崎岖地面 (rough.world) 应允许导航器动态修改 stepHeight。")

H2("6.3 未来工作")
BULLET("将 fake_scan 与真实 LiDAR 输出做对照,验证算法对噪声的容忍度。")
BULLET("把 OmniDWA 的 (vx, vy, ω) 输出用 IK 直接下发给 ASK-3 的 RL 步态层,"
       "去掉 mydog_control 的 joystick 桥接。")
BULLET("引入分层规划:在 hybrid A* 之前加一层基于地形高度图的"
       "可通行性筛选,以便切换到 stairs.world / rough.world 等场景。")

# =====================================================================
# 7. 复现指南
# =====================================================================
H1("7. 复现指南(remote: cjx@10.15.205.85)")
P("步骤摘要,详细命令见 README:")
CODE("# 1. 编译\n"
     "cd /home/cjx/catkin_ws && catkin_make -j2\n"
     "source devel/setup.bash\n\n"
     "# 2. 一键启动\n"
     "roslaunch dog_hybrid_planner dog_all.launch\n\n"
     "# 3. 离线基准(可选)\n"
     "rosrun dog_hybrid_planner benchmark_dog_astar.py\n"
     "rosrun dog_hybrid_planner benchmark_omni_dwa.py\n\n"
     "# 4. 在线集成实验\n"
     "rosrun dog_hybrid_planner run_navigation_experiment.py\n\n"
     "# 5. 重生成本报告\n"
     "python3 /home/cjx/catkin_ws/src/dog_hybrid_planner/tools/build_dog_report.py")

doc.save(OUT)
print(f"Report -> {OUT}")
