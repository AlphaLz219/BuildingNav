#!/usr/bin/env python3
"""Generate principle figures and insert them into the ablation report docx."""

from __future__ import annotations

import math
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from matplotlib import font_manager
from matplotlib.patches import Circle, FancyArrowPatch, PathPatch, Polygon, Rectangle
from matplotlib.path import Path as MplPath
from mpl_toolkits.mplot3d import Axes3D  # noqa: F401


REPORT = Path("/media/cjx/D683-C616/四足机器人算法改进消融实验报告.docx")
OUT_DOCX = Path("/tmp/四足机器人算法改进消融实验报告_原理图增强版.docx")
OUT_DIR = Path("/tmp/四足机器人算法原理图_增强版")
FONT = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
FONT_PROP = font_manager.FontProperties(fname=FONT)


COLORS = {
    "bg": "#f7f9fb",
    "grid": "#d8dee8",
    "obs": "#1f2933",
    "path": "#ef4b3f",
    "path2": "#2f80ed",
    "safe": "#ffd166",
    "safe2": "#a8dadc",
    "good": "#2a9d8f",
    "bad": "#d62828",
    "gray": "#8d99ae",
    "blue": "#457b9d",
    "purple": "#7b61ff",
    "orange": "#f4a261",
    "green": "#52b788",
}


def setup_font() -> None:
    if hasattr(font_manager.fontManager, "addfont"):
        font_manager.fontManager.addfont(FONT)
    plt.rcParams["font.family"] = "sans-serif"
    plt.rcParams["font.sans-serif"] = ["DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def save_fig(fig: plt.Figure, path: Path) -> None:
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def clean_axes(ax: plt.Axes, xlim=(0, 10), ylim=(0, 6)) -> None:
    ax.set_xlim(*xlim)
    ax.set_ylim(*ylim)
    ax.set_aspect("equal")
    ax.set_facecolor(COLORS["bg"])
    ax.set_xticks([])
    ax.set_yticks([])
    for spine in ax.spines.values():
        spine.set_visible(False)


def add_grid(ax: plt.Axes, xlim=(0, 10), ylim=(0, 6), step=0.5) -> None:
    xs = np.arange(xlim[0], xlim[1] + 0.001, step)
    ys = np.arange(ylim[0], ylim[1] + 0.001, step)
    for x in xs:
        ax.plot([x, x], ylim, color=COLORS["grid"], lw=0.4, alpha=0.45, zorder=0)
    for y in ys:
        ax.plot(xlim, [y, y], color=COLORS["grid"], lw=0.4, alpha=0.45, zorder=0)


def add_obstacles(ax: plt.Axes, obstacles: list[tuple[float, float, float, float]]) -> None:
    for x, y, w, h in obstacles:
        ax.add_patch(Rectangle((x, y), w, h, color=COLORS["obs"], zorder=4))


def arrow(ax: plt.Axes, p0, p1, color, lw=2.5, ms=16, style="-|>", z=8, alpha=1.0) -> None:
    ax.add_patch(
        FancyArrowPatch(
            p0,
            p1,
            arrowstyle=style,
            mutation_scale=ms,
            lw=lw,
            color=color,
            alpha=alpha,
            zorder=z,
            shrinkA=0,
            shrinkB=0,
        )
    )


def plot_polyline(ax: plt.Axes, points, color, lw=3, ls="-", z=8, alpha=1.0) -> None:
    xs, ys = zip(*points)
    ax.plot(xs, ys, color=color, lw=lw, ls=ls, solid_capstyle="round", solid_joinstyle="round", zorder=z, alpha=alpha)


def robot(ax: plt.Axes, center, yaw, color="#ffffff", edge="#263238", scale=1.0, alpha=1.0) -> None:
    cx, cy = center
    length, width = 0.85 * scale, 0.42 * scale
    corners = np.array(
        [
            [length / 2, width / 2],
            [length / 2, -width / 2],
            [-length / 2, -width / 2],
            [-length / 2, width / 2],
        ]
    )
    rot = np.array([[math.cos(yaw), -math.sin(yaw)], [math.sin(yaw), math.cos(yaw)]])
    pts = corners @ rot.T + np.array([cx, cy])
    ax.add_patch(Polygon(pts, closed=True, facecolor=color, edgecolor=edge, lw=1.8, zorder=9, alpha=alpha))
    nose = np.array([[length * 0.52, 0], [length * 0.25, width * 0.20], [length * 0.25, -width * 0.20]])
    npts = nose @ rot.T + np.array([cx, cy])
    ax.add_patch(Polygon(npts, closed=True, facecolor=COLORS["orange"], edgecolor="none", zorder=10, alpha=alpha))


def title(ax: plt.Axes, text: str) -> None:
    ax.text(
        0.02,
        0.98,
        text,
        transform=ax.transAxes,
        va="top",
        ha="left",
        fontsize=17,
        weight="bold",
        color="#17212b",
        fontproperties=FONT_PROP,
    )


def small_label(ax: plt.Axes, xy, text: str, color="#17212b", size=12, ha="center") -> None:
    ax.text(*xy, text, fontsize=size, color=color, ha=ha, va="center", zorder=20, fontproperties=FONT_PROP)


def draw_fig_01(path: Path) -> None:
    fig, axes = plt.subplots(1, 2, figsize=(13, 5.1))
    obstacles = [(2.1, 1.0, 0.3, 3.0), (4.2, 0.7, 0.35, 2.0), (4.2, 3.7, 0.35, 1.1), (6.0, 2.0, 1.1, 0.32), (7.4, 3.0, 0.35, 1.5)]
    start, goal = (0.8, 0.7), (9.1, 5.0)
    for ax in axes:
        clean_axes(ax)
        add_grid(ax)
        add_obstacles(ax, obstacles)
        ax.scatter(*start, s=115, color="#10a23a", zorder=15)
        ax.scatter(*goal, s=170, marker="*", color="red", zorder=15)
    title(axes[0], "传统A星")
    title(axes[1], "改进A星")
    path1 = [
        (0.8, 0.7),
        (1.45, 1.35),
        (1.65, 2.0),
        (1.65, 4.3),
        (2.75, 4.3),
        (3.65, 3.35),
        (4.05, 3.25),
        (5.20, 3.65),
        (6.10, 4.75),
        (7.15, 4.85),
        (8.2, 5.0),
        (9.1, 5.0),
    ]
    path2 = [
        (0.8, 0.7),
        (1.50, 1.50),
        (1.50, 4.35),
        (2.80, 4.35),
        (4.00, 3.30),
        (5.55, 3.40),
        (6.70, 4.72),
        (8.00, 4.88),
        (9.1, 5.0),
    ]
    for p in path1:
        ax = axes[0]
        ax.add_patch(Circle(p, 0.23, edgecolor="#f2c94c", facecolor="#f2c94c", alpha=0.18, lw=1, zorder=2))
    for p in path2:
        ax = axes[1]
        ax.add_patch(Circle(p, 0.25, edgecolor="#f2c94c", facecolor="#f2c94c", alpha=0.22, lw=1, zorder=2))
    plot_polyline(axes[0], path1, COLORS["path"], lw=3.4)
    plot_polyline(axes[1], path2, COLORS["path"], lw=3.4)
    arrow(axes[1], start, (1.50, 1.50), COLORS["path2"], lw=2.2, alpha=0.85)
    arrow(axes[1], goal, (8.00, 4.88), COLORS["path2"], lw=2.2, alpha=0.85)
    small_label(axes[1], (5.2, 5.55), "双向搜索 + 平滑", COLORS["blue"], 13)
    save_fig(fig, path)


def draw_fig_02(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    clean_axes(ax)
    add_grid(ax)
    obstacles = [(3.1, 0.5, 0.45, 4.7), (6.4, 0.8, 0.45, 4.3), (0.8, 4.8, 1.4, 0.35), (8.1, 0.7, 1.0, 0.35)]
    add_obstacles(ax, obstacles)
    for x, y, w, h in obstacles:
        ax.add_patch(Rectangle((x - 0.35, y - 0.35), w + 0.7, h + 0.7, facecolor=COLORS["safe"], edgecolor="none", alpha=0.25, zorder=1))
        ax.add_patch(Rectangle((x - 0.65, y - 0.65), w + 1.3, h + 1.3, facecolor=COLORS["safe2"], edgecolor="none", alpha=0.22, zorder=0))
    unsafe = [(0.7, 0.9), (2.3, 1.45), (3.0, 1.65)]
    safe = [(0.7, 0.9), (1.7, 2.2), (2.7, 5.55), (4.5, 5.55), (6.9, 5.55), (7.9, 5.25), (8.8, 4.8)]
    plot_polyline(ax, unsafe, COLORS["bad"], lw=2.5, ls="--", alpha=0.7)
    plot_polyline(ax, safe, COLORS["good"], lw=4)
    ax.scatter(3.0, 1.65, s=150, marker="x", color=COLORS["bad"], linewidths=3, zorder=18)
    robot(ax, (4.5, 5.55), 0.0, scale=0.95)
    ax.scatter([0.7], [0.9], s=120, color="#10a23a", zorder=15)
    ax.scatter([8.8], [4.8], s=170, marker="*", color="red", zorder=15)
    title(ax, "足迹膨胀与安全软代价")
    small_label(ax, (4.2, 5.55), "硬边界", COLORS["orange"], 12)
    small_label(ax, (5.35, 5.55), "软代价区", COLORS["blue"], 12)
    save_fig(fig, path)


def draw_fig_03(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    clean_axes(ax)
    add_grid(ax)
    obstacles = [(3.0, 1.0, 0.35, 3.4), (3.0, 1.0, 2.8, 0.35), (6.5, 2.1, 0.35, 2.9), (4.9, 4.65, 2.0, 0.35)]
    add_obstacles(ax, obstacles)
    a, b = (2.2, 0.8), (3.0, 1.25)
    c, d = (2.2, 0.8), (2.2, 4.85)
    f = (4.55, 5.35)
    plot_polyline(ax, [a, b], COLORS["bad"], lw=4, ls="--", alpha=0.85)
    plot_polyline(ax, [c, d, f, (8.7, 5.35)], COLORS["good"], lw=4)
    ax.scatter(*a, s=115, color="#10a23a", zorder=15)
    ax.scatter(8.7, 5.35, s=170, marker="*", color="red", zorder=15)
    for p in [a, b]:
        ax.scatter(*p, s=52, color=COLORS["path2"], zorder=16)
    ax.add_patch(Circle((3.0, 1.0), 0.5, edgecolor=COLORS["bad"], facecolor="none", lw=2.5, zorder=12))
    ax.scatter(3.0, 1.25, s=150, marker="x", color=COLORS["bad"], linewidths=3, zorder=18)
    title(ax, "线段安全检测")
    small_label(ax, (3.9, 1.25), "穿角", COLORS["bad"], 13)
    small_label(ax, (4.0, 5.45), "逐格验线", COLORS["good"], 13)
    save_fig(fig, path)


def draw_fig_04(path: Path) -> None:
    fig, axes = plt.subplots(1, 2, figsize=(13, 5.1))
    for ax in axes:
        clean_axes(ax)
        add_grid(ax)
        add_obstacles(ax, [(2.2, 2.65, 2.5, 0.35), (5.6, 3.10, 2.6, 0.35), (4.65, 0.75, 0.35, 1.35), (5.05, 4.05, 0.35, 1.20)])
        ax.scatter(0.8, 1.2, s=115, color="#10a23a", zorder=15)
        ax.scatter(9.0, 4.7, s=170, marker="*", color="red", zorder=15)
    title(axes[0], "无转角惩罚")
    title(axes[1], "加入转角惩罚")
    p1 = [
        (0.8, 1.2),
        (1.35, 1.45),
        (1.85, 1.35),
        (2.45, 1.65),
        (3.10, 1.55),
        (3.85, 2.05),
        (4.55, 2.40),
        (5.15, 2.95),
        (5.25, 3.75),
        (6.05, 3.75),
        (6.75, 4.08),
        (7.45, 4.18),
        (8.25, 4.50),
        (9.0, 4.7),
    ]
    p2 = [(0.8, 1.2), (3.7, 2.05), (4.95, 2.75), (5.35, 3.72), (7.05, 4.12), (9.0, 4.7)]
    plot_polyline(axes[0], p1, COLORS["path"], lw=3.3)
    plot_polyline(axes[1], p2, COLORS["good"], lw=4)
    for p in p1[1:-1]:
        axes[0].scatter(*p, s=22, color=COLORS["orange"], zorder=14)
    for p in p2[1:-1]:
        axes[1].scatter(*p, s=42, color=COLORS["good"], zorder=14)
    small_label(axes[0], (5.0, 5.45), "频繁小转向", COLORS["bad"], 13)
    small_label(axes[1], (5.0, 5.45), "少转角", COLORS["good"], 13)
    save_fig(fig, path)


def bezier(p0, p1, p2, p3, n=60):
    t = np.linspace(0, 1, n)[:, None]
    return (1 - t) ** 3 * p0 + 3 * (1 - t) ** 2 * t * p1 + 3 * (1 - t) * t**2 * p2 + t**3 * p3


def draw_fig_05(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    clean_axes(ax)
    add_grid(ax)
    add_obstacles(ax, [(2.6, 1.1, 0.35, 2.75), (5.8, 2.15, 1.7, 0.35), (6.8, 3.0, 0.35, 1.35)])
    raw = np.array(
        [
            (0.8, 0.8),
            (1.25, 1.25),
            (1.75, 2.0),
            (2.15, 3.0),
            (2.20, 4.35),
            (3.40, 5.15),
            (4.65, 4.95),
            (5.60, 4.75),
            (6.70, 4.95),
            (7.60, 5.05),
            (8.9, 5.0),
        ]
    )
    pruned = [(0.8, 0.8), (2.20, 4.35), (5.60, 4.75), (8.9, 5.0)]
    curve = bezier(np.array([0.8, 0.8]), np.array([1.75, 4.35]), np.array([5.8, 5.45]), np.array([8.9, 5.0]), 90)
    plot_polyline(ax, raw, COLORS["gray"], lw=2.2, ls="--", alpha=0.75)
    plot_polyline(ax, pruned, COLORS["orange"], lw=2.7, alpha=0.8)
    ax.plot(curve[:, 0], curve[:, 1], color=COLORS["good"], lw=4, solid_capstyle="round", zorder=10)
    sample_idx = np.linspace(0, len(curve) - 1, 13).astype(int)
    ax.scatter(curve[sample_idx, 0], curve[sample_idx, 1], s=34, color="white", edgecolor=COLORS["good"], lw=1.6, zorder=14)
    ax.scatter(0.8, 0.8, s=115, color="#10a23a", zorder=15)
    ax.scatter(8.9, 5.0, s=170, marker="*", color="red", zorder=15)
    title(ax, "剪枝 + 圆角平滑 + 重采样")
    small_label(ax, (2.0, 5.45), "原始点", COLORS["gray"], 12)
    small_label(ax, (4.05, 5.45), "关键点", COLORS["orange"], 12)
    small_label(ax, (6.3, 5.45), "等距参考线", COLORS["good"], 12)
    save_fig(fig, path)


def draw_fig_06(path: Path) -> None:
    fig = plt.figure(figsize=(11, 5.5))
    ax = fig.add_subplot(121, projection="3d")
    ax.set_facecolor(COLORS["bg"])
    ax.view_init(elev=22, azim=-42)
    vx = np.linspace(0.0, 0.8, 5)
    vy = np.linspace(-0.35, 0.35, 5)
    wz = np.linspace(-0.8, 0.8, 5)
    pts = np.array(np.meshgrid(vx, vy, wz)).reshape(3, -1).T
    score = pts[:, 0] - 0.35 * np.abs(pts[:, 1]) - 0.12 * np.abs(pts[:, 2])
    ax.scatter(pts[:, 0], pts[:, 1], pts[:, 2], c=score, cmap="viridis", s=18, alpha=0.8)
    ax.set_xlabel("$v_x$")
    ax.set_ylabel("$v_y$")
    ax.set_zlabel("$\\omega$")
    ax.set_title("三维速度采样", fontsize=16, pad=12, fontproperties=FONT_PROP)
    ax.grid(True, alpha=0.2)
    ax2 = fig.add_subplot(122)
    clean_axes(ax2)
    add_grid(ax2)
    add_obstacles(ax2, [(4.9, 1.5, 0.5, 2.0), (6.8, 4.0, 1.4, 0.35)])
    robot(ax2, (2.0, 2.4), 0.0, scale=1.1)
    colors = [COLORS["bad"], COLORS["orange"], COLORS["good"], COLORS["path2"], COLORS["purple"]]
    paths = [
        [(2.35, 2.4), (3.2, 2.4), (4.65, 2.35)],
        [(2.35, 2.4), (3.0, 3.05), (4.2, 3.65), (5.7, 3.75), (6.55, 3.85)],
        [(2.35, 2.4), (3.1, 1.85), (4.25, 1.30), (5.75, 1.25), (7.2, 2.4)],
        [(2.35, 2.4), (3.2, 2.85), (4.4, 3.55), (5.65, 3.85), (6.55, 3.90)],
        [(2.35, 2.4), (3.0, 1.65), (4.2, 1.05), (5.4, 1.05)],
    ]
    for c, p in zip(colors, paths):
        plot_polyline(ax2, p, c, lw=2.5, alpha=0.85)
    arrow(ax2, (2.0, 2.4), (3.1, 2.4), COLORS["good"], lw=2.2)
    arrow(ax2, (2.0, 2.4), (2.0, 3.3), COLORS["path2"], lw=2.2)
    ax2.text(2.92, 2.18, "$v_x$", fontsize=13, color=COLORS["good"])
    ax2.text(1.68, 3.15, "$v_y$", fontsize=13, color=COLORS["path2"])
    title(ax2, "候选轨迹")
    save_fig(fig, path)


def draw_fig_07(path: Path) -> None:
    fig, axes = plt.subplots(1, 2, figsize=(13, 5.1))
    for ax in axes:
        clean_axes(ax)
        add_grid(ax)
        add_obstacles(ax, [(4.4, 0.6, 0.4, 1.8), (4.4, 3.5, 0.4, 1.8)])
        plot_polyline(ax, [(0.8, 3.0), (2.7, 3.0), (4.1, 3.0), (5.6, 3.0), (8.9, 3.0)], COLORS["path2"], lw=2.8, ls="--", alpha=0.65)
    title(axes[0], "无头部偏好")
    title(axes[1], "头部对齐优先")
    robot(axes[0], (2.3, 2.3), math.pi / 2, scale=1.1)
    plot_polyline(axes[0], [(2.3, 2.3), (2.3, 3.0), (3.3, 3.0), (4.15, 3.0)], COLORS["bad"], lw=3.5)
    arrow(axes[0], (2.3, 2.3), (2.3, 3.0), COLORS["bad"], lw=2.5)
    small_label(axes[0], (3.3, 4.8), "侧移优先", COLORS["bad"], 13)
    robot(axes[1], (2.3, 2.3), math.pi / 2, scale=1.1, alpha=0.45)
    arrow(axes[1], (2.3, 2.3), (2.3, 3.05), COLORS["orange"], lw=2.5)
    robot(axes[1], (3.15, 3.0), 0.0, scale=1.1)
    plot_polyline(axes[1], [(3.15, 3.0), (4.2, 3.0), (5.6, 3.0), (8.9, 3.0)], COLORS["good"], lw=4)
    small_label(axes[1], (3.7, 4.8), "先对齐 再前进", COLORS["good"], 13)
    save_fig(fig, path)


def draw_fig_08(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    clean_axes(ax)
    add_grid(ax)
    add_obstacles(ax, [(2.2, 1.0, 0.35, 3.4), (4.9, 2.55, 2.0, 0.35), (7.6, 0.9, 0.35, 3.3)])
    global_path = np.array(
        [
            (0.8, 0.8),
            (1.35, 2.20),
            (1.75, 4.75),
            (2.55, 5.30),
            (4.00, 5.20),
            (5.75, 5.00),
            (7.10, 4.72),
            (8.05, 4.62),
            (9.0, 4.9),
        ]
    )
    ax.plot(global_path[:, 0], global_path[:, 1], color=COLORS["path2"], lw=16, alpha=0.18, solid_capstyle="round", zorder=1)
    ax.plot(global_path[:, 0], global_path[:, 1], color=COLORS["path2"], lw=2.6, ls="--", alpha=0.75, zorder=3)
    local = [(1.1, 1.0), (1.35, 2.15), (1.75, 4.65), (2.65, 5.08), (4.10, 4.98), (5.20, 4.82)]
    plot_polyline(ax, local, COLORS["good"], lw=4)
    for idx, p in enumerate(global_path[1:6], start=1):
        ax.scatter(*p, s=70, facecolor="white", edgecolor=COLORS["path2"], lw=2, zorder=12)
        small_label(ax, (p[0], p[1] + 0.32), f"{idx}", COLORS["path2"], 10)
    robot(ax, (1.35, 2.15), 1.15, scale=0.95)
    ax.scatter(0.8, 0.8, s=115, color="#10a23a", zorder=15)
    ax.scatter(9.0, 4.9, s=170, marker="*", color="red", zorder=15)
    arrow(ax, (2.0, 1.0), (2.9, 1.8), COLORS["orange"], lw=2.5, alpha=0.85)
    arrow(ax, (6.1, 5.0), (6.9, 4.25), COLORS["orange"], lw=2.5, alpha=0.85)
    title(ax, "A星参考走廊与单调目标")
    small_label(ax, (4.4, 5.55), "走廊约束", COLORS["path2"], 13)
    small_label(ax, (7.3, 5.55), "索引只前进", COLORS["orange"], 13)
    save_fig(fig, path)


def draw_fig_09(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    clean_axes(ax, xlim=(0, 12), ylim=(0, 6))
    xs = np.linspace(0.9, 11.2, 120)
    speed_hold = 4.75 + 0.06 * np.sin(xs * 1.4)
    speed_raw = 4.75 - 1.15 * (1.0 - np.exp(-0.24 * (xs - 0.9)))
    ax.plot(xs, speed_raw, color=COLORS["gray"], lw=2.4, ls="--", alpha=0.85)
    ax.plot(xs, speed_hold, color=COLORS["good"], lw=3.4)
    ax.text(0.7, 5.55, "速度保持", fontsize=17, weight="bold", color="#17212b", va="top", fontproperties=FONT_PROP)
    ax.text(9.55, 5.05, "保持", fontsize=12, color=COLORS["good"], fontproperties=FONT_PROP)
    ax.text(9.55, 3.78, "衰减", fontsize=12, color=COLORS["gray"], fontproperties=FONT_PROP)
    ax.plot([0.9, 11.3], [2.6, 2.6], color="#d5dbe5", lw=2)
    blocks = [
        (1.1, 1.65, 1.7, 0.78, "DWA"),
        (3.6, 1.65, 1.7, 0.78, "可行轨迹"),
        (6.1, 1.65, 1.7, 0.78, "进展"),
        (8.6, 1.65, 1.7, 0.78, "正常控制"),
    ]
    for x, y, w, h, txt in blocks:
        ax.add_patch(Rectangle((x, y), w, h, facecolor="white", edgecolor=COLORS["blue"], lw=2, zorder=5))
        ax.text(
            x + w / 2,
            y + h / 2,
            txt,
            fontsize=13,
            color="#17212b",
            ha="center",
            va="center",
            zorder=6,
            fontproperties=FONT_PROP,
        )
    for x in [2.8, 5.3, 7.8]:
        arrow(ax, (x, 2.04), (x + 0.55, 2.04), COLORS["blue"], lw=2.3, ms=14)
    ax.add_patch(Rectangle((4.4, 0.55), 3.2, 0.6, facecolor="#fff3cd", edgecolor=COLORS["orange"], lw=1.6, zorder=5))
    ax.text(6.0, 0.85, "仅在失效时Recovery", fontsize=13, color="#8a5a00", ha="center", va="center", zorder=6, fontproperties=FONT_PROP)
    arrow(ax, (6.0, 1.65), (6.0, 1.16), COLORS["orange"], lw=2.2, ms=13)
    ax.text(1.0, 0.72, "后退", fontsize=11, color=COLORS["orange"], fontproperties=FONT_PROP)
    ax.text(8.5, 0.72, "侧步/小转角", fontsize=11, color=COLORS["orange"], fontproperties=FONT_PROP)
    save_fig(fig, path)


FIGURES = [
    ("1.1 基础A星与参考论文改进A星对比", "图1-1 基础A星与参考论文改进A星搜索策略原理图", draw_fig_01),
    ("2.1 足迹膨胀与安全距离软代价", "图2-1 足迹膨胀与安全距离软代价原理图", draw_fig_02),
    ("2.2 线段安全检测", "图2-2 线段安全检测原理图", draw_fig_03),
    ("2.3 转角惩罚", "图2-3 转角惩罚原理图", draw_fig_04),
    ("2.4 安全剪枝、圆角平滑与重采样", "图2-4 安全剪枝、圆角平滑与重采样原理图", draw_fig_05),
    ("2.5 DWA三维速度空间", "图2-5 DWA三维速度空间原理图", draw_fig_06),
    ("2.6 前向姿态偏好与DWA路径头部对齐", "图2-6 前向姿态偏好与路径头部对齐原理图", draw_fig_07),
    ("2.7 A星参考走廊与局部目标单调推进", "图2-7 A星参考走廊与局部目标单调推进原理图", draw_fig_08),
    ("2.8 速度保持与受限Recovery", "图2-8 速度保持与受限Recovery原理图", draw_fig_09),
]


def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def insert_figure_after_heading(doc: Document, heading: str, image_path: Path, caption: str) -> bool:
    for para in doc.paragraphs:
        if para.text.strip() == heading:
            pic_para = insert_paragraph_after(para)
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_para.add_run()
            run.add_picture(str(image_path), width=Inches(5.95))

            cap_para = insert_paragraph_after(pic_para, caption)
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap_para.runs:
                run.font.size = Pt(10.5)
                run.font.name = "宋体"
            return True
    return False


def main() -> None:
    setup_font()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for old_png in OUT_DIR.glob("*.png"):
        old_png.unlink()

    generated = []
    for idx, (_, caption, drawer) in enumerate(FIGURES, start=1):
        fig_path = OUT_DIR / f"{idx:02d}_{caption}.png"
        drawer(fig_path)
        generated.append(fig_path)

    doc = Document(str(REPORT))
    missing = []
    for (heading, caption, _), fig_path in zip(FIGURES, generated):
        if not insert_figure_after_heading(doc, heading, fig_path, caption):
            missing.append(heading)

    if missing:
        raise RuntimeError("未找到以下标题，未能插图：" + "；".join(missing))

    doc.save(str(OUT_DOCX))
    print(OUT_DOCX)
    print(OUT_DIR)


if __name__ == "__main__":
    main()
