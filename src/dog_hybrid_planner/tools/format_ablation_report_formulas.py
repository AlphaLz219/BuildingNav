#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Render plain-text formulas in the ablation report as clear math images.

The source report uses backtick-wrapped formula lines. This script replaces
those formula-only paragraphs with centered high-resolution math renderings,
while keeping all Chinese explanations, figures and tables unchanged.
"""

from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image


INPUT = Path("/tmp/四足机器人算法改进消融实验报告_work.docx")
OUTPUT = Path("/tmp/四足机器人算法改进消融实验报告_公式版.docx")
IMG_DIR = Path("/tmp/ablation_formula_images")


FORMULA_LATEX = {
    "f(n)=g(n)+h(n)": r"f(n)=g(n)+h(n)",
    "f(n)=g(n)+w(n)h(n)": r"f(n)=g(n)+w(n)h(n)",
    "w(n)=1+lambda_d * d(n,goal)/d(start,goal)": r"w(n)=1+\lambda_d\frac{d(n,\mathrm{goal})}{d(\mathrm{start},\mathrm{goal})}",
    "F_f(n)=g_f(n)+w_f(n)h_f(n), F_b(n)=g_b(n)+w_b(n)h_b(n)": r"F_f(n)=g_f(n)+w_f(n)h_f(n),\quad F_b(n)=g_b(n)+w_b(n)h_b(n)",
    "c(n_i,n_j)=sqrt((x_i-x_j)^2+(y_i-y_j)^2)": r"c(n_i,n_j)=\sqrt{(x_i-x_j)^2+(y_i-y_j)^2}",
    "O_inflated={p | d_obs(p) <= r_body + r_margin}": r"O_{\mathrm{inflated}}=\{p\mid d_{\mathrm{obs}}(p)\leq r_{\mathrm{body}}+r_{\mathrm{margin}}\}",
    "C_clear(p)=lambda_c * (1 - d_obs(p)/r_safe)^2, 0 < d_obs(p) < r_safe": r"C_{\mathrm{clear}}(p)=\lambda_c\left(1-\frac{d_{\mathrm{obs}}(p)}{r_{\mathrm{safe}}}\right)^2,\quad 0<d_{\mathrm{obs}}(p)<r_{\mathrm{safe}}",
    "C_clear(p)=0, d_obs(p) >= r_safe": r"C_{\mathrm{clear}}(p)=0,\quad d_{\mathrm{obs}}(p)\geq r_{\mathrm{safe}}",
    "g(n_j)=g(n_i)+c(n_i,n_j)*(1+C_clear(n_j))": r"g(n_j)=g(n_i)+c(n_i,n_j)\left(1+C_{\mathrm{clear}}(n_j)\right)",
    "p(u)=(1-u)p_i+u p_j, u in [0,1]": r"p(u)=(1-u)p_i+u p_j,\quad u\in[0,1]",
    "Safe(e)=1 <=> M(round(p(u)))=0, forall u in [0,1]": r"\mathrm{Safe}(e)=1\Leftrightarrow M(\mathrm{round}(p(u)))=0,\quad \forall u\in[0,1]",
    "u_k=k/K, k=0,1,...,K": r"u_k=\frac{k}{K},\quad k=0,1,\ldots,K",
    "K=ceil(L(e)/Delta_s)": r"K=\left\lceil\frac{L(e)}{\Delta_s}\right\rceil",
    "d_obs(p(u_k)) > r_check, k=0,1,...,K": r"d_{\mathrm{obs}}(p(u_k))>r_{\mathrm{check}},\quad k=0,1,\ldots,K",
    "N_risk=sum I(Safe(e_m)=0)": r"N_{\mathrm{risk}}=\sum_m \mathrm{I}\left(\mathrm{Safe}(e_m)=0\right)",
    "v_1=p_k-p_{k-1}, v_2=p_{k+1}-p_k": r"\mathbf{v}_1=p_k-p_{k-1},\quad \mathbf{v}_2=p_{k+1}-p_k",
    "Delta_theta=acos((v_1 dot v_2)/(|v_1||v_2|))": r"\Delta\theta=\arccos\left(\frac{\mathbf{v}_1\cdot\mathbf{v}_2}{\left\|\mathbf{v}_1\right\|\left\|\mathbf{v}_2\right\|}\right)",
    "C_turn=lambda_theta*(1-cos(Delta_theta))": r"C_{\mathrm{turn}}=\lambda_\theta\left(1-\cos(\Delta\theta)\right)",
    "g(n_{k+1})=g(n_k)+c(n_k,n_{k+1})+C_turn": r"g(n_{k+1})=g(n_k)+c(n_k,n_{k+1})+C_{\mathrm{turn}}",
    "Theta_sum=sum |wrap(theta_{i+1}-theta_i)|": r"\Theta_{\mathrm{sum}}=\sum_i\left|\mathrm{wrap}\left(\theta_{i+1}-\theta_i\right)\right|",
    "LineFree(p_i,p_j)=1": r"\mathrm{LineFree}(p_i,p_j)=1",
    "j*=max{j | j>i and LineFree(p_i,p_j)=1 and d_obs(line(p_i,p_j))>r_prune}": r"j^*=\max\{j\mid j>i,\ \mathrm{LineFree}(p_i,p_j)=1,\ d_{\mathrm{obs}}(\mathrm{line}(p_i,p_j))>r_{\mathrm{prune}}\}",
    "q_i=0.75p_i+0.25p_{i+1}": r"q_i=0.75p_i+0.25p_{i+1}",
    "r_i=0.25p_i+0.75p_{i+1}": r"r_i=0.25p_i+0.75p_{i+1}",
    "P_smooth accepted <=> Safe(P_smooth)=1": r"P_{\mathrm{smooth}}\ \mathrm{accepted}\Leftrightarrow \mathrm{Safe}(P_{\mathrm{smooth}})=1",
    "s_k=k*Delta_l, k=0,1,...,floor(L/Delta_l)": r"s_k=k\Delta_l,\quad k=0,1,\ldots,\left\lfloor\frac{L}{\Delta_l}\right\rfloor",
    "V_2={(v_x, omega)}": r"V_2=\{(v_x,\omega)\}",
    "V_3={(v_x,v_y,omega) | v_x in [v_x_min,v_x_max], v_y in [v_y_min,v_y_max], omega in [omega_min,omega_max]}": r"V_3=\{(v_x,v_y,\omega)\mid v_x\in[v_{x,\min},v_{x,\max}],\ v_y\in[v_{y,\min},v_{y,\max}],\ \omega\in[\omega_{\min},\omega_{\max}]\}",
    "v_x in [v_x^t-a_x Delta_t, v_x^t+a_x Delta_t]": r"v_x\in[v_x^t-a_x\Delta t,\ v_x^t+a_x\Delta t]",
    "v_y in [v_y^t-a_y Delta_t, v_y^t+a_y Delta_t]": r"v_y\in[v_y^t-a_y\Delta t,\ v_y^t+a_y\Delta t]",
    "omega in [omega^t-alpha Delta_t, omega^t+alpha Delta_t]": r"\omega\in[\omega^t-\alpha\Delta t,\ \omega^t+\alpha\Delta t]",
    "x_{t+Delta_t}=x_t+(v_x cos theta - v_y sin theta)Delta_t": r"x_{t+\Delta t}=x_t+(v_x\cos\theta-v_y\sin\theta)\Delta t",
    "y_{t+Delta_t}=y_t+(v_x sin theta + v_y cos theta)Delta_t": r"y_{t+\Delta t}=y_t+(v_x\sin\theta+v_y\cos\theta)\Delta t",
    "theta_{t+Delta_t}=theta_t+omega Delta_t": r"\theta_{t+\Delta t}=\theta_t+\omega\Delta t",
    "G(v)=w_h H(v)+w_c C(v)+w_v V(v)+w_p P(v)-w_s S(v)": r"G(v)=w_hH(v)+w_cC(v)+w_vV(v)+w_pP(v)-w_sS(v)",
    "psi_traj=atan2(y_T-y_0, x_T-x_0)": r"\psi_{\mathrm{traj}}=\mathrm{atan2}(y_T-y_0,\ x_T-x_0)",
    "e_psi=wrap(psi_traj-theta)": r"e_\psi=\mathrm{wrap}(\psi_{\mathrm{traj}}-\theta)",
    "C_forward=lambda_psi |e_psi| + lambda_y |v_y|/v_ymax": r"C_{\mathrm{forward}}=\lambda_\psi|e_\psi|+\lambda_y\frac{|v_y|}{v_{y,\max}}",
    "|e_psi| > theta_align and d_free > d_align": r"|e_\psi|>\theta_{\mathrm{align}},\quad d_{\mathrm{free}}>d_{\mathrm{align}}",
    "v_x=0, v_y=0, omega=clip(k_psi e_psi, -omega_max, omega_max)": r"v_x=0,\quad v_y=0,\quad \omega=\mathrm{clip}(k_\psi e_\psi,-\omega_{\max},\omega_{\max})",
    "P_g={p_0,p_1,...,p_N}": r"P_g=\{p_0,p_1,\ldots,p_N\}",
    "T={q_0,q_1,...,q_M}": r"T=\{q_0,q_1,\ldots,q_M\}",
    "d_corr(T,P_g)=1/M * sum_{m=1}^{M} min_i ||q_m-p_i||": r"d_{\mathrm{corr}}(T,P_g)=\frac{1}{M}\sum_{m=1}^{M}\min_i\left\|q_m-p_i\right\|",
    "G'(v)=G(v)-lambda_corr d_corr(T(v),P_g)": r"G'(v)=G(v)-\lambda_{\mathrm{corr}}d_{\mathrm{corr}}(T(v),P_g)",
    "k_t=max(k_{t-1}, argmin_{k in [k_{t-1}, k_{t-1}+K]} ||p_k-x_t||)": r"k_t=\max\left(k_{t-1},\ \mathrm{argmin}_{k\in[k_{t-1},k_{t-1}+K]}\left\|p_k-x_t\right\|\right)",
    "p_goal^local=p_{min(k_t+L_look, N)}": r"p_{\mathrm{goal}}^{\mathrm{local}}=p_{\min(k_t+L_{\mathrm{look}},N)}",
    "v_ref=alpha_v v_cmd^{t-1}+(1-alpha_v)v_odom^t": r"v_{\mathrm{ref}}=\alpha_v v_{\mathrm{cmd}}^{t-1}+(1-\alpha_v)v_{\mathrm{odom}}^t",
    "V_dw=[v_ref-a_max Delta_t, v_ref+a_max Delta_t]": r"V_{\mathrm{dw}}=[v_{\mathrm{ref}}-a_{\max}\Delta t,\ v_{\mathrm{ref}}+a_{\max}\Delta t]",
    "Recovery=1 <=> (N_feasible=0) or (d_front<d_stop and Delta s<T_prog)": r"\mathrm{Recovery}=1\Leftrightarrow (N_{\mathrm{feasible}}=0)\ \mathrm{or}\ (d_{\mathrm{front}}<d_{\mathrm{stop}}\ \mathrm{and}\ \Delta s<T_{\mathrm{prog}})",
    "u_rec in {(-v_back,0,0), (0,0,omega_rec), (0,v_side,0)}": r"u_{\mathrm{rec}}\in\{(-v_{\mathrm{back}},0,0),(0,0,\omega_{\mathrm{rec}}),(0,v_{\mathrm{side}},0)\}",
}


def render_formula(latex: str, index: int) -> Path:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    out = IMG_DIR / f"formula_{index:02d}.png"

    fig = plt.figure(figsize=(8.0, 0.70), dpi=300)
    fig.patch.set_alpha(0.0)
    fig.text(0.5, 0.50, f"${latex}$", ha="center", va="center", fontsize=18)
    fig.savefig(out, dpi=300, bbox_inches="tight", pad_inches=0.05, transparent=True)
    plt.close(fig)
    return out


def image_width_inches(image_path: Path) -> float:
    with Image.open(image_path) as im:
        px_w, _ = im.size
    natural = px_w / 300.0
    return min(5.85, max(1.35, natural * 1.08))


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def main() -> None:
    doc = Document(INPUT)
    replaced = 0
    missing: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not (text.startswith("`") and text.endswith("`")):
            continue
        source = text[1:-1]
        latex = FORMULA_LATEX.get(source)
        if latex is None:
            missing.append(source)
            continue

        img = render_formula(latex, replaced + 1)
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run()
        run.add_picture(str(img), width=Inches(image_width_inches(img)))
        replaced += 1

    if missing:
        raise RuntimeError("Unmapped formula(s):\n" + "\n".join(missing))

    doc.save(OUTPUT)
    print(f"replaced={replaced}")
    print(OUTPUT)


if __name__ == "__main__":
    main()
