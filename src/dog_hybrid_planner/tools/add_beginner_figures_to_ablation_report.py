#!/usr/bin/env python3
"""Generate beginner-friendly principle figures for the ablation report."""

from __future__ import annotations

import math
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from matplotlib import font_manager
from matplotlib.patches import Circle, FancyArrowPatch, Polygon, Rectangle


REPORT = Path("/media/cjx/D683-C616/四足机器人算法改进消融实验报告.docx")
OUT_DOCX = Path("/tmp/四足机器人算法改进消融实验报告_通俗原理图版.docx")
OUT_DIR = Path("/tmp/四足机器人算法原理图_通俗版")
FONT = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
FONT_PROP = font_manager.FontProperties(fname=FONT)


C = {
    "bg": "#f7f9fc",
    "grid": "#dde5ef",
    "wall": "#1f2933",
    "bad": "#e53935",
    "good": "#1f9d8a",
    "blue": "#2f80ed",
    "orange": "#f4a261",
    "yellow": "#ffd166",
    "soft": "#cdeef1",
    "safe": "#fff3bf",
    "text": "#17212b",
    "muted": "#718096",
}


def setup_font() -> None:
    if hasattr(font_manager.fontManager, "addfont"):
        font_manager.fontManager.addfont(FONT)
    plt.rcParams["font.family"] = "sans-serif"
    plt.rcParams["font.sans-serif"] = ["DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def clean(ax, title: str, tone: str = "text") -> None:
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.set_aspect("equal")
    ax.set_facecolor(C["bg"])
    ax.set_xticks([])
    ax.set_yticks([])
    for spine in ax.spines.values():
        spine.set_visible(False)
    for x in np.arange(0, 10.01, 1.0):
        ax.plot([x, x], [0, 6], color=C["grid"], lw=0.45, alpha=0.6, zorder=0)
    for y in np.arange(0, 6.01, 1.0):
        ax.plot([0, 10], [y, y], color=C["grid"], lw=0.45, alpha=0.6, zorder=0)
    ax.text(0.35, 5.62, title, fontsize=15, weight="bold", color=C[tone], va="top", fontproperties=FONT_PROP)


def save(fig: plt.Figure, path: Path) -> None:
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def wall(ax, rects) -> None:
    for x, y, w, h in rects:
        ax.add_patch(Rectangle((x, y), w, h, facecolor=C["wall"], edgecolor="none", zorder=5))


def path(ax, pts, color, lw=4, ls="-", alpha=1.0, z=10) -> None:
    xs, ys = zip(*pts)
    ax.plot(xs, ys, color=color, lw=lw, ls=ls, solid_capstyle="round", solid_joinstyle="round", alpha=alpha, zorder=z)


def arrow(ax, p0, p1, color, lw=3, alpha=1.0, z=12) -> None:
    ax.add_patch(
        FancyArrowPatch(
            p0,
            p1,
            arrowstyle="-|>",
            mutation_scale=18,
            lw=lw,
            color=color,
            alpha=alpha,
            shrinkA=0,
            shrinkB=0,
            zorder=z,
        )
    )


def start_goal(ax, start=(1.0, 1.0), goal=(9.0, 5.0)) -> None:
    ax.scatter(*start, s=140, color="#0ca33a", edgecolors="white", linewidths=1.6, zorder=18)
    ax.scatter(*goal, s=210, marker="*", color="red", edgecolors="white", linewidths=1.0, zorder=18)


def mark(ax, xy, ok=True, size=24) -> None:
    text = "✓" if ok else "×"
    color = C["good"] if ok else C["bad"]
    ax.text(xy[0], xy[1], text, fontsize=size, weight="bold", color=color, ha="center", va="center", zorder=30, fontproperties=FONT_PROP)


def short(ax, xy, text, color=None, size=12, ha="center") -> None:
    ax.text(xy[0], xy[1], text, fontsize=size, color=color or C["text"], ha=ha, va="center", zorder=30, fontproperties=FONT_PROP)


def robot(ax, center, yaw=0.0, scale=1.0, alpha=1.0, color="white") -> None:
    cx, cy = center
    length = 1.05 * scale
    width = 0.55 * scale
    base = np.array(
        [
            [length / 2, width / 2],
            [length / 2, -width / 2],
            [-length / 2, -width / 2],
            [-length / 2, width / 2],
        ]
    )
    rot = np.array([[math.cos(yaw), -math.sin(yaw)], [math.sin(yaw), math.cos(yaw)]])
    body = base @ rot.T + np.array([cx, cy])
    ax.add_patch(Polygon(body, closed=True, facecolor=color, edgecolor="#263238", lw=2.1, alpha=alpha, zorder=16))
    nose = np.array([[length * 0.56, 0], [length * 0.25, width * 0.22], [length * 0.25, -width * 0.22]]) @ rot.T + np.array([cx, cy])
    ax.add_patch(Polygon(nose, closed=True, facecolor=C["orange"], edgecolor="none", alpha=alpha, zorder=17))


def figure_layout():
    fig, axes = plt.subplots(1, 3, figsize=(14.2, 4.8))
    plt.subplots_adjust(wspace=0.055)
    return fig, axes


def draw_fig_01(out: Path) -> None:
    fig, ax = figure_layout()
    obstacles = [(3.0, 0.8, 0.45, 3.1), (6.0, 2.4, 1.9, 0.42), (6.2, 3.3, 0.42, 1.4)]
    titles = [("原问题：折线多", "bad"), ("改进：两头同时找", "blue"), ("效果：路径更顺", "good")]
    for a, (t, tone) in zip(ax, titles):
        clean(a, t, tone)
        wall(a, obstacles)
        start_goal(a)
    for p in [(1.0, 1.0), (1.8, 1.5), (2.4, 2.0), (2.55, 4.3), (4.0, 4.55), (5.7, 3.1), (7.0, 3.1), (8.2, 4.2), (9.0, 5.0)]:
        ax[0].add_patch(Circle(p, 0.22, facecolor=C["yellow"], edgecolor=C["orange"], lw=1, alpha=0.45, zorder=4))
    path(ax[0], [(1, 1), (1.8, 1.5), (2.4, 2.0), (2.55, 4.3), (4.0, 4.55), (5.7, 3.1), (7.0, 3.1), (8.2, 4.2), (9, 5)], C["bad"], lw=3.6)
    path(ax[1], [(1, 1), (2.2, 2.0), (2.4, 4.45), (4.0, 4.75)], C["blue"], lw=3.3)
    path(ax[1], [(9, 5), (8.1, 4.55), (6.85, 4.55), (5.3, 4.3), (4.0, 4.75)], C["blue"], lw=3.3)
    short(ax[1], (4.4, 3.5), "相遇", C["blue"], 13)
    mark(ax[1], (4.0, 4.75), True, 22)
    path(ax[2], [(1, 1), (2.2, 2.0), (2.4, 4.45), (4.0, 4.75), (5.5, 4.35), (6.85, 4.55), (8.1, 4.55), (9, 5)], C["good"], lw=4.4)
    short(ax[2], (5.5, 1.0), "少搜索  少折返", C["good"], 13)
    save(fig, out)


def draw_fig_02(out: Path) -> None:
    fig, ax = figure_layout()
    wall_rects = [(3.4, 0.7, 0.55, 3.75), (6.7, 0.7, 0.55, 3.75)]
    for a, title, tone in zip(ax, ["原问题：离墙太近", "改进：留出宽度", "效果：风险降低"], ["bad", "blue", "good"]):
        clean(a, title, tone)
        wall(a, wall_rects)
        start_goal(a, (1.0, 0.9), (9.0, 4.75))
    path(ax[0], [(1, 0.9), (2.7, 2.0), (3.15, 3.2)], C["bad"], lw=3.2, ls="--")
    robot(ax[0], (3.05, 3.1), 0.62, scale=0.95)
    mark(ax[0], (3.9, 3.1), False)
    for x, y, w, h in wall_rects:
        ax[1].add_patch(Rectangle((x - 0.75, y - 0.45), w + 1.5, h + 0.9, facecolor=C["soft"], edgecolor="none", alpha=0.55, zorder=1))
        ax[1].add_patch(Rectangle((x - 0.35, y - 0.22), w + 0.7, h + 0.44, facecolor=C["safe"], edgecolor="none", alpha=0.75, zorder=2))
    wall(ax[1], wall_rects)
    short(ax[1], (5.2, 4.55), "安全区", C["blue"], 13)
    path(ax[2], [(1, 0.9), (2.0, 1.75), (2.2, 4.85), (5.0, 4.85), (7.7, 4.75), (9, 4.75)], C["good"], lw=4.2)
    robot(ax[2], (5.0, 4.85), 0.0, scale=0.9)
    mark(ax[2], (8.4, 3.9), True)
    save(fig, out)


def draw_fig_03(out: Path) -> None:
    fig, ax = figure_layout()
    walls = [(3.1, 0.8, 0.48, 3.4), (3.1, 0.8, 3.0, 0.48)]
    for a, title, tone in zip(ax, ["原问题：一步跨过墙角", "改进：整段都检查", "效果：绕开墙角"], ["bad", "blue", "good"]):
        clean(a, title, tone)
        wall(a, walls)
        start_goal(a, (1.0, 1.0), (8.8, 5.0))
    path(ax[0], [(1.0, 1.0), (4.5, 2.9)], C["bad"], lw=3.5, ls="--")
    short(ax[0], (4.4, 2.2), "只看终点", C["bad"], 12)
    mark(ax[0], (3.25, 1.0), False)
    path(ax[1], [(1.0, 1.0), (1.8, 1.4), (2.6, 1.9), (3.0, 2.5)], C["blue"], lw=3.4)
    for p in [(1.8, 1.4), (2.6, 1.9), (3.0, 2.5)]:
        ax[1].scatter(*p, s=68, facecolor="white", edgecolor=C["blue"], lw=2.0, zorder=15)
    short(ax[1], (4.9, 4.9), "逐格验线", C["blue"], 13)
    path(ax[2], [(1, 1), (1.6, 4.75), (4.1, 5.15), (8.8, 5.0)], C["good"], lw=4.2)
    mark(ax[2], (5.0, 4.25), True)
    save(fig, out)


def draw_fig_04(out: Path) -> None:
    fig, ax = figure_layout()
    walls = [(4.2, 2.15, 0.55, 2.25), (6.45, 0.85, 0.55, 1.75)]
    for a, title, tone in zip(ax, ["原问题：频繁转身", "改进：转弯要付代价", "效果：动作更自然"], ["bad", "blue", "good"]):
        clean(a, title, tone)
        wall(a, walls)
        start_goal(a)
    zig = [
        (1, 1),
        (1.55, 1.35),
        (2.05, 1.20),
        (2.65, 1.70),
        (3.20, 1.95),
        (3.65, 4.72),
        (4.90, 4.88),
        (5.35, 4.60),
        (6.15, 4.85),
        (7.10, 4.78),
        (8.0, 4.95),
        (9, 5),
    ]
    smooth = [(1, 1), (2.65, 1.70), (3.65, 4.72), (5.35, 4.72), (7.25, 4.82), (9, 5)]
    path(ax[0], zig, C["bad"], lw=3.5)
    for p in zig[1:-1]:
        ax[0].scatter(*p, s=34, color=C["orange"], zorder=15)
    ax[1].add_patch(Circle((3.65, 4.72), 0.55, edgecolor=C["blue"], facecolor="none", lw=2.4, zorder=14))
    arrow(ax[1], (2.65, 1.70), (3.55, 4.25), C["blue"])
    arrow(ax[1], (3.95, 4.72), (5.25, 4.72), C["blue"])
    short(ax[1], (5.7, 3.5), "少转弯少扣分", C["blue"], 12)
    path(ax[2], smooth, C["good"], lw=4.3)
    robot(ax[2], (5.35, 4.72), 0.05, scale=0.9)
    mark(ax[2], (7.3, 4.55), True)
    save(fig, out)


def draw_fig_05(out: Path) -> None:
    fig, ax = plt.subplots(figsize=(12.6, 5.6))
    clean(ax, "同一地图上的三条轨迹对比", "text")
    walls = [(3.45, 1.0, 0.50, 3.25), (6.85, 1.7, 0.50, 2.60)]
    wall(ax, walls)
    start_goal(ax, (0.9, 0.8), (9.1, 5.0))
    raw = [
        (0.9, 0.8),
        (1.25, 1.15),
        (1.60, 1.72),
        (1.98, 2.15),
        (2.42, 2.78),
        (2.88, 3.55),
        (3.20, 4.55),
        (4.45, 4.85),
        (5.55, 4.65),
        (6.35, 4.88),
        (7.70, 4.78),
        (9.1, 5.0),
    ]
    key = [(0.9, 0.8), (3.20, 4.55), (5.55, 4.65), (9.1, 5.0)]
    smooth = [(0.9, 0.8), (1.45, 1.55), (2.05, 2.65), (3.20, 4.45), (4.65, 4.72), (6.20, 4.70), (7.80, 4.78), (9.1, 5.0)]
    path(ax, raw, C["bad"], lw=2.6, alpha=0.72, z=9)
    ax.scatter(*zip(*raw), s=34, color=C["bad"], alpha=0.78, zorder=13)
    path(ax, key, C["blue"], lw=3.2, ls="--", alpha=0.88, z=10)
    ax.scatter(*zip(*key), s=92, facecolor="white", edgecolor=C["blue"], lw=2.3, zorder=14)
    path(ax, smooth, C["good"], lw=4.4, alpha=0.95, z=12)
    for p in smooth[1:-1]:
        ax.scatter(*p, s=54, facecolor="white", edgecolor=C["good"], lw=1.8, zorder=15)
    short(ax, (5.2, 0.75), "红：原始碎点", C["bad"], 12)
    short(ax, (6.9, 0.75), "蓝：关键点", C["blue"], 12)
    short(ax, (8.35, 0.75), "绿：平滑轨迹", C["good"], 12)
    save(fig, out)


def draw_fig_06(out: Path) -> None:
    fig, ax = figure_layout()
    walls = [(5.0, 1.25, 0.60, 3.1)]
    for a, title, tone in zip(ax, ["原问题：只能绕大弯", "改进：前进/侧移/转向", "效果：近处灵活避障"], ["bad", "blue", "good"]):
        clean(a, title, tone)
        wall(a, walls)
        robot(a, (2.0, 3.0), 0.0, scale=1.0)
    path(ax[0], [(2.55, 3.0), (3.15, 4.6), (4.6, 5.0), (6.9, 5.0), (8.7, 3.9)], C["bad"], lw=3.7, ls="--")
    short(ax[0], (6.7, 5.25), "绕远", C["bad"], 12)
    mark(ax[0], (4.75, 3.0), False)
    arrow(ax[1], (2.6, 3.0), (4.0, 3.0), C["blue"])
    arrow(ax[1], (2.0, 3.25), (2.0, 4.45), C["blue"])
    arrow(ax[1], (2.0, 2.75), (2.0, 1.55), C["blue"])
    ax[1].add_patch(Circle((3.1, 2.0), 0.42, edgecolor=C["blue"], facecolor="none", lw=2.4, zorder=14))
    arrow(ax[1], (2.85, 1.75), (3.28, 2.25), C["blue"], lw=2.2)
    short(ax[1], (4.45, 3.25), "前进", C["blue"], 12)
    short(ax[1], (2.55, 4.35), "侧移", C["blue"], 12)
    short(ax[1], (3.95, 1.70), "转向", C["blue"], 12)
    path(ax[2], [(2.55, 3.0), (3.2, 2.15), (4.55, 1.35), (6.5, 1.45), (8.7, 2.9)], C["good"], lw=4.3)
    arrow(ax[2], (2.1, 2.9), (2.1, 2.0), C["orange"], lw=2.4)
    short(ax[2], (6.7, 1.05), "侧一步再前进", C["good"], 12)
    mark(ax[2], (8.0, 4.65), True)
    save(fig, out)


def draw_fig_07(out: Path) -> None:
    fig, ax = figure_layout()
    walls = [(5.0, 1.0, 0.45, 1.65), (5.0, 3.55, 0.45, 1.65)]
    for a, title, tone in zip(ax, ["原问题：横着走", "改进：先把头转正", "效果：头朝前通过"], ["bad", "blue", "good"]):
        clean(a, title, tone)
        wall(a, walls)
        path(a, [(1, 3), (9, 3)], C["blue"], lw=2.5, ls="--", alpha=0.55, z=3)
    robot(ax[0], (2.25, 3.0), math.pi / 2, scale=1.1)
    for x in [2.25, 3.05, 3.85]:
        robot(ax[0], (x, 3.0), math.pi / 2, scale=0.92, alpha=0.28)
    path(ax[0], [(2.25, 3.0), (4.65, 3.0)], C["bad"], lw=4.2)
    arrow(ax[0], (2.65, 2.55), (4.55, 2.55), C["bad"], lw=2.6)
    short(ax[0], (3.55, 4.65), "头朝上，却横移", C["bad"], 12)
    mark(ax[0], (4.35, 4.10), False)
    robot(ax[1], (2.6, 2.25), math.pi / 2, alpha=0.45)
    arrow(ax[1], (2.6, 2.25), (2.6, 3.0), C["orange"], lw=2.8)
    robot(ax[1], (3.4, 3.0), 0.0)
    short(ax[1], (5.6, 4.6), "先对齐", C["blue"], 14)
    robot(ax[2], (2.6, 3.0), 0.0)
    path(ax[2], [(3.1, 3.0), (4.7, 3.0), (6.1, 3.0), (9.0, 3.0)], C["good"], lw=4.4)
    mark(ax[2], (8.0, 4.45), True)
    save(fig, out)


def draw_fig_08(out: Path) -> None:
    fig, ax = figure_layout()
    walls = [(3.15, 1.0, 0.55, 3.35), (7.0, 1.0, 0.55, 3.45)]
    for a, title, tone in zip(ax, ["原问题：目标来回跳", "改进：设置参考走廊", "效果：只向前推进"], ["bad", "blue", "good"]):
        clean(a, title, tone)
        wall(a, walls)
        start_goal(a, (1.0, 0.8), (9.0, 5.1))
    route = [(1.0, 0.8), (2.25, 4.75), (4.75, 5.0), (6.25, 4.82), (8.0, 4.72), (9.0, 5.1)]
    path(ax[0], route, C["blue"], lw=2.5, ls="--", alpha=0.45)
    robot(ax[0], (1.55, 2.1), 1.22, scale=0.82)
    arrow(ax[0], (1.8, 2.3), (2.45, 4.0), C["bad"], lw=2.8)
    arrow(ax[0], (1.8, 2.3), (1.15, 1.0), C["bad"], lw=2.8)
    ax[0].add_patch(Circle((1.8, 2.3), 0.75, edgecolor=C["bad"], facecolor="none", lw=2.3, zorder=15))
    short(ax[0], (5.2, 1.15), "原地犹豫", C["bad"], 13)
    corridor = route
    path(ax[1], corridor, C["blue"], lw=12, alpha=0.18, z=2)
    path(ax[1], corridor, C["blue"], lw=3.2, ls="--", alpha=0.9, z=10)
    for i, p in enumerate(corridor[1:-1], 1):
        ax[1].scatter(*p, s=80, facecolor="white", edgecolor=C["blue"], lw=2.0, zorder=15)
        short(ax[1], (p[0], p[1] + 0.45), str(i), C["blue"], 11)
    short(ax[1], (5.25, 1.0), "DWA只在走廊内微调", C["blue"], 12)
    path(ax[2], [(1.0, 0.8), (1.55, 2.15), (2.25, 4.58), (4.0, 4.78), (5.6, 4.63), (6.35, 4.45)], C["good"], lw=4.4)
    robot(ax[2], (1.55, 2.15), 1.18, scale=0.86)
    arrow(ax[2], (4.1, 1.2), (5.6, 2.2), C["orange"], lw=2.8)
    short(ax[2], (5.55, 1.0), "目标序号不后退", C["orange"], 12)
    mark(ax[2], (8.0, 4.7), True)
    save(fig, out)


def draw_fig_09(out: Path) -> None:
    fig, ax = figure_layout()
    for a, title, tone in zip(ax, ["原问题：越跑越慢", "改进：Recovery有条件", "效果：不停顿前进"], ["bad", "blue", "good"]):
        clean(a, title, tone)
    xs = np.linspace(0.8, 9.0, 80)
    ax[0].plot(xs, 4.9 - 2.2 * (1 - np.exp(-0.28 * (xs - 0.8))), color=C["bad"], lw=4, ls="--", zorder=10)
    short(ax[0], (6.7, 3.1), "速度衰减", C["bad"], 13)
    mark(ax[0], (8.1, 2.7), False)
    boxes = [(0.8, "DWA"), (3.5, "有路?"), (6.2, "有进展?")]
    for x, txt in boxes:
        ax[1].add_patch(Rectangle((x, 2.65), 1.65, 0.85, facecolor="white", edgecolor=C["blue"], lw=2.2, zorder=12))
        short(ax[1], (x + 0.82, 3.08), txt, C["text"], 12)
    arrow(ax[1], (2.45, 3.08), (3.35, 3.08), C["blue"], lw=2.4)
    arrow(ax[1], (5.15, 3.08), (6.05, 3.08), C["blue"], lw=2.4)
    ax[1].add_patch(Rectangle((3.2, 1.05), 3.9, 0.82, facecolor="#fff4d6", edgecolor=C["orange"], lw=2.0, zorder=12))
    short(ax[1], (5.15, 1.46), "失败才脱困", "#9a5a00", 12)
    xs = np.linspace(0.8, 9.0, 80)
    ax[2].plot(xs, 4.2 + 0.08 * np.sin(xs * 1.5), color=C["good"], lw=4, zorder=10)
    robot(ax[2], (2.1, 4.2), 0.05, scale=0.85)
    arrow(ax[2], (3.0, 4.2), (8.4, 4.2), C["good"], lw=3.3)
    mark(ax[2], (8.4, 4.85), True)
    save(fig, out)


FIGURES = [
    ("1.1 基础A星与参考论文改进A星对比", "图1-1 基础A星与参考论文改进A星通俗原理图", draw_fig_01),
    ("2.1 足迹膨胀与安全距离软代价", "图2-1 足迹膨胀与安全距离软代价通俗原理图", draw_fig_02),
    ("2.2 线段安全检测", "图2-2 线段安全检测通俗原理图", draw_fig_03),
    ("2.3 转角惩罚", "图2-3 转角惩罚通俗原理图", draw_fig_04),
    ("2.4 安全剪枝、圆角平滑与重采样", "图2-4 安全剪枝、圆角平滑与重采样通俗原理图", draw_fig_05),
    ("2.5 DWA三维速度空间", "图2-5 DWA三维速度空间通俗原理图", draw_fig_06),
    ("2.6 前向姿态偏好与DWA路径头部对齐", "图2-6 前向姿态偏好与路径头部对齐通俗原理图", draw_fig_07),
    ("2.7 A星参考走廊与局部目标单调推进", "图2-7 A星参考走廊与局部目标单调推进通俗原理图", draw_fig_08),
    ("2.8 速度保持与受限Recovery", "图2-8 速度保持与受限Recovery通俗原理图", draw_fig_09),
]


def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def insert_figure_after_heading(doc: Document, heading: str, image_path: Path, caption: str) -> bool:
    for para in doc.paragraphs:
        if para.text.strip() == heading:
            pic_para = insert_paragraph_after(para)
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_para.add_run()
            run.add_picture(str(image_path), width=Inches(6.15))

            cap_para = insert_paragraph_after(pic_para, caption)
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap_para.runs:
                run.font.size = Pt(10.5)
                run.font.name = "宋体"
            return True
    return False


def main() -> None:
    setup_font()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for old in OUT_DIR.glob("*.png"):
        old.unlink()

    generated = []
    for idx, (_, caption, drawer) in enumerate(FIGURES, start=1):
        image_path = OUT_DIR / f"{idx:02d}_{caption}.png"
        drawer(image_path)
        generated.append(image_path)

    doc = Document(str(REPORT))
    missing = []
    for (heading, caption, _), image_path in zip(FIGURES, generated):
        if not insert_figure_after_heading(doc, heading, image_path, caption):
            missing.append(heading)
    if missing:
        raise RuntimeError("未找到以下标题，未能插图：" + "；".join(missing))

    doc.save(str(OUT_DOCX))
    print(OUT_DOCX)
    print(OUT_DIR)


if __name__ == "__main__":
    main()
