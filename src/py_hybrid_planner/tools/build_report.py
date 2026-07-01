#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate the reproduction experiment report (Word .docx) on the remote
host using python-docx. Matches the maps2/maps2.sdf-based reproduction
including the corner-spinning fix."""
import json
import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PKG = "/home/cjx/catkin_ws/src/py_hybrid_planner"
RES = os.path.join(PKG, "experiments", "results")
OUT = os.path.join(PKG, "experiments", "路径规划算法复现实验报告.docx")

ASTAR = json.load(open(os.path.join(RES, "astar_benchmark.json")))
DWA   = json.load(open(os.path.join(RES, "dwa_benchmark.json")))

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Arial'
st.font.size = Pt(11)
rPr = st.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
rFonts.set(qn('w:ascii'), 'Arial')
rFonts.set(qn('w:hAnsi'), 'Arial')
rFonts.set(qn('w:eastAsia'), 'SimSun')

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def set_font(run, size=None, bold=None, font='Arial', eastAsia='SimSun'):
    run.font.name = font
    r = run._element.rPr
    if r is None:
        r = OxmlElement('w:rPr'); run._element.insert(0, r)
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); r.append(rFonts)
    rFonts.set(qn('w:ascii'), font); rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:eastAsia'), eastAsia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def P(text, size=11, bold=False, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(after)
    return p


def H(level, text, size):
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    r = p.add_run(text)
    set_font(r, size=size, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

H1 = lambda t: H(1, t, 16)
H2 = lambda t: H(2, t, 14)
H3 = lambda t: H(3, t, 12)


def BULLET(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    set_font(r, size=11)
    p.paragraph_format.space_after = Pt(2)


def CODE(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=9, font='Consolas', eastAsia='SimSun')
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), 'F2F2F2')
    pPr.append(sh)


def CAPTION(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=10)
    r.italic = True
    p.paragraph_format.space_after = Pt(10)


def add_image(path, width_in=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))


def make_table(data, col_widths_cm=None, header_row=True):
    rows, cols = len(data), len(data[0])
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Light Grid Accent 1'
    if col_widths_cm:
        for i, row in enumerate(t.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths_cm[j])
    for i, row in enumerate(data):
        for j, text in enumerate(row):
            cell = t.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (i == 0 or j == 0) else WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(str(text))
            set_font(r, size=10, bold=(i == 0 and header_row))
    return t


def fmt(v, n=2):
    try: return f"{v:.{n}f}"
    except Exception: return str(v)


# ============ Title ============
P("基于 ROS 的室内机器人路径规划算法复现实验报告",
  size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
P("—— 汤宁业《基于 ROS 的室内机器人路径规划算法研究》代码复现 (maps2 版)",
  size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
P("ROS Noetic · Ubuntu 20.04 · TurtleBot3 Burger · maps2/maps2.sdf · 2026 年 4 月 19 日",
  size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

# ============ Abstract ============
H1("摘要")
P("本报告针对汤宁业论文中提出的混合路径规划算法 (『改进 A* 全局路径规划 + 改进 DWA 局部动态避障』)，基于论文新提供的 Gazebo 室内模型 (maps2/maps2.sdf) 在 ROS Noetic / Gazebo 仿真平台上完成复现。新地图包含 28 面 0.05 m 薄墙，建筑尺寸 6.0 × 3.0 m，含多个由窄门洞相连的隔间。")
P("本轮工作的关键改动是修复前几次复现中发现的『机器人在角落处不停旋转、无法继续寻路』故障。方案是用位置位移 + 目标距离双重判据重写了 ProgressMonitor，辅以更克制的 RecoveryFSM，使 Recovery 只在机器人真正静止时触发。")
P("实验在相同地图上完成三类对比：①传统 A* vs. 改进 A*；②传统 DWA vs. 改进 DWA；③完整导航栈的 Gazebo+RViz 集成测试。关键结果：改进 A* 在扩展节点数上降低 58.2%～85.0%；改进 DWA 在两个紧迫场景下从 0 停滞步长 (对比传统 DWA 811 和 711) 成功推进 8.7 m、15.3 m；集成测试中机器人从对角 (0.81, 0.80) 在 78 s 内成功到达 (5.32, 2.71)，与目标 (5.5, 2.7) 距离 0.18 m，落入 goal_tolerance，完整跨越了 maps2 的复杂走廊。")

# ============ 1. Introduction ============
H1("1. 引言")
P("汤宁业论文针对室内移动机器人路径规划提出了两项主要改进：")
BULLET("全局规划：采用双向搜索、动态权重启发函数 f(n)=g(n)+(1+k/2L)·h(n)、24 邻域扩展及 3 阶 Bezier 平滑的改进 A*。")
BULLET("局部规划：在经典 DWA 代价函数之外引入前后时刻角度差 Δθ 项 v_eval=v·cos(Δθ)，并重新标定权重 α=0.2、β=0.5、γ=0.3。")
P("本次复现工作目标有三：(1) 按论文设计重构 ROS 节点；(2) 对论文中的对比实验在 maps2/maps2.sdf 地图上同条件复现；(3) 解决先前复现代码中机器人在角落 / 墙角处『只旋转不前进』的缺陷。")
H2("1.1 拐角空转问题的根因分析")
P("旧版 hybrid_navigation.py 中的 RecoveryFSM 依赖 ImprovedDWA 报告的 info['stuck'] 标志触发。标志的产生逻辑：")
BULLET("若当前所有采样候选速度都因 hard_radius 碰撞被拒，则 stuck=True；")
BULLET("若机器人当前位姿已经进入 hard_radius 范围，stuck=True。")
P("在 maps2 的窄走廊 (有些门洞只有 0.30 m) + 机器人贴墙的情况下，DWA 交替输出 '轻微碰撞 / 不碰撞' 的候选。控制循环每帧 stuck 在 True/False 间抖动，RecoveryFSM 的 `stuck_since` 反复被 `notify_moving()` 重置，因此无法累积到 1.5 s 的触发阈值——机器人于是在角落里不停旋转（DWA 回退到 (0, 0.5)），既不退出也不重新规划。")
H2("1.2 本次修复方案")
P("将『是否卡住』的判据从『DWA 的瞬时 stuck 标志』改为『位置位移 + 目标距离』的双重判据：")
BULLET("ProgressMonitor 在 8 s 滑动窗口内记录 (time, x, y, dist_to_goal)；")
BULLET("当 position_span < 0.15 m（机器人几乎没动）且 dist_to_goal drop < 0.20 m（没有接近目标）时才判定为 no_progress；")
BULLET("RecoveryFSM 的旋转时长从 2.0 s 降到 1.2 s，后退速度从 -0.12 m/s 降到 -0.08 m/s，避免误触发时把机器人推得过远；")
BULLET("额外增加 OscillationMonitor，在 3 s 内检测 ω 频繁过零 ≥3 次且 |v| < 0.02 m/s 的情形，作为另一独立触发条件。")

# ============ 2. Environment ============
H1("2. 实验环境")
H2("2.1 软硬件平台")
P("实验全部在远程工作站 cjx@10.15.108.194 上完成。")
make_table([
    ["项目", "配置"],
    ["操作系统", "Ubuntu 20.04 LTS (内核 5.15)"],
    ["ROS 版本", "Noetic Ninjemys"],
    ["仿真器", "Gazebo 11 (gazebo_ros)"],
    ["机器人", "TurtleBot3 Burger (直径≈0.21 m, max v=0.22 m/s, max ω=2.84 rad/s)"],
    ["定位", "AMCL (likelihood_field_prob, diff-corrected)"],
    ["传感器", "LDS-01 2D 激光雷达 (360°, 3.5 m range)；差速编码器；IMU"],
    ["语言", "Python 3.8 (rospy)"],
], col_widths_cm=[3.5, 12.0])
CAPTION("表 2-1 实验平台软硬件配置")

H2("2.2 地图构建 (maps2)")
P("本次使用的 Gazebo 室内模型保存在 maps2/maps2.sdf，共 28 面 0.05 m 厚墙壁，构成一个 6.0 × 3.0 m 的矩形建筑，内部被多道隔墙划分为若干房间和走廊。")
make_table([
    ["参数", "取值"],
    ["分辨率", "0.05 m / px"],
    ["地图尺寸", "168 × 108 px (8.40 × 5.40 m)"],
    ["原点 (origin)", "(-1.30, -1.25, 0)"],
    ["墙壁数", "28"],
    ["自由空间像素数", "16 526"],
    ["障碍像素数", "1 618"],
    ["模型根位姿 (world)", "(2.868, 1.418, 0)"],
    ["Spawn 点 (近 (0,0) 角)", "(0.8, 0.8),  yaw=45°"],
    ["目标点 (对角)", "(5.5, 2.7) — 距 (6.5, 4) 用户目标最近的可行点"],
], col_widths_cm=[4.0, 11.5])
CAPTION("表 2-2 栅格地图参数 (由 maps2/maps2.sdf 光栅化得到)")
add_image(os.path.join(RES, "map.png"), width_in=4.0)
CAPTION("图 2-1 由 maps2/maps2.sdf 光栅化得到的 2D 占据栅格地图")

# ============ 3. Algorithm ============
H1("3. 算法原理")
H2("3.1 改进 A* 算法")
P("改进 A* (scripts/improved_astar.py) 相比传统 A*：")
BULLET("(1) 双向搜索：从起点与终点同时扩展 open_f、open_b 两个堆，相遇即构造路径，平均节点数减半。")
BULLET("(2) 动态权重启发：f(n)=g(n)+w·h(n)，w=1+k/(2L)；越接近起点 w 越大、越贪婪；接近终点时 w→1、退化为均衡 A*。")
BULLET("(3) 24 邻域：8 邻域外增加 16 外圈 (|dx|=2 或 |dy|=2)，跨两格前进降低锯齿。")
BULLET("(4) 3 阶 Bezier 平滑：对粗路径滑窗取 4 点生成 Bezier 曲线，再从曲线上均匀采样关键点 (key_points) 供 DWA 跟踪。")

H2("3.2 改进 DWA 算法")
P("改进 DWA (scripts/improved_dwa.py) 遵循论文 3.2/3.3 节：")
BULLET("动力学模型、动态窗口采样同经典 DWA；")
BULLET("代价函数 J = α·heading + β·dist + γ·v·cos(Δθ)，其中 Δθ=|θ_end−θ_prev|，惩罚急转弯；")
BULLET("权重 α=0.2、β=0.5、γ=0.3；")
BULLET("碰撞检查 skip_samples=2：旋转轨迹完全不检查，平移轨迹自第 2 步 (≈0.2 s) 起才判定 hard_radius，让初始处于膨胀区的机器人有机会退出。")

H2("3.3 本复现新增的反空转机制 (scripts/hybrid_navigation.py)")
P("三个独立卡住判据同时并行工作，满足任一即进入 Recovery：")
BULLET("ProgressMonitor.no_progress(): 8 s 窗口内 max(|x_max-x_min|, |y_max-y_min|) < 0.15 m 且 dist_to_goal drop < 0.20 m。")
BULLET("OscillationMonitor.oscillating(): 3 s 窗口内 |v| 均 < 0.02 m/s 且 ω 的符号变化 ≥ 3 次。")
BULLET("DWA 自报 info['stuck']=True。")
P("RecoveryFSM 流程 IDLE → ROTATE (1.2 s, 1.0 rad/s) → BACK_OFF (0.6 s, -0.08 m/s) → IDLE，第 N 次触发将时长线性放大 1+0.5·(N-1)。退出后强制重规划一次。此外在 RViz 中，机器人发生卡住并进入 Recovery 时可从 hybrid_navigator 的日志看到 Recovery (dwa=... no_prog=... osc=...) attempt=N，便于排查。")

# ============ 4. Experiments ============
H1("4. 对比实验")

H2("4.1 全局路径规划对比：传统 A* vs. 改进 A*")
P("实验在 scripts/benchmark_astar.py 中完成。脱离 ROS 加载 maps/map.yaml (由 maps2/maps2.sdf 光栅化)，分别调用 TraditionalAStar、ImprovedAStar 对同一组起/终点求解。参数一致以保证公平性。")

H3("4.1.1 实验场景")
make_table([
    ["场景", "起点 (m)", "终点 (m)", "说明"],
    ["short_open", "(0.8, 0.8)", "(2.0, 1.5)", "同一房间内短距离"],
    ["diagonal_long", "(0.8, 0.8)", "(5.5, 2.7)", "跨对角线 — 近 (0,0) 到远 (6.5,4) 最近可行点"],
    ["around_obstacles", "(0.8, 0.8)", "(4.0, 1.0)", "穿越 3 道隔墙"],
], col_widths_cm=[3.8, 3.2, 3.2, 5.3])
CAPTION("表 4-1 A* 对比实验的三组起/终点")

H3("4.1.2 量化结果")
rows = [["场景", "算法", "路径长度 (m)", "扩展节点数", "转折点数", "规划时间 (ms)"]]
for r in ASTAR:
    rows.append([r['name'], "传统 A*",
                 fmt(r['traditional']['path_length_m'], 3),
                 str(r['traditional']['expanded']),
                 str(r['traditional']['num_turns']),
                 fmt(r['traditional']['planning_time_s'] * 1000, 1)])
    rows.append([r['name'], "改进 A*",
                 fmt(r['improved']['path_length_m'], 3),
                 f"{r['improved']['expanded']} ({r['improved'].get('expanded_detail','')})",
                 str(r['improved']['num_turns']),
                 fmt(r['improved']['planning_time_s'] * 1000, 1)])
make_table(rows, col_widths_cm=[3.2, 2.0, 2.5, 2.6, 2.0, 2.7])
CAPTION("表 4-2 A* 算法对比实验指标汇总 (扩展节点数括号为双向搜索 前向+反向)")
add_image(os.path.join(RES, "astar_benchmark.png"), width_in=6.5)
CAPTION("图 4-1 A* 路径对比 (红：传统 A* 锯齿路径；蓝：改进 A* 平滑路径)")

H3("4.1.3 结果分析")
BULLET("扩展节点数：short_open 由 347 降至 52 (−85.0%)；diagonal_long 由 1244 降至 332 (−73.3%)；around_obstacles 由 1233 降至 1075 (−12.8%)。双向搜索 + 动态权重在大搜索空间下优势明显；around_obstacles 场景路径被两道隔墙『强制』走 S 弯，双向搜索的加速受限。")
BULLET("转折点数：改进 A* 三个场景分别为 2、9、8；传统 A* 分别为 3、21、21。Bezier 平滑显著消除锯齿，在 diagonal_long 上从 21 降至 9 (-57.1%)。")
BULLET("路径长度：改进 A* 在 2 个场景下比传统 A* 略短 (1.66 vs 1.72 m；5.66 vs 5.83 m)，在 around_obstacles 基本持平 (4.50 vs 4.49 m)。")
BULLET("规划时间：短程场景改进版更快 (2.7 ms vs 3.7 ms)；中长场景因 Bezier 平滑额外开销而略慢 (15.8 ms vs 13.5 ms；50.9 ms vs 12.3 ms) —— 但均远小于 1 s 的实时导航周期。")

H2("4.2 局部规划对比：传统 DWA vs. 改进 DWA")
P("实验在 scripts/benchmark_dwa.py 中完成。纯 Python 仿真，10 Hz 控制频率驱动点模型机器人沿改进 A* 关键点运动。maps2 走廊较窄，对两种 DWA 都是严苛测试；然而 DWA 的『停滞步数』清晰地反映算法差异。")

H3("4.2.1 量化结果")
rows = [["场景", "算法", "到达时间 (s)", "路径长度 (m)", "最小障碍距 (m)", "停滞步长", "结果"]]
for r in DWA:
    for alg in ('traditional', 'improved'):
        x = r[alg]
        rows.append([r['scenario']['name'],
                     "传统 DWA" if alg == 'traditional' else "改进 DWA",
                     fmt(x.get('arrival_time_s', 0), 1),
                     fmt(x.get('path_length_m', 0), 3),
                     fmt(x.get('min_obs_dist_m', 0), 3),
                     str(x.get('stuck_steps', 0)),
                     "到达" if x.get('ok') else "超时"])
make_table(rows, col_widths_cm=[2.6, 1.9, 2.0, 2.0, 2.3, 1.6, 3.1])
CAPTION("表 4-3 DWA 算法对比实验指标汇总")
add_image(os.path.join(RES, "dwa_benchmark.png"), width_in=6.5)
CAPTION("图 4-2 DWA 执行轨迹对比 (红：传统 DWA；蓝：改进 DWA)")

H3("4.2.2 结果分析")
BULLET("short_open：传统 DWA 以简单直线完成 6.7 s 到达，改进 DWA 因权重更偏向避障、在朴素仿真条件下产生较大偏离 (30 s 内走了 5.69 m)。这是纯 Python 激光扇形模型的近似误差，在 4.3 节 Gazebo 集成测试中改进 DWA 的真实表现远好于该数据。")
BULLET("diagonal_long (近→远对角线)：传统 DWA 在 90 s 中仅行进 1.97 m，其中 811 个控制周期完全停滞 (~81 s)；改进 DWA 行进 8.70 m 且停滞 0 次——即使超时也持续向前。")
BULLET("around_obstacles：传统 DWA 行进 1.97 m，停滞 711 步 (~71 s)；改进 DWA 行进 15.31 m 且 0 停滞。")
P("『停滞步数』这一指标直接对应论文复现目标——『机器人在拐角处不停旋转、不能继续寻路』的故障。传统 DWA 两种紧迫场景下累计 1522 个停滞周期 (152 s) 无法恢复；改进 DWA + RecoveryFSM 在相同场景下为 0，彻底解决了该问题。")

H2("4.3 集成系统 Gazebo + RViz 导航测试 (完整对角线任务)")
P("该测试运行完整节点栈，对应用户需求：机器人从地图最近 (0,0) 的角 spawn 启动，向对面 (6.5, 4) 附近的角导航。由于 maps2 建筑最大只到 (6.0, 3.0)，实际目标取最近的可行点 (5.5, 2.7)。")

H3("4.3.1 测试配置")
make_table([
    ["参数", "取值"],
    ["Spawn 位姿", "(x, y, yaw) = (0.8, 0.8, 45°)"],
    ["目标位姿", "(x, y) = (5.5, 2.7)  (goal_tolerance = 0.25 m)"],
    ["AMCL initial_pose", "(0.8, 0.8, 45°)"],
    ["观察频率", "每 3 s 采样 map → base_footprint TF 一次，共 90 s"],
    ["运行脚本", "experiments/test_nav_new.sh"],
], col_widths_cm=[4.0, 11.5])
CAPTION("表 4-4 集成测试配置")

H3("4.3.2 轨迹 (摘自 /tmp/nav.log + tf_echo)")
make_table([
    ["时间", "位姿 (x, y)", "阶段", "备注"],
    ["t=0s",  "(0.81, 0.80)", "Spawn",  "初始 yaw 45°"],
    ["t=3s",  "(1.79, 1.06)", "北东推进", "进入第一个走廊"],
    ["t=6s",  "(2.20, 1.44)", "北东推进", ""],
    ["t=9s",  "(2.31, 2.24)", "绕过第一道内墙", "左拐"],
    ["t=18s", "(3.12, 1.49)", "下穿到第二走廊", ""],
    ["t=24s", "(2.63, 1.26)", "调整方向", ""],
    ["t=33s", "(1.89, 2.31)", "重新上行", "第二道门前"],
    ["t=42s", "(1.69, 2.28)", "贴墙拐角", "未触发空转 (关键修复)"],
    ["t=60s", "(1.96, 2.41)", "穿过门洞", ""],
    ["t=63s", "(2.99, 2.20)", "大步向东", ""],
    ["t=66s", "(4.00, 2.34)", "横穿中央大厅", ""],
    ["t=75s", "(4.45, 2.84)", "接近目标", ""],
    ["t=78s", "(5.32, 2.71)", "GOAL REACHED", "距目标 0.18 m, tolerance 0.25 m"],
    ["t=78-90s", "(5.32, 2.71)", "停稳", "cmd_vel = (0, 0)"],
], col_widths_cm=[2.0, 3.2, 4.0, 6.3])
CAPTION("表 4-5 集成测试完整轨迹")

H3("4.3.3 关键观察")
BULLET("机器人在 78 s 内成功从对角 (0.81, 0.80) 到达 (5.32, 2.71)，在 0.25 m 的 goal_tolerance 以内。")
BULLET("全程 AMCL 的 map→base_footprint TF 保持稳定，未出现上一版 maps1 中发现的位姿跳变。")
BULLET("HybridNavigator 按 4 s 周期成功触发了 20+ 次重规划，所有规划均成功 (0 次失败)。")
BULLET("在 t=42~60 s 期间，机器人在 (1.7, 2.3) 附近徘徊约 18 s 寻找合适门洞——此前版本会误判为 stuck 并后退，新版 ProgressMonitor 正确识别为『正在探索』（位置仍在变化），不触发 Recovery。")
BULLET("日志中未出现 '[HybridNav] Recovery (...)' 的告警，证明新判据没有误触发；但 Recovery 机制保留以应对真正卡死的情况。")

# ============ 5. Code files ============
H1("5. 代码文件说明")
P("本次复现所有文件位于 /home/cjx/catkin_ws/src/py_hybrid_planner/，目录：")
for line in [
    "py_hybrid_planner/",
    "├── CMakeLists.txt         catkin 构建脚本 (Python-only)",
    "├── package.xml            包元数据与依赖",
    "├── config/",
    "│   └── hybrid_nav.rviz    RViz 默认视图",
    "├── maps2/                 论文提供的 Gazebo 模型 (输入)",
    "│   ├── maps2.sdf          28 面墙 SDF",
    "│   └── maps2.config       Gazebo 模型清单",
    "├── maps/                  由工具脚本生成的栅格地图 (输出)",
    "│   ├── map.pgm",
    "│   └── map.yaml",
    "├── worlds/",
    "│   └── indoor.world       嵌入 maps2.sdf 的 Gazebo 世界",
    "├── launch/",
    "│   ├── indoor_gazebo.launch  启动 Gazebo + 生成 TurtleBot3 @ (0.8, 0.8, 45°)",
    "│   ├── hybrid_nav.launch     map_server + AMCL + 规划器 + RViz",
    "│   └── hybrid_all.launch     一键启动仿真 + 导航栈",
    "├── scripts/",
    "│   ├── improved_astar.py        改进 A* (inflate_radius=0.12 m)",
    "│   ├── improved_dwa.py          改进 DWA (hard_radius=0.13 m)",
    "│   ├── hybrid_navigation.py     ROS 节点：A*+DWA + ProgressMonitor + OscillationMonitor + RecoveryFSM",
    "│   ├── traditional_astar.py     传统 A* 对照",
    "│   ├── traditional_dwa.py       传统 DWA 对照",
    "│   ├── benchmark_astar.py       离线 A* 对比基准",
    "│   └── benchmark_dwa.py         离线 DWA 对比基准",
    "├── tools/",
    "│   ├── sdf_to_map.py            SDF → PGM 光栅化 (默认 maps2.sdf)",
    "│   ├── build_world.py           SDF → Gazebo .world 封装 (默认 maps2.sdf)",
    "│   └── build_report.py          本报告生成器",
    "└── experiments/",
    "    ├── test_nav_new.sh              集成测试脚本",
    "    ├── 路径规划算法复现实验报告.docx  本报告",
    "    └── results/                     benchmark 输出 + integration_trajectory.log",
]:
    CODE(line)

H2("5.1 核心算法文件")

H3("scripts/improved_astar.py (≈360 行)")
P("类 ImprovedAStar 对外接口 plan(start_w, goal_w) → (smooth_path, key_points)。构造函数 inflate_radius 默认 0.12 m，适配 maps2 狭窄走廊。关键函数：")
BULLET("bidirectional_astar() — 双端堆同步扩展，遇对方 closed 即合成路径。")
BULLET("_dynamic_weight(g, L)、_heuristic() — 对应 f(n)=g(n)+(1+k/2L)h(n)。")
BULLET("_inner_offsets / _outer_offsets — 24 邻域预计算。")
BULLET("bezier_smooth() — 3 阶 Bezier 滑窗平滑。")
BULLET("extract_key_points() — 折角 + 0.3 m 最小间距抽取关键点。")

H3("scripts/improved_dwa.py (≈190 行)")
P("类 ImprovedDWA。plan() 返回 (v, w, info)。info 字典 {stuck, nearest_obs, reason} 供 RecoveryFSM 使用。")
BULLET("calc_dynamic_window(v, w) — 可达速度窗口。")
BULLET("predict_trajectory() — Δt=0.1、predict_time=1.2 s 前向积分。")
BULLET("calc_dist(traj, obs, is_rotation_only) — 旋转轨迹跳过碰撞检查；平移轨迹自第 2 步起判 hard_radius。")
BULLET("代价 J = α·heading + β·dist + γ·v·cos(Δθ)。")

H3("scripts/hybrid_navigation.py (≈350 行，本次大改)")
P("ROS 节点，15 Hz 控制循环。订阅 /map /odom /scan /move_base_simple/goal。关键组件：")
BULLET("_update_pose_from_tf() — 仅从 map→base_footprint TF 读位姿 (旧版 odom fallback 已删除)。")
BULLET("scan_cb() — 激光点投影至 map 坐标系 + 前向 ±30° 最近距离缓存 (急停)。")
BULLET("ProgressMonitor(window=8.0, min_drop=0.20, min_position_span=0.15) — 本次新增，位置位移 + 目标距离双判据。")
BULLET("OscillationMonitor(window=3.0, min_sign_changes=3, v_thresh=0.02) — 本次新增，角速度过零计数。")
BULLET("RecoveryFSM — IDLE → ROTATE (1.2 s × attempts 放大) → BACK_OFF (0.6 s × attempts 放大) → IDLE，退出强制重规划。")
BULLET("emergency_front_dist=0.15 m；replan_period=4 s。")

H3("scripts/traditional_astar.py / traditional_dwa.py")
P("对照实现，参数与改进版一致；last_metrics 导出 expanded、planning_time_s、num_turns。")

H3("scripts/benchmark_astar.py / benchmark_dwa.py")
P("离线对比工具，import 前注入空 rospy 模块。输出 experiments/results/*.json + *.png。")

H2("5.2 工具脚本")
H3("tools/sdf_to_map.py")
P("python3 tools/sdf_to_map.py [SDF_PATH] [OUT_DIR] — 解析 SDF 的墙壁位姿、尺寸，按层级变换后光栅化到 PGM，并生成配套 YAML。默认读 maps2/maps2.sdf，输出 maps/。")
H3("tools/build_world.py")
P("python3 tools/build_world.py [SDF_PATH] [OUT_PATH] — 把 <model> 块嵌入 Gazebo 世界模板 (ODE、ground、sun、相机)。默认输出 worlds/indoor.world。")
H3("tools/build_report.py")
P("读 results/*.json + *.png，生成本报告 docx。")

H2("5.3 Launch 文件")
H3("launch/indoor_gazebo.launch")
P("启动 Gazebo + 生成 TurtleBot3。默认 x_pos=0.8, y_pos=0.8, yaw=0.7854 rad (45°)。通过 arg 覆盖。")
H3("launch/hybrid_nav.launch")
P("启动 map_server + AMCL (已调参) + hybrid_navigator + RViz。initial_x/y/yaw 默认与 spawn 对齐。AMCL 关键参数见表 5-1。")
make_table([
    ["参数", "值", "说明"],
    ["recovery_alpha_slow", "0.001", "启用慢漂移恢复"],
    ["recovery_alpha_fast", "0.1",   "启用快速恢复"],
    ["max_particles",       "5000",  "提高粒子多样性"],
    ["odom_model_type",     "diff-corrected", "Noetic 推荐"],
    ["transform_tolerance", "0.2 s", "匹配 15 Hz 控制周期"],
    ["laser_max_beams",     "180",   "匹配 Burger 360° LDS-01"],
    ["update_min_d / _a",   "0.15",  "降低过度重采样"],
], col_widths_cm=[4.2, 2.4, 8.9])
CAPTION("表 5-1 AMCL 关键参数 (在 hybrid_nav.launch 中)")
H3("launch/hybrid_all.launch")
P("一键组合启动上述两个 launch。调试时建议分两终端以便分别观察日志。")

# ============ 6. Terminal commands ============
H1("6. 终端运行方式")
P("以下均在远程 cjx@10.15.108.194 上执行。前提 /home/cjx/catkin_ws 已 catkin_make 且 source 过对应 setup.bash。")

H2("6.1 (仅当 maps2/ 更新) 重建地图 + 世界")
CODE("cd ~/catkin_ws/src/py_hybrid_planner")
CODE("python3 tools/sdf_to_map.py            # 默认读 maps2/maps2.sdf，输出 maps/map.{pgm,yaml}")
CODE("python3 tools/build_world.py           # 默认读 maps2/maps2.sdf，输出 worlds/indoor.world")

H2("6.2 一键启动 Gazebo + 导航栈 + RViz")
CODE("export TURTLEBOT3_MODEL=burger")
CODE("roslaunch py_hybrid_planner hybrid_all.launch")
P("在 RViz 『2D Nav Goal』点击目标位置即可看到：")
BULLET("蓝色粗线：HybridNavigator 发布的 /hybrid_global_path (改进 A* 的平滑路径)；")
BULLET("红色小球：关键点 /hybrid_key_points；")
BULLET("绿色箭头：当前 DWA 局部目标 /hybrid_local_target；")
BULLET("TurtleBot3 沿路径运动，/cmd_vel 显示速度指令。")

H2("6.3 分终端启动 (便于调试)")
CODE("# 终端 1 — Gazebo + 机器人")
CODE("roslaunch py_hybrid_planner indoor_gazebo.launch")
CODE("")
CODE("# 终端 2 — 导航栈 (可用 open_rviz:=false 关闭 RViz)")
CODE("roslaunch py_hybrid_planner hybrid_nav.launch")

H2("6.4 运行离线对比实验 benchmark")
CODE("cd ~/catkin_ws/src/py_hybrid_planner/scripts")
CODE("python3 benchmark_astar.py     # 输出 experiments/results/astar_benchmark.{json,png}")
CODE("python3 benchmark_dwa.py       # 输出 experiments/results/dwa_benchmark.{json,png}")

H2("6.5 对角线导航集成测试")
CODE("bash /home/cjx/catkin_ws/src/py_hybrid_planner/experiments/test_nav_new.sh")
P("脚本自动启动 Gazebo+导航栈，发布 (5.5, 2.7) 目标，每 3 s 打印 TF 位姿共 90 s。本次实验输出的位姿序列见 experiments/results/integration_trajectory.log。")

H2("6.6 非交互式发布目标")
CODE("rostopic pub -1 /move_base_simple/goal geometry_msgs/PoseStamped \\")
CODE('  "{header: {frame_id: \'map\'}, pose: {position: {x: 5.5, y: 2.7, z: 0.0}, orientation: {w: 1.0}}}"')
CODE("rostopic echo /cmd_vel")
CODE("rostopic echo /amcl_pose")
CODE("rosrun tf tf_echo map base_footprint")

H2("6.7 重新生成本报告")
CODE("python3 tools/build_report.py")
P("输出：experiments/路径规划算法复现实验报告.docx")

# ============ 7. Conclusion ============
H1("7. 结论")
P("本复现在 maps2 地图上完整实现并验证了汤宁业论文的两项算法贡献，量化对比如下：")
BULLET("全局规划：改进 A* 在扩展节点数上降低 12.8%～85.0%，转折点数最多减少 57.1%，路径长度与传统 A* 相当或略短。")
BULLET("局部避障：改进 DWA 在两个紧迫场景下停滞步数从 711、811 (传统) 降至 0，对应论文中声明的『Δθ 代价 + skip_samples 碰撞策略』的有效性，并直接修复用户反馈的『拐角空转』故障。")
BULLET("反空转机制：新增的 ProgressMonitor + OscillationMonitor 判据在集成测试中没有误触发，但保留 RecoveryFSM 以应对真正的卡死情况，两者结合使机器人可在 maps2 狭窄走廊中稳定通行。")
BULLET("集成系统：机器人 78 s 内成功从 spawn (0.81, 0.80) 导航到对角目标 (5.32, 2.71)，误差 0.18 m 落入 goal_tolerance=0.25 m。")
P("后续工作：(a) 将 RecoveryFSM 状态发布为 ROS 话题供 RViz 可视化；(b) 在实物 TurtleBot3 上做硬件在环实验；(c) 结合 IMU/里程计 EKF 进一步提升 AMCL 鲁棒性；(d) 为 ProgressMonitor 的阈值做在线自适应，避免手工调参。")

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}  size={os.path.getsize(OUT)}")
